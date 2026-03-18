from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Colores ITO Cloud ───────────────────────────────────────────────
AZUL_OSCURO  = (0x1A, 0x37, 0x6C)   # encabezados principales
AZUL_MEDIO   = (0x2E, 0x6D, 0xA4)   # encabezados de tabla
AZUL_CLARO   = (0xD6, 0xE4, 0xF0)   # fila par / relleno suave
GRIS_CLARO   = (0xF5, 0xF5, 0xF5)   # fila impar
BLANCO       = (0xFF, 0xFF, 0xFF)
NEGRO        = (0x00, 0x00, 0x00)
VERDE        = (0x1E, 0x7E, 0x34)   # PK
ROJO         = (0xC0, 0x39, 0x2B)   # NOT NULL
NARANJA      = (0xE6, 0x7E, 0x22)   # FK

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    # rgb puede ser RGBColor (indexable) o tupla
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), str(val.get('sz', 4)))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(*AZUL_OSCURO)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(*AZUL_MEDIO)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(*AZUL_OSCURO)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_dict(doc, table_info):
    cols = table_info['columns']
    constraints = table_info.get('constraints', [])

    # Encabezado de tabla
    add_heading(doc, f"Tabla: {table_info['name']}", level=2)
    p = doc.add_paragraph()
    p.add_run("Descripción: ").bold = True
    p.add_run(table_info['description'])
    p.paragraph_format.space_after = Pt(4)

    # Tabla de columnas
    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Anchos de columnas (cm)
    widths = [3.2, 2.8, 1.4, 1.0, 1.0, 1.6, 5.8]
    headers = ['Columna', 'Tipo', 'Nulable', 'PK', 'FK', 'Default', 'Descripción']

    # Fila encabezado
    hdr_row = tbl.rows[0]
    for i, (cell, hdr, w) in enumerate(zip(hdr_row.cells, headers, widths)):
        cell.width = Cm(w)
        set_cell_bg(cell, AZUL_MEDIO)
        p2 = cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*BLANCO)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Filas de datos
    for idx, col in enumerate(cols):
        row = tbl.add_row()
        bg = AZUL_CLARO if idx % 2 == 0 else GRIS_CLARO

        fk_str = col.get('fk') or ''
        def_str = col.get('default') or ''
        nullable_str = 'Sí' if col.get('nullable', True) else 'No'
        pk_str = '✓' if col.get('pk') else ''
        fk_show = '✓' if fk_str else ''

        values = [
            col['name'],
            col['type'],
            nullable_str,
            pk_str,
            fk_show,
            def_str,
            col.get('description', ''),
        ]

        for i, (cell, val) in enumerate(zip(row.cells, values)):
            set_cell_bg(cell, bg)
            p3 = cell.paragraphs[0]
            run = p3.add_run(str(val))
            run.font.size = Pt(8)
            # Colores especiales
            if i == 0:  # nombre columna
                if col.get('pk'):
                    run.font.color.rgb = RGBColor(*VERDE)
                    run.bold = True
                elif not col.get('nullable', True):
                    run.font.color.rgb = RGBColor(*ROJO)
                else:
                    run.font.color.rgb = RGBColor(*NEGRO)
            elif i == 4 and fk_show:  # FK check
                run.font.color.rgb = RGBColor(*NARANJA)
                run.bold = True
            else:
                run.font.color.rgb = RGBColor(*NEGRO)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Constraints
    if constraints:
        p4 = doc.add_paragraph()
        p4.paragraph_format.space_before = Pt(2)
        run = p4.add_run("Restricciones adicionales: ")
        run.bold = True
        run.font.size = Pt(9)
        for c in constraints:
            pc = doc.add_paragraph(style='List Bullet')
            pc.add_run(c).font.size = Pt(8)
            pc.paragraph_format.space_before = Pt(0)
            pc.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()  # separador

# ─── DATOS ───────────────────────────────────────────────────────────
TABLES = [
  {
    "name": "tenants",
    "description": "Empresas (tenants) registradas en la plataforma SaaS. Cada tenant representa una organización cliente independiente con su propia configuración, logo y límites de uso.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único del tenant"},
      {"name":"name","type":"VARCHAR(200)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre completo de la empresa"},
      {"name":"slug","type":"VARCHAR(100)","nullable":False,"pk":False,"fk":None,"unique":True,"default":None,"description":"Identificador URL-friendly único del tenant (ej. 'constructora-abc')"},
      {"name":"plan","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'basic'","description":"Plan de suscripción contratado"},
      {"name":"is_active","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si el tenant está activo en la plataforma"},
      {"name":"logo_url","type":"VARCHAR(500)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"URL del logotipo del tenant almacenado en MinIO"},
      {"name":"primary_color","type":"VARCHAR(20)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Color primario de marca en formato hex (ej. '#1A376C')"},
      {"name":"max_users","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"10","description":"Límite máximo de usuarios activos permitidos por plan"},
      {"name":"max_projects","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"5","description":"Límite máximo de proyectos activos permitidos por plan"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha y hora de creación del tenant"},
      {"name":"updated_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha y hora de última actualización"},
    ],
    "constraints":["chk_tenants_plan: plan IN ('basic','professional','enterprise')"]
  },
  {
    "name": "users",
    "description": "Usuarios del sistema (ASP.NET Identity extendido). Cada usuario pertenece a un tenant. Almacena credenciales, perfil, configuración móvil y campos de auditoría.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único del usuario"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant al que pertenece el usuario"},
      {"name":"user_name","type":"VARCHAR(256)","nullable":False,"pk":False,"fk":None,"unique":True,"default":None,"description":"Nombre de usuario único (Identity)"},
      {"name":"normalized_user_name","type":"VARCHAR(256)","nullable":True,"pk":False,"fk":None,"unique":True,"default":None,"description":"Nombre de usuario normalizado para búsquedas (Identity)"},
      {"name":"email","type":"VARCHAR(256)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Correo electrónico del usuario"},
      {"name":"normalized_email","type":"VARCHAR(256)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Email normalizado para búsquedas (Identity)"},
      {"name":"email_confirmed","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si el email fue confirmado"},
      {"name":"password_hash","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Hash de contraseña (ASP.NET Identity)"},
      {"name":"security_stamp","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Sello de seguridad para invalidar tokens (Identity)"},
      {"name":"concurrency_stamp","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Sello de concurrencia optimista (Identity)"},
      {"name":"phone_number","type":"VARCHAR(50)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Número de teléfono del usuario"},
      {"name":"phone_number_confirmed","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si el teléfono fue confirmado"},
      {"name":"two_factor_enabled","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si el 2FA está habilitado"},
      {"name":"lockout_end","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de fin del bloqueo de cuenta"},
      {"name":"lockout_enabled","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si el bloqueo por intentos fallidos está habilitado"},
      {"name":"access_failed_count","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"0","description":"Contador de intentos fallidos de acceso"},
      {"name":"first_name","type":"VARCHAR(100)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre del usuario"},
      {"name":"last_name","type":"VARCHAR(100)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Apellido del usuario"},
      {"name":"avatar_url","type":"VARCHAR(500)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"URL del avatar del usuario en MinIO"},
      {"name":"is_active","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si el usuario está activo"},
      {"name":"fcm_token","type":"VARCHAR(500)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Token FCM para notificaciones push en Android"},
      {"name":"last_login_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha y hora del último inicio de sesión"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación del usuario"},
      {"name":"updated_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de última actualización"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
    ],
    "constraints":[]
  },
  {
    "name": "roles",
    "description": "Roles del sistema por tenant (ASP.NET Identity extendido). Define los permisos agrupados que se asignan a usuarios.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único del rol"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant al que pertenece el rol"},
      {"name":"name","type":"VARCHAR(256)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre del rol"},
      {"name":"normalized_name","type":"VARCHAR(256)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre normalizado para búsquedas (Identity)"},
      {"name":"concurrency_stamp","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Sello de concurrencia optimista (Identity)"},
      {"name":"description","type":"VARCHAR(500)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Descripción del rol y sus permisos"},
      {"name":"is_system","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si es un rol del sistema (no editable por el tenant)"},
    ],
    "constraints":[]
  },
  {
    "name": "user_roles",
    "description": "Tabla de unión entre usuarios y roles (ASP.NET Identity). Define qué roles tiene asignado cada usuario.",
    "columns": [
      {"name":"user_id","type":"UUID","nullable":False,"pk":True,"fk":"users.id","unique":False,"default":None,"description":"Usuario al que se asigna el rol"},
      {"name":"role_id","type":"UUID","nullable":False,"pk":True,"fk":"roles.id","unique":False,"default":None,"description":"Rol asignado al usuario"},
    ],
    "constraints":["PK compuesta: (user_id, role_id)"]
  },
  {
    "name": "user_claims",
    "description": "Claims (atributos) asignados directamente a usuarios (ASP.NET Identity). Permite permisos granulares por usuario.",
    "columns": [
      {"name":"id","type":"SERIAL","nullable":False,"pk":True,"fk":None,"unique":False,"default":"SERIAL","description":"Identificador autoincremental del claim"},
      {"name":"user_id","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario al que pertenece el claim"},
      {"name":"claim_type","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tipo del claim (ej. 'permission')"},
      {"name":"claim_value","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Valor del claim (ej. 'inspections.create')"},
    ],
    "constraints":[]
  },
  {
    "name": "role_claims",
    "description": "Claims asignados a roles (ASP.NET Identity). Todos los usuarios con ese rol heredan estos claims.",
    "columns": [
      {"name":"id","type":"SERIAL","nullable":False,"pk":True,"fk":None,"unique":False,"default":"SERIAL","description":"Identificador autoincremental"},
      {"name":"role_id","type":"UUID","nullable":False,"pk":False,"fk":"roles.id","unique":False,"default":None,"description":"Rol al que pertenece el claim"},
      {"name":"claim_type","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tipo del claim"},
      {"name":"claim_value","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Valor del claim"},
    ],
    "constraints":[]
  },
  {
    "name": "user_tokens",
    "description": "Tokens de autenticación externos por usuario (ASP.NET Identity). Almacena tokens de proveedores OAuth.",
    "columns": [
      {"name":"user_id","type":"UUID","nullable":False,"pk":True,"fk":"users.id","unique":False,"default":None,"description":"Usuario propietario del token"},
      {"name":"login_provider","type":"VARCHAR(450)","nullable":False,"pk":True,"fk":None,"unique":False,"default":None,"description":"Proveedor de login (ej. 'Google')"},
      {"name":"name","type":"VARCHAR(450)","nullable":False,"pk":True,"fk":None,"unique":False,"default":None,"description":"Nombre del token"},
      {"name":"value","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Valor del token"},
    ],
    "constraints":["PK compuesta: (user_id, login_provider, name)"]
  },
  {
    "name": "user_logins",
    "description": "Inicios de sesión externos asociados a usuarios (ASP.NET Identity). Vincula cuentas OAuth con usuarios locales.",
    "columns": [
      {"name":"login_provider","type":"VARCHAR(450)","nullable":False,"pk":True,"fk":None,"unique":False,"default":None,"description":"Proveedor de login externo"},
      {"name":"provider_key","type":"VARCHAR(450)","nullable":False,"pk":True,"fk":None,"unique":False,"default":None,"description":"Clave del usuario en el proveedor externo"},
      {"name":"provider_display_name","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre amigable del proveedor"},
      {"name":"user_id","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario local vinculado"},
    ],
    "constraints":["PK compuesta: (login_provider, provider_key)"]
  },
  {
    "name": "companies",
    "description": "Empresas mandantes/clientes que contratan inspecciones. Cada empresa pertenece a un tenant y puede tener múltiples proyectos.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la empresa"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"name","type":"VARCHAR(300)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Razón social o nombre de la empresa"},
      {"name":"rut","type":"VARCHAR(20)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"RUT o identificador tributario de la empresa"},
      {"name":"address","type":"VARCHAR(500)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Dirección fiscal de la empresa"},
      {"name":"phone","type":"VARCHAR(50)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Teléfono de contacto"},
      {"name":"email","type":"VARCHAR(256)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Email de contacto"},
      {"name":"contact_name","type":"VARCHAR(200)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre de la persona de contacto"},
      {"name":"logo_url","type":"VARCHAR(500)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"URL del logo de la empresa en MinIO"},
      {"name":"is_active","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si la empresa está activa"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"updated_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de última actualización"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó el registro"},
      {"name":"updated_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que actualizó el registro"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"deleted_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que realizó el borrado lógico"},
    ],
    "constraints":[]
  },
  {
    "name": "projects",
    "description": "Proyectos de construcción/obra que serán inspeccionados. Contiene la información general del proyecto, empresa mandante, ITO responsable, fechas y estado.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único del proyecto"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"company_id","type":"UUID","nullable":False,"pk":False,"fk":"companies.id","unique":False,"default":None,"description":"Empresa mandante del proyecto"},
      {"name":"code","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Código único del proyecto (generado automáticamente)"},
      {"name":"name","type":"VARCHAR(300)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre del proyecto"},
      {"name":"description","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Descripción detallada del proyecto"},
      {"name":"address","type":"VARCHAR(500)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Dirección de la obra"},
      {"name":"commune","type":"VARCHAR(100)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Comuna de la obra"},
      {"name":"region","type":"VARCHAR(100)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Región de la obra"},
      {"name":"latitude","type":"DECIMAL(10,7)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Latitud GPS de la obra"},
      {"name":"longitude","type":"DECIMAL(10,7)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Longitud GPS de la obra"},
      {"name":"start_date","type":"DATE","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de inicio planificada del proyecto"},
      {"name":"end_date","type":"DATE","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de término planificada del proyecto"},
      {"name":"status","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'activo'","description":"Estado actual del proyecto"},
      {"name":"ito_manager_id","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"ITO jefe responsable del proyecto"},
      {"name":"budget","type":"DECIMAL(18,2)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Presupuesto del proyecto"},
      {"name":"currency","type":"VARCHAR(10)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'CLP'","description":"Moneda del presupuesto"},
      {"name":"cover_image_url","type":"VARCHAR(500)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"URL de imagen de portada del proyecto en MinIO"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"updated_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de última actualización"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó el registro"},
      {"name":"updated_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que actualizó el registro"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"deleted_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que realizó el borrado lógico"},
    ],
    "constraints":[
      "chk_projects_status: status IN ('activo','pausado','finalizado','cancelado')",
      "UNIQUE(tenant_id, code)"
    ]
  },
  {
    "name": "project_members",
    "description": "Miembros del equipo ITO asignados a un proyecto. Define el rol que cada usuario cumple dentro del proyecto.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único del registro"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"project_id","type":"UUID","nullable":False,"pk":False,"fk":"projects.id","unique":False,"default":None,"description":"Proyecto al que pertenece el miembro"},
      {"name":"user_id","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario miembro del proyecto"},
      {"name":"project_role","type":"VARCHAR(100)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'inspector'","description":"Rol del usuario en el proyecto"},
      {"name":"is_active","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si el miembro está activo en el proyecto"},
      {"name":"joined_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de incorporación al proyecto"},
      {"name":"left_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de salida del proyecto"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación del registro"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que añadió al miembro"},
    ],
    "constraints":[
      "chk_pm_role: project_role IN ('ito_jefe','ito_inspector','ito_coordinador','supervisor','observador')",
      "UNIQUE(tenant_id, project_id, user_id)"
    ]
  },
  {
    "name": "project_stages",
    "description": "Etapas del proyecto (ej. Fundaciones, Estructura, Terminaciones). Permite organizar las inspecciones por fase de obra.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la etapa"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"project_id","type":"UUID","nullable":False,"pk":False,"fk":"projects.id","unique":False,"default":None,"description":"Proyecto al que pertenece la etapa"},
      {"name":"name","type":"VARCHAR(200)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre de la etapa"},
      {"name":"description","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Descripción de la etapa"},
      {"name":"order_index","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"0","description":"Orden de presentación de la etapa"},
      {"name":"start_date","type":"DATE","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de inicio planificada"},
      {"name":"end_date","type":"DATE","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de término planificada"},
      {"name":"status","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'pendiente'","description":"Estado de la etapa"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó la etapa"},
    ],
    "constraints":["chk_pstage_status: status IN ('pendiente','en_progreso','completada','cancelada')"]
  },
  {
    "name": "project_sectors",
    "description": "Sectores jerárquicos dentro del proyecto (ej. Torre A, Piso 3, Bloque Norte). Permite organizar espacialmente las unidades inspeccionables.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único del sector"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"project_id","type":"UUID","nullable":False,"pk":False,"fk":"projects.id","unique":False,"default":None,"description":"Proyecto al que pertenece el sector"},
      {"name":"parent_id","type":"UUID","nullable":True,"pk":False,"fk":"project_sectors.id","unique":False,"default":None,"description":"Sector padre (jerarquía: Torre → Piso → Depto)"},
      {"name":"name","type":"VARCHAR(200)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre del sector"},
      {"name":"level","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"1","description":"Nivel de jerarquía (1=raíz, 2=subnivel, etc.)"},
      {"name":"order_index","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"0","description":"Orden de presentación"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó el sector"},
    ],
    "constraints":[]
  },
  {
    "name": "project_units",
    "description": "Unidades inspeccionables del proyecto (ej. Depto 101, Área de estacionamiento). Cada unidad puede ser inspeccionada independientemente.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la unidad"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"project_id","type":"UUID","nullable":False,"pk":False,"fk":"projects.id","unique":False,"default":None,"description":"Proyecto al que pertenece la unidad"},
      {"name":"sector_id","type":"UUID","nullable":True,"pk":False,"fk":"project_sectors.id","unique":False,"default":None,"description":"Sector al que pertenece la unidad"},
      {"name":"stage_id","type":"UUID","nullable":True,"pk":False,"fk":"project_stages.id","unique":False,"default":None,"description":"Etapa asociada a la unidad"},
      {"name":"code","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Código identificador de la unidad"},
      {"name":"name","type":"VARCHAR(200)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre descriptivo de la unidad"},
      {"name":"unit_type","type":"VARCHAR(100)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tipo de unidad (depto, casa, local, bodega, etc.)"},
      {"name":"area_m2","type":"DECIMAL(10,2)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Superficie de la unidad en m²"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó la unidad"},
    ],
    "constraints":["UNIQUE(tenant_id, project_id, code)"]
  },
  {
    "name": "specialties",
    "description": "Especialidades técnicas de inspección (ej. Estructural, Eléctrica, Sanitaria). Pertenecen a un tenant y se usan para clasificar inspecciones y contratistas.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la especialidad"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"name","type":"VARCHAR(200)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre de la especialidad"},
      {"name":"code","type":"VARCHAR(50)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Código corto de la especialidad"},
      {"name":"description","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Descripción de la especialidad técnica"},
      {"name":"color","type":"VARCHAR(20)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Color para identificación visual en la UI"},
      {"name":"icon","type":"VARCHAR(100)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre de ícono para la UI"},
      {"name":"is_active","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si la especialidad está activa"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó la especialidad"},
    ],
    "constraints":["UNIQUE(tenant_id, code) WHERE code IS NOT NULL"]
  },
  {
    "name": "contractors",
    "description": "Empresas contratistas que ejecutan obras en los proyectos. Pueden tener múltiples especialidades y ser asignadas a proyectos.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único del contratista"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"name","type":"VARCHAR(300)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Razón social o nombre del contratista"},
      {"name":"rut","type":"VARCHAR(20)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"RUT o identificador tributario"},
      {"name":"contact_name","type":"VARCHAR(200)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre de la persona de contacto"},
      {"name":"contact_email","type":"VARCHAR(256)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Email de contacto"},
      {"name":"contact_phone","type":"VARCHAR(50)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Teléfono de contacto"},
      {"name":"address","type":"VARCHAR(500)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Dirección del contratista"},
      {"name":"is_active","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si el contratista está activo"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"updated_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de última actualización"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó el contratista"},
      {"name":"updated_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que actualizó el registro"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"deleted_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que realizó el borrado lógico"},
    ],
    "constraints":[]
  },
  {
    "name": "contractor_specialties",
    "description": "Relación entre contratistas y sus especialidades técnicas. Define en qué áreas puede trabajar cada contratista.",
    "columns": [
      {"name":"contractor_id","type":"UUID","nullable":False,"pk":True,"fk":"contractors.id","unique":False,"default":None,"description":"Contratista"},
      {"name":"specialty_id","type":"UUID","nullable":False,"pk":True,"fk":"specialties.id","unique":False,"default":None,"description":"Especialidad del contratista"},
    ],
    "constraints":["PK compuesta: (contractor_id, specialty_id)"]
  },
  {
    "name": "project_contractors",
    "description": "Contratistas asignados a un proyecto específico. Registra el contrato, fechas y estado de la asignación.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la asignación"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"project_id","type":"UUID","nullable":False,"pk":False,"fk":"projects.id","unique":False,"default":None,"description":"Proyecto al que se asigna el contratista"},
      {"name":"contractor_id","type":"UUID","nullable":False,"pk":False,"fk":"contractors.id","unique":False,"default":None,"description":"Contratista asignado"},
      {"name":"specialty_id","type":"UUID","nullable":True,"pk":False,"fk":"specialties.id","unique":False,"default":None,"description":"Especialidad específica en este proyecto"},
      {"name":"contract_number","type":"VARCHAR(100)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Número de contrato"},
      {"name":"start_date","type":"DATE","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de inicio del contrato"},
      {"name":"end_date","type":"DATE","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de término del contrato"},
      {"name":"is_active","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si la asignación está activa"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó la asignación"},
    ],
    "constraints":[]
  },
  {
    "name": "inspection_templates",
    "description": "Plantillas reutilizables para inspecciones. Define la estructura base de una pauta de inspección: nombre, especialidad, tipo de proyecto y configuración.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la plantilla"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"specialty_id","type":"UUID","nullable":True,"pk":False,"fk":"specialties.id","unique":False,"default":None,"description":"Especialidad técnica de la plantilla"},
      {"name":"name","type":"VARCHAR(300)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre de la plantilla"},
      {"name":"description","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Descripción del alcance de la plantilla"},
      {"name":"template_type","type":"VARCHAR(100)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'general'","description":"Tipo de plantilla"},
      {"name":"is_public","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si la plantilla es visible para todos los tenants"},
      {"name":"is_active","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si la plantilla está activa"},
      {"name":"current_version","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"1","description":"Número de versión actual de la plantilla"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"updated_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de última actualización"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó la plantilla"},
      {"name":"updated_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que actualizó la plantilla"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"deleted_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que realizó el borrado lógico"},
    ],
    "constraints":[]
  },
  {
    "name": "template_versions",
    "description": "Historial de versiones de las plantillas de inspección. Cada vez que se publica un cambio se crea una nueva versión, preservando el historial.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la versión"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"template_id","type":"UUID","nullable":False,"pk":False,"fk":"inspection_templates.id","unique":False,"default":None,"description":"Plantilla a la que pertenece la versión"},
      {"name":"version_number","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Número secuencial de versión"},
      {"name":"change_summary","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Resumen de los cambios realizados en esta versión"},
      {"name":"snapshot","type":"JSONB","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Snapshot completo de la estructura de la plantilla en esta versión"},
      {"name":"published_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de publicación de la versión"},
      {"name":"published_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que publicó la versión"},
    ],
    "constraints":["UNIQUE(template_id, version_number)"]
  },
  {
    "name": "template_sections",
    "description": "Secciones de una plantilla de inspección (ej. 'Estructura', 'Terminaciones'). Agrupan las preguntas por área temática.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la sección"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"template_id","type":"UUID","nullable":False,"pk":False,"fk":"inspection_templates.id","unique":False,"default":None,"description":"Plantilla a la que pertenece la sección"},
      {"name":"name","type":"VARCHAR(300)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre de la sección"},
      {"name":"description","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Descripción de la sección"},
      {"name":"order_index","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"0","description":"Orden de presentación de la sección"},
      {"name":"is_required","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si la sección es obligatoria"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
    ],
    "constraints":[]
  },
  {
    "name": "template_questions",
    "description": "Preguntas de cada sección de la plantilla. Soporta múltiples tipos de respuesta: texto, numérico, booleano, selección, foto, firma y fecha.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la pregunta"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"section_id","type":"UUID","nullable":False,"pk":False,"fk":"template_sections.id","unique":False,"default":None,"description":"Sección a la que pertenece la pregunta"},
      {"name":"template_id","type":"UUID","nullable":False,"pk":False,"fk":"inspection_templates.id","unique":False,"default":None,"description":"Plantilla a la que pertenece la pregunta"},
      {"name":"question_text","type":"TEXT","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Texto de la pregunta a responder"},
      {"name":"question_type","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tipo de respuesta esperada"},
      {"name":"order_index","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"0","description":"Orden dentro de la sección"},
      {"name":"is_required","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si la respuesta es obligatoria"},
      {"name":"weight","type":"DECIMAL(5,2)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"1.0","description":"Peso de la pregunta en el puntaje total"},
      {"name":"min_value","type":"DECIMAL(10,2)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Valor mínimo aceptable (para tipo numérico)"},
      {"name":"max_value","type":"DECIMAL(10,2)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Valor máximo aceptable (para tipo numérico)"},
      {"name":"unit","type":"VARCHAR(50)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Unidad de medida (para tipo numérico, ej. 'cm', 'kg')"},
      {"name":"requires_evidence","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si la pregunta requiere foto u otro archivo adjunto"},
      {"name":"help_text","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Texto de ayuda o instrucciones adicionales para el inspector"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
    ],
    "constraints":[
      "chk_tq_type: question_type IN ('text','numeric','boolean','multiple_choice','single_choice','photo','signature','date')"
    ]
  },
  {
    "name": "template_question_options",
    "description": "Opciones de respuesta para preguntas de tipo multiple_choice o single_choice. Cada opción puede tener un puntaje asociado.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la opción"},
      {"name":"question_id","type":"UUID","nullable":False,"pk":False,"fk":"template_questions.id","unique":False,"default":None,"description":"Pregunta a la que pertenece la opción"},
      {"name":"option_text","type":"VARCHAR(500)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Texto de la opción de respuesta"},
      {"name":"score","type":"DECIMAL(5,2)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"0","description":"Puntaje asignado si se selecciona esta opción"},
      {"name":"order_index","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"0","description":"Orden de presentación de la opción"},
      {"name":"is_correct","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si es la respuesta correcta/esperada"},
    ],
    "constraints":[]
  },
  {
    "name": "inspections",
    "description": "Inspecciones ejecutadas sobre una unidad del proyecto. Registra el estado, puntaje, inspector, fechas y evidencias del proceso de inspección.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la inspección"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"project_id","type":"UUID","nullable":False,"pk":False,"fk":"projects.id","unique":False,"default":None,"description":"Proyecto al que pertenece la inspección"},
      {"name":"unit_id","type":"UUID","nullable":True,"pk":False,"fk":"project_units.id","unique":False,"default":None,"description":"Unidad inspeccionada (opcional)"},
      {"name":"stage_id","type":"UUID","nullable":True,"pk":False,"fk":"project_stages.id","unique":False,"default":None,"description":"Etapa del proyecto inspeccionada"},
      {"name":"template_id","type":"UUID","nullable":False,"pk":False,"fk":"inspection_templates.id","unique":False,"default":None,"description":"Plantilla utilizada para la inspección"},
      {"name":"template_version","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"1","description":"Versión de la plantilla usada al momento de la inspección"},
      {"name":"specialty_id","type":"UUID","nullable":True,"pk":False,"fk":"specialties.id","unique":False,"default":None,"description":"Especialidad técnica de la inspección"},
      {"name":"contractor_id","type":"UUID","nullable":True,"pk":False,"fk":"contractors.id","unique":False,"default":None,"description":"Contratista cuyo trabajo se inspecciona"},
      {"name":"code","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Código correlativo único de la inspección (ej. INS-2026-0001)"},
      {"name":"title","type":"VARCHAR(300)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Título descriptivo de la inspección"},
      {"name":"description","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Descripción del alcance y objetivo de la inspección"},
      {"name":"status","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'programada'","description":"Estado actual de la inspección"},
      {"name":"result","type":"VARCHAR(50)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Resultado final de la inspección"},
      {"name":"scheduled_date","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha y hora programada para la inspección"},
      {"name":"started_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha y hora de inicio efectivo"},
      {"name":"finished_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha y hora de finalización efectiva"},
      {"name":"total_score","type":"DECIMAL(5,2)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Puntaje total obtenido en la inspección"},
      {"name":"max_score","type":"DECIMAL(5,2)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Puntaje máximo posible de la plantilla"},
      {"name":"score_percentage","type":"DECIMAL(5,2)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Porcentaje de cumplimiento (total_score/max_score × 100)"},
      {"name":"inspector_id","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Inspector responsable de ejecutar la inspección"},
      {"name":"reviewer_id","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que revisó y aprobó el informe"},
      {"name":"general_observations","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Observaciones generales del inspector sobre la inspección"},
      {"name":"is_offline","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si la inspección se realizó offline en la app móvil"},
      {"name":"synced_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de sincronización con el servidor (para inspecciones offline)"},
      {"name":"signature_url","type":"VARCHAR(500)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"URL de la firma digital del inspector en MinIO"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación del registro"},
      {"name":"updated_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de última actualización"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó el registro"},
      {"name":"updated_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que actualizó el registro"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"deleted_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que realizó el borrado lógico"},
    ],
    "constraints":[
      "chk_insp_status: status IN ('programada','en_progreso','completada','cancelada','pendiente_revision')",
      "chk_insp_result: result IS NULL OR result IN ('aprobada','aprobada_con_obs','rechazada')",
      "UNIQUE(tenant_id, code)"
    ]
  },
  {
    "name": "inspection_answers",
    "description": "Respuestas a cada pregunta de la plantilla en una inspección específica. Almacena el valor respondido, puntaje obtenido y si genera observación.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la respuesta"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"inspection_id","type":"UUID","nullable":False,"pk":False,"fk":"inspections.id","unique":False,"default":None,"description":"Inspección a la que pertenece la respuesta"},
      {"name":"question_id","type":"UUID","nullable":False,"pk":False,"fk":"template_questions.id","unique":False,"default":None,"description":"Pregunta respondida"},
      {"name":"answer_text","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Respuesta en texto libre (para tipo text)"},
      {"name":"answer_numeric","type":"DECIMAL(10,2)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Respuesta numérica (para tipo numeric)"},
      {"name":"answer_boolean","type":"BOOLEAN","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Respuesta booleana (para tipo boolean)"},
      {"name":"answer_option_id","type":"UUID","nullable":True,"pk":False,"fk":"template_question_options.id","unique":False,"default":None,"description":"Opción seleccionada (para single_choice)"},
      {"name":"answer_options","type":"UUID[]","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Opciones seleccionadas (para multiple_choice, array de UUIDs)"},
      {"name":"answer_date","type":"DATE","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Respuesta de tipo fecha"},
      {"name":"score_obtained","type":"DECIMAL(5,2)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Puntaje obtenido en esta respuesta"},
      {"name":"creates_observation","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si esta respuesta generó automáticamente una observación"},
      {"name":"notes","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Notas del inspector sobre esta respuesta"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación de la respuesta"},
      {"name":"updated_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de última actualización"},
    ],
    "constraints":["UNIQUE(inspection_id, question_id)"]
  },
  {
    "name": "inspection_evidence",
    "description": "Archivos multimedia adjuntos a una respuesta de inspección (fotos, videos, documentos). Almacenados en MinIO.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la evidencia"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"inspection_id","type":"UUID","nullable":False,"pk":False,"fk":"inspections.id","unique":False,"default":None,"description":"Inspección a la que pertenece la evidencia"},
      {"name":"answer_id","type":"UUID","nullable":True,"pk":False,"fk":"inspection_answers.id","unique":False,"default":None,"description":"Respuesta específica a la que está vinculada (opcional)"},
      {"name":"file_name","type":"VARCHAR(300)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre del archivo"},
      {"name":"file_path","type":"VARCHAR(1000)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Ruta en MinIO"},
      {"name":"file_size_bytes","type":"BIGINT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tamaño del archivo en bytes"},
      {"name":"mime_type","type":"VARCHAR(100)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tipo MIME del archivo"},
      {"name":"evidence_type","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'photo'","description":"Tipo de evidencia"},
      {"name":"caption","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Descripción o pie de foto de la evidencia"},
      {"name":"latitude","type":"DECIMAL(10,7)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Latitud GPS donde se tomó la evidencia"},
      {"name":"longitude","type":"DECIMAL(10,7)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Longitud GPS donde se tomó la evidencia"},
      {"name":"taken_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha y hora en que se capturó la evidencia"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación del registro"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que subió la evidencia"},
    ],
    "constraints":[
      "chk_ev_type: evidence_type IN ('photo','video','document','audio','signature')"
    ]
  },
  {
    "name": "observations",
    "description": "Observaciones o No Conformidades levantadas durante o después de una inspección. Registra el hallazgo, gravedad, responsable de corrección y fechas.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la observación"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"project_id","type":"UUID","nullable":False,"pk":False,"fk":"projects.id","unique":False,"default":None,"description":"Proyecto al que pertenece la observación"},
      {"name":"inspection_id","type":"UUID","nullable":True,"pk":False,"fk":"inspections.id","unique":False,"default":None,"description":"Inspección que originó la observación (puede ser manual)"},
      {"name":"answer_id","type":"UUID","nullable":True,"pk":False,"fk":"inspection_answers.id","unique":False,"default":None,"description":"Respuesta específica que originó la observación"},
      {"name":"contractor_id","type":"UUID","nullable":True,"pk":False,"fk":"contractors.id","unique":False,"default":None,"description":"Contratista responsable de corregir"},
      {"name":"specialty_id","type":"UUID","nullable":True,"pk":False,"fk":"specialties.id","unique":False,"default":None,"description":"Especialidad técnica de la observación"},
      {"name":"code","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Código correlativo único (ej. OBS-2026-0001)"},
      {"name":"title","type":"VARCHAR(300)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Título de la observación"},
      {"name":"description","type":"TEXT","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Descripción detallada del hallazgo"},
      {"name":"severity","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'media'","description":"Nivel de gravedad de la observación"},
      {"name":"status","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'abierta'","description":"Estado actual de la observación"},
      {"name":"due_date","type":"DATE","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha límite para resolver la observación"},
      {"name":"resolved_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha y hora en que fue resuelta"},
      {"name":"resolved_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que marcó la observación como resuelta"},
      {"name":"resolution_notes","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Notas sobre la resolución y acciones tomadas"},
      {"name":"assigned_to","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario responsable de gestionar la observación"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"updated_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de última actualización"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó la observación"},
      {"name":"updated_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que actualizó el registro"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"deleted_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que realizó el borrado lógico"},
    ],
    "constraints":[
      "chk_obs_severity: severity IN ('critica','alta','media','baja')",
      "chk_obs_status: status IN ('abierta','en_proceso','resuelta','cerrada','vencida')",
      "UNIQUE(tenant_id, code)"
    ]
  },
  {
    "name": "observation_history",
    "description": "Historial de cambios de estado de las observaciones. Registra cada transición con quién la realizó, cuándo y el comentario.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único del registro histórico"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"observation_id","type":"UUID","nullable":False,"pk":False,"fk":"observations.id","unique":False,"default":None,"description":"Observación cuyo historial se registra"},
      {"name":"from_status","type":"VARCHAR(50)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Estado anterior de la observación"},
      {"name":"to_status","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nuevo estado de la observación"},
      {"name":"comment","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Comentario del usuario sobre el cambio de estado"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha y hora del cambio de estado"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que realizó el cambio de estado"},
    ],
    "constraints":[]
  },
  {
    "name": "reinspections",
    "description": "Reinspecciones realizadas para verificar la corrección de observaciones. Cada reinspección verifica si el contratista subsanó los hallazgos.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la reinspección"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"observation_id","type":"UUID","nullable":False,"pk":False,"fk":"observations.id","unique":False,"default":None,"description":"Observación que se está reinspeccionando"},
      {"name":"inspector_id","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Inspector que ejecuta la reinspección"},
      {"name":"code","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Código correlativo único (ej. REINSP-2026-0001)"},
      {"name":"scheduled_date","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha programada para la reinspección"},
      {"name":"executed_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha y hora de ejecución efectiva"},
      {"name":"status","type":"VARCHAR(50)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'pendiente'","description":"Estado de la reinspección"},
      {"name":"result","type":"VARCHAR(50)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Resultado de la reinspección"},
      {"name":"notes","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Notas del inspector sobre la reinspección"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"updated_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de última actualización"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que creó la reinspección"},
      {"name":"updated_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que actualizó el registro"},
    ],
    "constraints":[
      "chk_reinsp_status: status IN ('pendiente','programada','ejecutada')",
      "chk_reinsp_result: result IS NULL OR result IN ('aprobada','rechazada')"
    ]
  },
  {
    "name": "project_documents",
    "description": "Documentos adjuntos a proyectos, inspecciones u observaciones: planos, especificaciones, contratos, procedimientos, informes, etc.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único del documento"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"project_id","type":"UUID","nullable":False,"pk":False,"fk":"projects.id","unique":False,"default":None,"description":"Proyecto al que pertenece el documento"},
      {"name":"inspection_id","type":"UUID","nullable":True,"pk":False,"fk":"inspections.id","unique":False,"default":None,"description":"Inspección relacionada (opcional)"},
      {"name":"observation_id","type":"UUID","nullable":True,"pk":False,"fk":"observations.id","unique":False,"default":None,"description":"Observación relacionada (opcional)"},
      {"name":"category","type":"VARCHAR(100)","nullable":False,"pk":False,"fk":None,"unique":False,"default":"'general'","description":"Categoría del documento"},
      {"name":"name","type":"VARCHAR(300)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre descriptivo del documento"},
      {"name":"description","type":"TEXT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Descripción del contenido"},
      {"name":"file_name","type":"VARCHAR(300)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Nombre del archivo en el sistema"},
      {"name":"file_path","type":"VARCHAR(1000)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Ruta de almacenamiento en MinIO"},
      {"name":"file_size_bytes","type":"BIGINT","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tamaño del archivo en bytes"},
      {"name":"mime_type","type":"VARCHAR(100)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tipo MIME del archivo"},
      {"name":"version","type":"VARCHAR(50)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Versión del documento (ej. Rev.2)"},
      {"name":"is_active","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"TRUE","description":"Indica si el documento está activo"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha de creación"},
      {"name":"updated_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de última actualización"},
      {"name":"created_by","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que subió el documento"},
      {"name":"updated_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que actualizó el registro"},
      {"name":"is_deleted","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Borrado lógico"},
      {"name":"deleted_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha de borrado lógico"},
      {"name":"deleted_by","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que realizó el borrado lógico"},
    ],
    "constraints":[
      "chk_doc_category: category IN ('plano','especificacion','contrato','procedimiento','informe','foto','otro','general')"
    ]
  },
  {
    "name": "notifications",
    "description": "Notificaciones internas del sistema enviadas a usuarios. Registra si fueron leídas, enviadas por email y por push móvil.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único de la notificación"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant propietario"},
      {"name":"user_id","type":"UUID","nullable":False,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario destinatario"},
      {"name":"type","type":"VARCHAR(100)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tipo de notificación (ej. inspeccion_asignada, observacion_vencida)"},
      {"name":"title","type":"VARCHAR(300)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Título de la notificación"},
      {"name":"body","type":"TEXT","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Cuerpo o mensaje de la notificación"},
      {"name":"entity_type","type":"VARCHAR(100)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tipo de entidad relacionada (polimórfico)"},
      {"name":"entity_id","type":"UUID","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"ID de la entidad relacionada (sin FK formal)"},
      {"name":"is_read","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si el usuario ya leyó la notificación"},
      {"name":"read_at","type":"TIMESTAMPTZ","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Fecha y hora en que fue leída"},
      {"name":"sent_email","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si se envió por correo electrónico"},
      {"name":"sent_push","type":"BOOLEAN","nullable":False,"pk":False,"fk":None,"unique":False,"default":"FALSE","description":"Indica si se envió como push móvil"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha y hora de creación"},
    ],
    "constraints":[]
  },
  {
    "name": "audit_logs",
    "description": "Registro de auditoría de todas las acciones relevantes del sistema. Almacena qué usuario hizo qué acción, sobre qué entidad, con los valores antes/después.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único del evento de auditoría"},
      {"name":"tenant_id","type":"UUID","nullable":True,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant (puede ser nulo para acciones del sistema)"},
      {"name":"user_id","type":"UUID","nullable":True,"pk":False,"fk":"users.id","unique":False,"default":None,"description":"Usuario que realizó la acción (nulo para acciones automáticas)"},
      {"name":"action","type":"VARCHAR(100)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Acción realizada (ej. CREATE, UPDATE, DELETE, LOGIN)"},
      {"name":"entity_type","type":"VARCHAR(200)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tipo de entidad afectada (ej. inspection, project)"},
      {"name":"entity_id","type":"VARCHAR(200)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"ID de la entidad afectada (como texto para flexibilidad)"},
      {"name":"old_values","type":"JSONB","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Valores anteriores al cambio (snapshot JSON)"},
      {"name":"new_values","type":"JSONB","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Valores nuevos después del cambio (snapshot JSON)"},
      {"name":"ip_address","type":"VARCHAR(50)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Dirección IP de la acción"},
      {"name":"user_agent","type":"VARCHAR(500)","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"User agent del cliente"},
      {"name":"additional_data","type":"JSONB","nullable":True,"pk":False,"fk":None,"unique":False,"default":None,"description":"Datos adicionales de contexto en JSON libre"},
      {"name":"created_at","type":"TIMESTAMPTZ","nullable":False,"pk":False,"fk":None,"unique":False,"default":"NOW()","description":"Fecha y hora del evento"},
    ],
    "constraints":[]
  },
  {
    "name": "sequence_counters",
    "description": "Generador de códigos correlativos por tenant, tipo de entidad y año. Usado para producir códigos del tipo INS-2026-0001, OBS-2026-0001.",
    "columns": [
      {"name":"id","type":"UUID","nullable":False,"pk":True,"fk":None,"unique":False,"default":"uuid_generate_v4()","description":"Identificador único del contador"},
      {"name":"tenant_id","type":"UUID","nullable":False,"pk":False,"fk":"tenants.id","unique":False,"default":None,"description":"Tenant al que pertenece el contador"},
      {"name":"entity_type","type":"VARCHAR(100)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Tipo de entidad (ej. inspection, observation, reinspection)"},
      {"name":"prefix","type":"VARCHAR(20)","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Prefijo del código (ej. INS, OBS, REINSP)"},
      {"name":"year","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":None,"description":"Año al que corresponde el correlativo"},
      {"name":"last_value","type":"INT","nullable":False,"pk":False,"fk":None,"unique":False,"default":"0","description":"Último número correlativo emitido"},
    ],
    "constraints":["UNIQUE(tenant_id, entity_type, year)"]
  },
]

# ─── CONSTRUIR EL DOCUMENTO ──────────────────────────────────────────
doc = Document()

# Márgenes
section = doc.sections[0]
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Fuente base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ── Portada ──────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = titulo.add_run('ITO CLOUD')
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = RGBColor(*AZUL_OSCURO)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Diccionario de Datos — Base de Datos')
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(*AZUL_MEDIO)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sub2.add_run('Sistema de Inspección Técnica de Obras (SaaS)')
r3.font.size = Pt(13)
r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)  # ya ok

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Motor: PostgreSQL 18   |   Versión esquema: 1.0   |   Fecha: 2026-03-18').font.size = Pt(10)

doc.add_page_break()

# ── Introducción ─────────────────────────────────────────────────────
add_heading(doc, '1. Introducción', level=1)
intro = doc.add_paragraph(
    'Este documento describe la estructura completa de la base de datos del sistema ITO Cloud, '
    'una plataforma SaaS de Inspección Técnica de Obras. El esquema está diseñado sobre '
    'PostgreSQL 18 con arquitectura multi-tenant lógica, claves primarias UUID, soft delete '
    'y auditoría completa en todas las entidades.'
)
intro.paragraph_format.space_after = Pt(6)

add_heading(doc, '2. Convenciones', level=1)
convs = [
    'PK  → Clave primaria (Primary Key). Columna resaltada en verde.',
    'FK  → Clave foránea (Foreign Key). Columna marcada con ✓ en columna FK.',
    'NOT NULL → Columna requerida. Columna resaltada en rojo cuando no es PK.',
    'UUID → Tipo de identificador usado en todas las PKs para soportar generación offline.',
    'tenant_id → Presente en todas las tablas de negocio. Soporta multi-tenancy lógico.',
    'is_deleted / deleted_at / deleted_by → Patrón de borrado lógico (soft delete).',
    'created_by / updated_by / deleted_by → Auditoría de usuario.',
    'TIMESTAMPTZ → Todos los timestamps incluyen zona horaria.',
]
for c in convs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(c).font.size = Pt(9)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)

add_heading(doc, '3. Resumen de Tablas', level=1)

# Tabla resumen
tbl_res = doc.add_table(rows=1, cols=3)
tbl_res.style = 'Table Grid'
hdr = tbl_res.rows[0]
for cell, txt in zip(hdr.cells, ['#', 'Tabla', 'Descripción corta']):
    set_cell_bg(cell, AZUL_OSCURO)
    p2 = cell.paragraphs[0]
    r = p2.add_run(txt)
    r.bold = True
    r.font.color.rgb = RGBColor(*BLANCO)
    r.font.size = Pt(9)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

domains = {
    'tenants':'Multi-tenancy','users':'Identidad','roles':'Identidad',
    'user_roles':'Identidad','user_claims':'Identidad','role_claims':'Identidad',
    'user_tokens':'Identidad','user_logins':'Identidad','companies':'Empresas',
    'projects':'Proyectos','project_members':'Proyectos','project_stages':'Proyectos',
    'project_sectors':'Proyectos','project_units':'Proyectos','specialties':'Catálogos',
    'contractors':'Contratistas','contractor_specialties':'Contratistas',
    'project_contractors':'Contratistas','inspection_templates':'Plantillas',
    'template_versions':'Plantillas','template_sections':'Plantillas',
    'template_questions':'Plantillas','template_question_options':'Plantillas',
    'inspections':'Inspecciones','inspection_answers':'Inspecciones',
    'inspection_evidence':'Inspecciones','observations':'Observaciones',
    'observation_history':'Observaciones','reinspections':'Reinspecciones',
    'project_documents':'Documentos','notifications':'Notificaciones',
    'audit_logs':'Auditoría','sequence_counters':'Sistema',
}

for i, t in enumerate(TABLES):
    row = tbl_res.add_row()
    bg = AZUL_CLARO if i % 2 == 0 else GRIS_CLARO
    for cell, val in zip(row.cells, [str(i+1), t['name'], t['description'][:80]+'...' if len(t['description'])>80 else t['description']]):
        set_cell_bg(cell, bg)
        cell.paragraphs[0].add_run(val).font.size = Pt(8)

doc.add_page_break()

# ── Diccionario por tabla ─────────────────────────────────────────────
add_heading(doc, '4. Definición de Tablas', level=1)

for t in TABLES:
    add_table_dict(doc, t)

# ── Leyenda de colores ────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '5. Leyenda', level=1)
leyenda = [
    ('Verde / negrita', 'Columna PK (clave primaria)'),
    ('Rojo', 'Columna NOT NULL (requerida, que no es PK)'),
    ('✓ naranja (col. FK)', 'Columna con clave foránea'),
    ('Fila azul claro', 'Fila par de columnas'),
    ('Fila gris claro', 'Fila impar de columnas'),
]
tbl_ley = doc.add_table(rows=1, cols=2)
tbl_ley.style = 'Table Grid'
for cell, txt in zip(tbl_ley.rows[0].cells, ['Color / Símbolo', 'Significado']):
    set_cell_bg(cell, AZUL_OSCURO)
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True; r.font.color.rgb = RGBColor(*BLANCO); r.font.size = Pt(9)
for sym, sig in leyenda:
    row = tbl_ley.add_row()
    row.cells[0].paragraphs[0].add_run(sym).font.size = Pt(9)
    row.cells[1].paragraphs[0].add_run(sig).font.size = Pt(9)

# ── Guardar ───────────────────────────────────────────────────────────
out = r'F:\SISTEMA DE INSPECCION DE OBRA\docs\ITO_Cloud_Diccionario_Datos.docx'
doc.save(out)
print(f'Guardado: {out}')
