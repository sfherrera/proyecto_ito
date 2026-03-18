"""
Genera el documento ejecutivo de funcionalidades de ITO Cloud en formato Word.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level, color_hex='1F3864'):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return h

def add_table_header_row(table, headers, bg='1F3864'):
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = header
        set_cell_bg(cell, bg)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0] if para.runs else para.add_run(header)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

def add_table_row(table, values, bold_first=False, bg=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        if bg:
            set_cell_bg(cell, bg)
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(val)
        if bold_first and i == 0:
            run.font.bold = True
        run.font.size = Pt(9)
    return row

# ─────────────────────────────────────────────
# Documento
# ─────────────────────────────────────────────

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── PORTADA ──────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('ITO CLOUD')
r.font.bold = True
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run('Plataforma de Inspección Técnica de Obras')
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x26, 0x96, 0x88)
r2.font.bold = True

doc.add_paragraph()

doc_title = doc.add_paragraph()
doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = doc_title.add_run('DOCUMENTO EJECUTIVO DE FUNCIONALIDADES')
r3.font.size = Pt(14)
r3.font.bold = True
r3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Versión: 1.0   |   Fecha: {datetime.date.today().strftime("%d de %B de %Y")}   |   Confidencial')
meta.runs[0].font.size = Pt(10)
meta.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ── 1. RESUMEN EJECUTIVO ─────────────────────
add_heading(doc, '1. Resumen Ejecutivo', 1)

p = doc.add_paragraph(
    'ITO Cloud es una plataforma SaaS B2B multi-empresa para la gestión digital completa del proceso de '
    'Inspección Técnica de Obras (ITO) en proyectos de construcción e inmobiliarias. '
    'La plataforma reemplaza el uso de formularios en papel, Excel, WhatsApp y correos electrónicos '
    'por un sistema trazable, centralizado y con capacidades de análisis en tiempo real.'
)
p.style.font.size = Pt(10)

doc.add_paragraph()

p2 = doc.add_paragraph(
    'El sistema estará compuesto por tres componentes principales:'
)
p2.style.font.size = Pt(10)

bullets = [
    'Portal Web de gestión y supervisión (Blazor Server)',
    'API REST segura para integraciones y app móvil (ASP.NET Core 8)',
    'Aplicación Android offline-first para trabajo en terreno (Kotlin + Jetpack Compose)',
]
for b in bullets:
    bp = doc.add_paragraph(b, style='List Bullet')
    bp.runs[0].font.size = Pt(10)

# ── 2. PROBLEMA QUE RESUELVE ─────────────────
add_heading(doc, '2. Problema que Resuelve', 1)

p = doc.add_paragraph(
    'La gestión de inspecciones técnicas en obra presenta actualmente los siguientes problemas:'
)
p.style.font.size = Pt(10)

problemas = [
    ('Registros manuales en papel o Excel', 'Sin trazabilidad, propenso a errores y pérdidas de información'),
    ('WhatsApp como herramienta de evidencia', 'Sin orden, sin acceso histórico, sin vínculo a observaciones'),
    ('Correos sueltos para comunicar observaciones', 'Sin flujo formal ni control de estado de resolución'),
    ('Ausencia de KPIs de calidad', 'No hay métricas de cumplimiento ni historial de rendimiento de contratistas'),
    ('Falta de trazabilidad', 'Imposible auditar quién inspeccionó qué, cuándo y con qué resultado'),
]

t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
add_table_header_row(t, ['Problema', 'Impacto'], '2E4057')
for prob, imp in problemas:
    add_table_row(t, [prob, imp], bold_first=True)

# ── 3. PROPUESTA DE VALOR ────────────────────
add_heading(doc, '3. Propuesta de Valor por Actor', 1)

actores = [
    ('ITO / Inspector', 'App móvil con funcionamiento offline, registro ágil en terreno con checklist digital y captura fotográfica integrada'),
    ('Supervisor de obra', 'Trazabilidad en tiempo real de hallazgos, observaciones y estado de correcciones'),
    ('Constructora / Empresa', 'Dashboard ejecutivo con KPIs, reportes automáticos PDF y Excel, historial completo por proyecto'),
    ('Mandante / Cliente final', 'Transparencia total del proceso de inspección y evidencia digital del cumplimiento'),
    ('Contratista', 'Portal para visualizar y gestionar las observaciones que le han sido asignadas para corrección'),
]

t2 = doc.add_table(rows=1, cols=2)
t2.style = 'Table Grid'
add_table_header_row(t2, ['Actor', 'Beneficio'], '2E4057')
for actor, beneficio in actores:
    add_table_row(t2, [actor, beneficio], bold_first=True)

# ── 4. MÓDULOS Y FUNCIONALIDADES ─────────────
add_heading(doc, '4. Módulos y Funcionalidades del Sistema', 1)

p = doc.add_paragraph(
    'Al completar todas las etapas de desarrollo, ITO Cloud contará con los siguientes 14 módulos funcionales:'
)
p.style.font.size = Pt(10)

modulos = [
    ('1', 'Identity & Tenants',
     ['Login seguro con ASP.NET Identity + JWT',
      'Gestión de empresas (tenants) en modelo multiempresa SaaS',
      'Roles del sistema: Administrador, Supervisor, Inspector, Contratista',
      'Aislamiento total de datos por empresa mediante tenant_id',
      'Planes de suscripción configurables por tenant']),

    ('2', 'Empresas & Proyectos',
     ['CRUD completo de empresa (razón social, RUT, logo, contactos)',
      'CRUD de proyectos de construcción con información detallada',
      'Gestión de etapas del proyecto (ej: obra gruesa, terminaciones, etc.)',
      'Sectores y unidades de inspección dentro de cada proyecto',
      'Gestión de contratistas vinculados a cada proyecto']),

    ('3', 'Usuarios & Permisos',
     ['Administración de usuarios por empresa',
      'Asignación de roles globales y roles específicos por proyecto',
      'Invitación de usuarios por correo electrónico',
      'Gestión de acceso: activar, desactivar, reasignar',
      'Historial de acciones por usuario (auditoría)']),

    ('4', 'Plantillas de Inspección',
     ['Builder dinámico para crear formularios de inspección personalizados',
      'Tipos de pregunta: Sí/No, Selección múltiple, Numérica, Texto libre, Rango/puntuación',
      'Organización por secciones y subsecciones',
      'Versionado de plantillas: cada modificación genera una nueva versión',
      'Biblioteca de plantillas reutilizables por empresa',
      'Clonación y edición de plantillas existentes']),

    ('5', 'Programación de Inspecciones',
     ['Creación y agendamiento de inspecciones por proyecto y sector',
      'Asignación de ITO/inspector responsable',
      'Selección de plantilla de inspección a aplicar',
      'Calendario de inspecciones programadas',
      'Notificaciones automáticas al inspector asignado',
      'Estados: Programada → En ejecución → Completada → Cerrada']),

    ('6', 'Ejecución en Terreno (App Android)',
     ['Aplicación Android nativa con modo offline-first (Room + WorkManager)',
      'Lista de inspecciones asignadas al inspector autenticado',
      'Ejecución completa de checklist digital sin necesidad de conexión a internet',
      'Captura de fotos directamente asociadas a cada respuesta del formulario',
      'Geolocalización automática al registrar la inspección',
      'Sincronización automática al recuperar conectividad (WorkManager)',
      'Resolución de conflictos automática (server-wins con timestamps)',
      'Compresión de imágenes antes del upload para optimizar almacenamiento']),

    ('7', 'Observaciones / No Conformidades',
     ['Registro de observaciones vinculadas a inspecciones específicas',
      'Clasificación por severidad: Crítica, Mayor, Menor, Observación',
      'Asignación de responsable (contratista o persona específica)',
      'Fecha límite de corrección configurable',
      'Estados: Abierta → En corrección → Resuelta → Cerrada / Rechazada',
      'Comentarios y adjuntos de evidencia de corrección',
      'Historial completo de cambios de estado con timestamps y usuario responsable']),

    ('8', 'Reinspecciones',
     ['Flujo formal de validación de correcciones a observaciones',
      'Asignación de inspector para reinspección',
      'Vinculación directa con la observación original',
      'Estados: Aprobada / Rechazada con justificación',
      'Trazabilidad completa: inspección → observación → reinspección',
      'Historial de reinspecciones por observación']),

    ('9', 'Gestión de Documentos',
     ['Repositorio de documentos adjuntos por proyecto',
      'Tipos de documento: planos, contratos, especificaciones técnicas, permisos',
      'Versionado de documentos',
      'Almacenamiento en MinIO (compatible S3), sin límite de vendor',
      'Control de acceso por rol y proyecto',
      'Descarga y previsualización desde el portal web']),

    ('10', 'Reportes',
     ['Informe PDF de inspección completo generado automáticamente (QuestPDF)',
      'Contenido del PDF: datos del proyecto, inspector, fecha, respuestas, fotos, observaciones',
      'Exportación a Excel con datos tabulados por inspección (ClosedXML)',
      'Reporte consolidado de observaciones por proyecto y período',
      'Reporte de rendimiento por contratista',
      'Descarga directa desde el portal web']),

    ('11', 'Dashboard Ejecutivo',
     ['Panel de KPIs en tiempo real vía SignalR (Blazor Server nativo)',
      'Indicadores: total de inspecciones, tasa de cumplimiento, NC abiertas/cerradas',
      'Gráficos de tendencia de observaciones por período',
      'Ranking de sectores / unidades con más NC',
      'Estado de avance de inspecciones por proyecto',
      'Filtros dinámicos por proyecto, período, inspector y contratista',
      'Widgets configurables por rol de usuario']),

    ('12', 'Notificaciones',
     ['Notificaciones por correo electrónico (SMTP)',
      'Alertas al inspector al ser asignado a una inspección',
      'Notificación al contratista al recibir una observación asignada',
      'Recordatorios automáticos de observaciones próximas a vencer',
      'Notificaciones de sincronización y errores en app Android',
      'Centro de notificaciones dentro del portal web']),

    ('13', 'Módulo IA (Etapa futura)',
     ['Interfaces y arquitectura preparadas para integración futura de IA',
      'Generación automática de descripciones de observaciones',
      'Clasificación automática de severidad de NC mediante NLP',
      'Resumen automático de informes de inspección',
      'Búsqueda semántica en el historial de observaciones',
      'Integración lista para conectar con Claude API u OpenAI']),

    ('14', 'Auditoría del Sistema',
     ['Registro inmutable de todas las acciones del sistema',
      'Trazabilidad: quién creó, modificó o eliminó cada registro',
      'Campos de auditoría en todas las tablas: created_by, updated_by, deleted_by + timestamps',
      'Soft delete en todas las entidades (sin pérdida de datos históricos)',
      'Log de acceso por usuario y tenant',
      'Historial de cambios de estado en inspecciones y observaciones']),
]

for num, nombre, funcs in modulos:
    add_heading(doc, f'4.{num}. Módulo {num}: {nombre}', 2)
    for f in funcs:
        bp = doc.add_paragraph(f, style='List Bullet')
        bp.runs[0].font.size = Pt(10)

# ── 5. COMPONENTES TECNOLÓGICOS ───────────────
add_heading(doc, '5. Componentes y Stack Tecnológico', 1)

stack = [
    ('Backend API',        'ASP.NET Core 8',              'REST API + JWT, Clean Architecture, CQRS con MediatR'),
    ('Frontend Web',       'Blazor Server + MudBlazor',   'SPA en servidor, dashboards en tiempo real con SignalR'),
    ('App Móvil',          'Android (Kotlin + Compose)',  'Offline-first, Room DB local, sincronización WorkManager'),
    ('Base de Datos',      'PostgreSQL 16',               'Multi-tenant lógico, EF Core 8, migraciones automáticas'),
    ('Almacenamiento',     'MinIO (S3 compatible)',        'Fotos de inspección, documentos, auto-hosteable'),
    ('Autenticación',      'ASP.NET Identity + JWT',      'Sin dependencias externas, control total'),
    ('Reportes PDF',       'QuestPDF',                    'Informes de inspección generados automáticamente'),
    ('Reportes Excel',     'ClosedXML',                   'Exportación tabular de datos, licencia MIT'),
    ('Infraestructura',    'Docker + docker-compose',     'Despliegue reproducible en cualquier servidor Linux'),
    ('Caché / Sesiones',   'Redis',                       'Escalabilidad horizontal de Blazor Server'),
    ('Logging',            'Serilog',                     'Logs estructurados, configurable por sinks'),
    ('Validación',         'FluentValidation',            'Validación expresiva y desacoplada del modelo'),
]

t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Table Grid'
add_table_header_row(t3, ['Componente', 'Tecnología', 'Propósito'], '2E4057')
for comp, tec, prop in stack:
    add_table_row(t3, [comp, tec, prop], bold_first=True)

# ── 6. ARQUITECTURA ──────────────────────────
add_heading(doc, '6. Arquitectura del Sistema', 1)

p = doc.add_paragraph(
    'ITO Cloud implementa Clean Architecture con separación en 4 capas:'
)
p.style.font.size = Pt(10)

capas = [
    ('ITO.Cloud.Domain',          'Entidades de negocio, interfaces, enums, value objects. Sin dependencias externas.'),
    ('ITO.Cloud.Application',     'Casos de uso con CQRS (Commands/Queries), validadores FluentValidation, DTOs, mappers.'),
    ('ITO.Cloud.Infrastructure',  'Implementaciones: EF Core + PostgreSQL, MinIO, SMTP, JWT, servicios externos.'),
    ('ITO.Cloud.API',             'Controladores REST, middlewares (tenant, excepciones), configuración de la app.'),
    ('ITO.Cloud.Web',             'Páginas y componentes Blazor Server para el portal de gestión.'),
    ('ITO.Cloud.Android',         'App Kotlin con Jetpack Compose, Room DB local, sincronización WorkManager.'),
]

t4 = doc.add_table(rows=1, cols=2)
t4.style = 'Table Grid'
add_table_header_row(t4, ['Capa / Proyecto', 'Responsabilidad'], '2E4057')
for capa, resp in capas:
    add_table_row(t4, [capa, resp], bold_first=True)

doc.add_paragraph()
p2 = doc.add_paragraph(
    'El modelo multi-tenant es lógico: base de datos compartida con discriminador tenant_id en todas '
    'las tablas y Global Query Filters en EF Core que garantizan el aislamiento automático de datos '
    'entre empresas. Todos los PKs son UUID, lo que permite la creación de registros en modo offline '
    'desde Android sin riesgo de colisión al sincronizar.'
)
p2.style.font.size = Pt(10)

# ── 7. ROADMAP ───────────────────────────────
add_heading(doc, '7. Roadmap de Desarrollo', 1)

roadmap = [
    ('Etapa 1', 'Arquitectura y Diseño',          '✅ Completada', 'Decisiones técnicas, ADRs, estructura del proyecto, arquitectura documentada'),
    ('Etapa 2', 'Dominio y Base de Datos',        '🔄 En curso',   'Entidades de dominio, EF Core, migrations, modelo PostgreSQL completo'),
    ('Etapa 3', 'Backend API',                    '⏳ Pendiente',  'CQRS completo, endpoints REST, validaciones, auth, multi-tenant'),
    ('Etapa 4', 'Frontend Web (Blazor)',           '⏳ Pendiente',  'Portal de gestión: dashboard, proyectos, plantillas, inspecciones, reportes'),
    ('Etapa 5', 'App Android',                    '⏳ Pendiente',  'Login, inspecciones asignadas, ejecución offline, fotos, sincronización'),
    ('Etapa 6', 'Reportes',                       '⏳ Pendiente',  'PDFs con QuestPDF, Excel con ClosedXML, reportes ejecutivos'),
    ('Etapa 7', 'Infraestructura y Despliegue',   '⏳ Pendiente',  'Docker compose, Redis, MinIO, CI/CD, deploy en servidor Linux'),
    ('Etapa 8', 'Módulo IA',                      '⏳ Futuro',     'Integración con modelo de lenguaje para clasificación y generación automática'),
]

t5 = doc.add_table(rows=1, cols=4)
t5.style = 'Table Grid'
add_table_header_row(t5, ['Etapa', 'Nombre', 'Estado', 'Alcance'], '2E4057')
for et, nom, est, alc in roadmap:
    add_table_row(t5, [et, nom, est, alc], bold_first=True)

# ── 8. SEGURIDAD ─────────────────────────────
add_heading(doc, '8. Seguridad y Cumplimiento', 1)

seguridad = [
    'Autenticación con JWT firmado, sin estado en servidor (stateless API)',
    'Multi-tenant seguro: Global Query Filters en EF Core impiden acceso cruzado entre empresas',
    'Tenant ID incluido en JWT firmado, no manipulable por el cliente',
    'Soft delete en todas las tablas: no se eliminan datos, se marcan como eliminados',
    'Auditoría completa: toda acción queda registrada con usuario, fecha y acción',
    'HTTPS obligatorio en todos los endpoints',
    'Validación de entradas con FluentValidation en todos los comandos',
    'Almacenamiento de archivos en MinIO con acceso controlado por rutas firmadas',
    'Roles y permisos granulares por proyecto y por módulo',
]

for s in seguridad:
    bp = doc.add_paragraph(s, style='List Bullet')
    bp.runs[0].font.size = Pt(10)

# ── 9. CAPACIDADES OFFLINE ───────────────────
add_heading(doc, '9. Capacidades Offline (App Android)', 1)

p = doc.add_paragraph(
    'La aplicación Android fue diseñada con arquitectura offline-first para garantizar la operatividad '
    'en obras con conectividad limitada o nula. Las capacidades offline incluyen:'
)
p.style.font.size = Pt(10)

offline = [
    'Descarga y almacenamiento local de proyectos, plantillas e inspecciones asignadas',
    'Ejecución completa de checklist sin conexión a internet (Room DB local / SQLite)',
    'Captura y almacenamiento local de fotografías de evidencia',
    'Cola de sincronización (SyncQueue) que persiste todas las operaciones pendientes',
    'Sincronización automática al detectar conectividad (WorkManager)',
    'Resolución automática de conflictos con estrategia server-wins por timestamp',
    'UUID como PK en todas las entidades permite creación offline sin colisión',
    'WorkManager garantiza la ejecución de la sincronización incluso si la app se cierra',
]

for o in offline:
    bp = doc.add_paragraph(o, style='List Bullet')
    bp.runs[0].font.size = Pt(10)

# ── 10. CONCLUSIÓN ───────────────────────────
add_heading(doc, '10. Conclusión', 1)

p = doc.add_paragraph(
    'ITO Cloud es una solución integral y de nivel enterprise para la digitalización del proceso de '
    'Inspección Técnica de Obras. Al completar sus 8 etapas de desarrollo, la plataforma ofrecerá:'
)
p.style.font.size = Pt(10)

conclusiones = [
    '14 módulos funcionales completamente integrados',
    'Portal web de gestión con dashboard ejecutivo en tiempo real',
    'Aplicación Android con capacidad de trabajo offline total',
    'API REST documentada y segura para futuras integraciones',
    'Arquitectura multi-empresa escalable con aislamiento de datos garantizado',
    'Reportes automáticos en PDF y Excel',
    'Infraestructura auto-hosteada sin dependencia de servicios cloud de terceros',
    'Arquitectura preparada para integración futura de Inteligencia Artificial',
]

for c in conclusiones:
    bp = doc.add_paragraph(c, style='List Bullet')
    bp.runs[0].font.size = Pt(10)

doc.add_paragraph()
p2 = doc.add_paragraph(
    'La plataforma está construida sobre tecnologías modernas de código abierto, con licencias MIT '
    'donde es posible, garantizando independencia tecnológica y control total del producto.'
)
p2.style.font.size = Pt(10)

# ── Guardar ───────────────────────────────────
output_path = r'F:\SISTEMA DE INSPECCION DE OBRA\docs\ITO_Cloud_Documento_Ejecutivo_Funcionalidades.docx'
doc.save(output_path)
print(f'Documento generado: {output_path}')
