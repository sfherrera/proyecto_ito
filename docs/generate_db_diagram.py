# -*- coding: utf-8 -*-
"""
Generador del Diagrama de Base de Datos ITO Cloud
Produce: imagen ERD (PNG) + documento Word con tablas de entidades
"""

import os
import datetime
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
import matplotlib.patheffects as pe

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────
# DEFINICION DEL MODELO
# ─────────────────────────────────────────────────────────────

# Cada entidad: (nombre_tabla, color_hex, grupo, [(col, tipo, constraints, descripcion)])
# PK/FK/NN = Primary Key / Foreign Key / Not Null

AUDIT_FIELDS = [
    ("tenant_id",   "UUID",          "FK, NN, IX",   "Aislamiento multi-tenant"),
    ("created_at",  "TIMESTAMPTZ",   "NN DEFAULT NOW()", "Fecha de creacion"),
    ("created_by",  "UUID",          "NN, FK→users", "Usuario que creo el registro"),
    ("updated_at",  "TIMESTAMPTZ",   "NULLABLE",      "Ultima modificacion"),
    ("updated_by",  "UUID",          "NULLABLE, FK→users", "Usuario que modifico"),
    ("deleted_at",  "TIMESTAMPTZ",   "NULLABLE",      "Fecha de borrado logico"),
    ("deleted_by",  "UUID",          "NULLABLE, FK→users", "Usuario que elimino"),
    ("is_deleted",  "BOOLEAN",       "NN DEFAULT false", "Soft delete flag"),
]

ENTITIES = [

    # ── IDENTIDAD ──────────────────────────────────────────────
    {
        "table": "tenants",
        "group": "Identidad",
        "color": "#1565C0",
        "description": "Empresa / organizacion que contrata el SaaS. Raiz del multi-tenancy.",
        "columns": [
            ("id",        "UUID",         "PK",          "Identificador unico"),
            ("slug",      "VARCHAR(50)",  "NN, UNIQUE",  "Identificador URL amigable"),
            ("nombre",    "VARCHAR(200)", "NN",          "Nombre de la empresa"),
            ("plan",      "VARCHAR(50)",  "NN",          "Plan SaaS: free / pro / enterprise"),
            ("activo",    "BOOLEAN",      "NN DEFAULT true", "Tenant habilitado"),
            ("created_at","TIMESTAMPTZ",  "NN DEFAULT NOW()", "Fecha de alta"),
        ]
    },
    {
        "table": "application_users",
        "group": "Identidad",
        "color": "#1565C0",
        "description": "Usuarios del sistema. Extiende ASP.NET Identity.",
        "columns": [
            ("id",          "UUID",         "PK",          "Identificador unico"),
            ("tenant_id",   "UUID",         "FK→tenants, NN, IX", "Tenant al que pertenece"),
            ("email",       "VARCHAR(256)", "NN, UNIQUE",  "Email de login"),
            ("nombre",      "VARCHAR(100)", "NN",          "Nombre del usuario"),
            ("apellido",    "VARCHAR(100)", "NN",          "Apellido del usuario"),
            ("rol",         "VARCHAR(50)",  "NN",          "Admin / Supervisor / Inspector / Contratista"),
            ("activo",      "BOOLEAN",      "NN DEFAULT true", "Cuenta habilitada"),
            ("created_at",  "TIMESTAMPTZ",  "NN",          "Fecha de registro"),
            ("updated_at",  "TIMESTAMPTZ",  "NULLABLE",    "Ultima actualizacion"),
        ]
    },
    {
        "table": "project_user_assignments",
        "group": "Identidad",
        "color": "#1565C0",
        "description": "Asignacion de usuarios a proyectos especificos con un rol en ese proyecto.",
        "columns": [
            ("id",          "UUID",    "PK",              "Identificador unico"),
            ("tenant_id",   "UUID",    "FK→tenants, NN",  "Aislamiento multi-tenant"),
            ("project_id",  "UUID",    "FK→projects, NN, IX", "Proyecto"),
            ("user_id",     "UUID",    "FK→users, NN, IX","Usuario asignado"),
            ("rol_proyecto","VARCHAR(50)","NN",            "Rol en este proyecto especifico"),
            ("created_at",  "TIMESTAMPTZ","NN",           "Fecha de asignacion"),
        ]
    },

    # ── EMPRESAS Y PROYECTOS ────────────────────────────────────
    {
        "table": "companies",
        "group": "Proyectos",
        "color": "#00695C",
        "description": "Empresas constructoras o inmobiliarias gestionadas por el tenant.",
        "columns": [
            ("id",              "UUID",         "PK",         "Identificador unico"),
            ("nombre",          "VARCHAR(200)", "NN",         "Razon social"),
            ("rut",             "VARCHAR(20)",  "NN",         "RUT o identificador fiscal"),
            ("direccion",       "VARCHAR(300)", "NULLABLE",   "Direccion de la empresa"),
            ("telefono",        "VARCHAR(30)",  "NULLABLE",   "Telefono de contacto"),
            ("email_contacto",  "VARCHAR(200)", "NULLABLE",   "Email de contacto"),
        ] + AUDIT_FIELDS
    },
    {
        "table": "projects",
        "group": "Proyectos",
        "color": "#00695C",
        "description": "Obra o proyecto de construccion.",
        "columns": [
            ("id",           "UUID",         "PK",              "Identificador unico"),
            ("company_id",   "UUID",         "FK→companies, NN","Empresa propietaria"),
            ("nombre",       "VARCHAR(200)", "NN",              "Nombre del proyecto"),
            ("descripcion",  "TEXT",         "NULLABLE",        "Descripcion"),
            ("ubicacion",    "VARCHAR(400)", "NULLABLE",        "Direccion de la obra"),
            ("latitud",      "NUMERIC(10,7)","NULLABLE",        "Coordenada GPS"),
            ("longitud",     "NUMERIC(10,7)","NULLABLE",        "Coordenada GPS"),
            ("fecha_inicio", "DATE",         "NULLABLE",        "Fecha inicio planificada"),
            ("fecha_fin",    "DATE",         "NULLABLE",        "Fecha fin planificada"),
            ("estado",       "VARCHAR(30)",  "NN DEFAULT 'Activo'", "Activo / Pausado / Cerrado"),
        ] + AUDIT_FIELDS
    },
    {
        "table": "project_stages",
        "group": "Proyectos",
        "color": "#00695C",
        "description": "Etapas del proyecto (ej: Fundacion, Estructura, Terminaciones).",
        "columns": [
            ("id",          "UUID",         "PK",              "Identificador unico"),
            ("project_id",  "UUID",         "FK→projects, NN, IX","Proyecto padre"),
            ("nombre",      "VARCHAR(100)", "NN",              "Nombre de la etapa"),
            ("orden",       "SMALLINT",     "NN",              "Orden de visualizacion"),
        ] + AUDIT_FIELDS
    },
    {
        "table": "project_sectors",
        "group": "Proyectos",
        "color": "#00695C",
        "description": "Sectores dentro de una etapa (ej: Torre A, Bloque 1).",
        "columns": [
            ("id",        "UUID",         "PK",              "Identificador unico"),
            ("stage_id",  "UUID",         "FK→project_stages, NN, IX","Etapa padre"),
            ("nombre",    "VARCHAR(100)", "NN",              "Nombre del sector"),
            ("orden",     "SMALLINT",     "NN",              "Orden de visualizacion"),
        ] + AUDIT_FIELDS
    },
    {
        "table": "project_units",
        "group": "Proyectos",
        "color": "#00695C",
        "description": "Unidades dentro de un sector (departamentos, casas, locales).",
        "columns": [
            ("id",          "UUID",        "PK",              "Identificador unico"),
            ("sector_id",   "UUID",        "FK→project_sectors, NN, IX","Sector padre"),
            ("nombre",      "VARCHAR(100)","NN",              "Nombre / numero de la unidad"),
            ("tipo",        "VARCHAR(50)", "NULLABLE",        "Depto / Casa / Local / Oficina"),
            ("piso",        "SMALLINT",    "NULLABLE",        "Numero de piso"),
        ] + AUDIT_FIELDS
    },
    {
        "table": "contractors",
        "group": "Proyectos",
        "color": "#00695C",
        "description": "Empresas contratistas que ejecutan trabajos en la obra.",
        "columns": [
            ("id",               "UUID",         "PK",         "Identificador unico"),
            ("nombre",           "VARCHAR(200)", "NN",         "Nombre del contratista"),
            ("rut",              "VARCHAR(20)",  "NULLABLE",   "RUT o identificador fiscal"),
            ("representante",    "VARCHAR(150)", "NULLABLE",   "Nombre del representante"),
            ("telefono",         "VARCHAR(30)",  "NULLABLE",   "Telefono de contacto"),
            ("email",            "VARCHAR(200)", "NULLABLE",   "Email de contacto"),
            ("especialidad",     "VARCHAR(100)", "NULLABLE",   "Esp: Electrico, Sanitario, etc."),
        ] + AUDIT_FIELDS
    },
    {
        "table": "project_contractors",
        "group": "Proyectos",
        "color": "#00695C",
        "description": "Relacion N:M entre proyectos y contratistas.",
        "columns": [
            ("id",             "UUID",  "PK",              "Identificador unico"),
            ("tenant_id",      "UUID",  "FK→tenants, NN",  "Aislamiento multi-tenant"),
            ("project_id",     "UUID",  "FK→projects, NN, IX", "Proyecto"),
            ("contractor_id",  "UUID",  "FK→contractors, NN, IX", "Contratista"),
            ("activo",         "BOOLEAN","NN DEFAULT true", "Relacion activa"),
            ("created_at",     "TIMESTAMPTZ","NN",          "Fecha de vinculacion"),
        ]
    },

    # ── PLANTILLAS ──────────────────────────────────────────────
    {
        "table": "inspection_templates",
        "group": "Plantillas",
        "color": "#6A1B9A",
        "description": "Plantillas maestras de inspeccion. Cada plantilla tiene versiones publicadas.",
        "columns": [
            ("id",              "UUID",         "PK",         "Identificador unico"),
            ("nombre",          "VARCHAR(200)", "NN",         "Nombre de la plantilla"),
            ("descripcion",     "TEXT",         "NULLABLE",   "Descripcion del proposito"),
            ("version_actual",  "SMALLINT",     "NN DEFAULT 1","Numero de version vigente"),
            ("publicada",       "BOOLEAN",      "NN DEFAULT false","Disponible para uso"),
        ] + AUDIT_FIELDS
    },
    {
        "table": "template_versions",
        "group": "Plantillas",
        "color": "#6A1B9A",
        "description": "Version especifica de una plantilla. Las inspecciones apuntan a una version fija.",
        "columns": [
            ("id",             "UUID",      "PK",                "Identificador unico"),
            ("template_id",    "UUID",      "FK→inspection_templates, NN, IX","Plantilla padre"),
            ("version",        "SMALLINT",  "NN",                "Numero de version"),
            ("publicada",      "BOOLEAN",   "NN DEFAULT false",  "Version disponible para uso"),
            ("fecha_publicacion","TIMESTAMPTZ","NULLABLE",       "Cuando se publico"),
            ("notas_version",  "TEXT",      "NULLABLE",          "Que cambio en esta version"),
        ] + AUDIT_FIELDS
    },
    {
        "table": "template_sections",
        "group": "Plantillas",
        "color": "#6A1B9A",
        "description": "Secciones dentro de una version de plantilla (ej: Estructura, Instalaciones).",
        "columns": [
            ("id",                 "UUID",         "PK",          "Identificador unico"),
            ("template_version_id","UUID",         "FK→template_versions, NN, IX","Version padre"),
            ("nombre",             "VARCHAR(200)", "NN",          "Nombre de la seccion"),
            ("descripcion",        "TEXT",         "NULLABLE",    "Descripcion de la seccion"),
            ("orden",              "SMALLINT",     "NN",          "Orden de visualizacion"),
        ] + AUDIT_FIELDS
    },
    {
        "table": "template_questions",
        "group": "Plantillas",
        "color": "#6A1B9A",
        "description": "Preguntas dentro de una seccion. Soporta multiples tipos de respuesta.",
        "columns": [
            ("id",          "UUID",         "PK",          "Identificador unico"),
            ("section_id",  "UUID",         "FK→template_sections, NN, IX","Seccion padre"),
            ("texto",       "TEXT",         "NN",          "Enunciado de la pregunta"),
            ("tipo",        "VARCHAR(30)",  "NN",          "SiNo / Texto / Numerico / Seleccion / Lista"),
            ("requerida",   "BOOLEAN",      "NN DEFAULT true","Respuesta obligatoria"),
            ("orden",       "SMALLINT",     "NN",          "Orden dentro de la seccion"),
            ("opciones",    "JSONB",        "NULLABLE",    "Opciones para tipo Seleccion/Lista"),
            ("valor_min",   "NUMERIC",      "NULLABLE",    "Minimo para tipo Numerico"),
            ("valor_max",   "NUMERIC",      "NULLABLE",    "Maximo para tipo Numerico"),
        ] + AUDIT_FIELDS
    },

    # ── INSPECCIONES ────────────────────────────────────────────
    {
        "table": "inspections",
        "group": "Inspecciones",
        "color": "#B71C1C",
        "description": "Registro central de cada inspeccion programada o ejecutada.",
        "columns": [
            ("id",                  "UUID",      "PK",              "Identificador unico"),
            ("project_id",          "UUID",      "FK→projects, NN, IX","Proyecto"),
            ("stage_id",            "UUID",      "FK→project_stages, NULLABLE","Etapa inspeccionada"),
            ("sector_id",           "UUID",      "FK→project_sectors, NULLABLE","Sector inspeccionado"),
            ("unit_id",             "UUID",      "FK→project_units, NULLABLE","Unidad inspeccionada"),
            ("template_version_id", "UUID",      "FK→template_versions, NN","Version de plantilla usada"),
            ("inspector_id",        "UUID",      "FK→users, NN, IX","Inspector asignado"),
            ("fecha_programada",    "TIMESTAMPTZ","NN",             "Fecha/hora planificada"),
            ("fecha_inicio",        "TIMESTAMPTZ","NULLABLE",       "Inicio real de la inspeccion"),
            ("fecha_fin",           "TIMESTAMPTZ","NULLABLE",       "Fin real de la inspeccion"),
            ("estado",              "VARCHAR(30)","NN",             "Programada/EnEjecucion/Ejecutada/Cancelada"),
            ("observaciones_generales","TEXT",   "NULLABLE",       "Notas generales del inspector"),
            ("latitud",             "NUMERIC(10,7)","NULLABLE",     "GPS al momento de inspeccion"),
            ("longitud",            "NUMERIC(10,7)","NULLABLE",     "GPS al momento de inspeccion"),
        ] + AUDIT_FIELDS
    },
    {
        "table": "inspection_answers",
        "group": "Inspecciones",
        "color": "#B71C1C",
        "description": "Respuestas a cada pregunta de la inspeccion.",
        "columns": [
            ("id",               "UUID",    "PK",              "Identificador unico"),
            ("inspection_id",    "UUID",    "FK→inspections, NN, IX","Inspeccion padre"),
            ("question_id",      "UUID",    "FK→template_questions, NN","Pregunta respondida"),
            ("respuesta_bool",   "BOOLEAN", "NULLABLE",        "Respuesta Si/No"),
            ("respuesta_texto",  "TEXT",    "NULLABLE",        "Respuesta texto libre"),
            ("respuesta_numero", "NUMERIC", "NULLABLE",        "Respuesta numerica"),
            ("respuesta_opciones","JSONB",  "NULLABLE",        "Respuesta seleccion multiple"),
            ("conforme",         "BOOLEAN", "NULLABLE",        "Marca rapida de conformidad"),
            ("comentario",       "TEXT",    "NULLABLE",        "Comentario adicional"),
        ] + AUDIT_FIELDS
    },
    {
        "table": "inspection_evidences",
        "group": "Inspecciones",
        "color": "#B71C1C",
        "description": "Archivos adjuntos (fotos, videos) vinculados a una respuesta o inspeccion.",
        "columns": [
            ("id",             "UUID",         "PK",              "Identificador unico"),
            ("inspection_id",  "UUID",         "FK→inspections, NN, IX","Inspeccion"),
            ("answer_id",      "UUID",         "FK→inspection_answers, NULLABLE","Respuesta vinculada"),
            ("file_path",      "VARCHAR(500)", "NN",              "Ruta en MinIO"),
            ("nombre_archivo", "VARCHAR(200)", "NN",              "Nombre original del archivo"),
            ("mime_type",      "VARCHAR(100)", "NN",              "image/jpeg, image/png, etc."),
            ("tamano_bytes",   "BIGINT",       "NN",              "Tamano del archivo"),
            ("descripcion",    "TEXT",         "NULLABLE",        "Descripcion de la evidencia"),
        ] + AUDIT_FIELDS
    },

    # ── OBSERVACIONES ───────────────────────────────────────────
    {
        "table": "observations",
        "group": "Observaciones",
        "color": "#E65100",
        "description": "No conformidades u observaciones registradas durante o despues de la inspeccion.",
        "columns": [
            ("id",              "UUID",         "PK",              "Identificador unico"),
            ("inspection_id",   "UUID",         "FK→inspections, NN, IX","Inspeccion origen"),
            ("answer_id",       "UUID",         "FK→inspection_answers, NULLABLE","Respuesta vinculada"),
            ("descripcion",     "TEXT",         "NN",              "Descripcion de la observacion"),
            ("severidad",       "VARCHAR(20)",  "NN",              "Leve / Grave / Critico"),
            ("estado",          "VARCHAR(30)",  "NN",              "Abierta/EnCorreccion/Resuelta/Cerrada/Rechazada"),
            ("contractor_id",   "UUID",         "FK→contractors, NULLABLE","Contratista responsable"),
            ("fecha_limite",    "DATE",         "NULLABLE",        "Plazo maximo de correccion"),
            ("fecha_cierre",    "TIMESTAMPTZ",  "NULLABLE",        "Cuando se cerro definitivamente"),
            ("notas_cierre",    "TEXT",         "NULLABLE",        "Descripcion del cierre"),
        ] + AUDIT_FIELDS
    },
    {
        "table": "observation_histories",
        "group": "Observaciones",
        "color": "#E65100",
        "description": "Historial inmutable de cambios de estado de una observacion.",
        "columns": [
            ("id",               "UUID",    "PK",              "Identificador unico"),
            ("tenant_id",        "UUID",    "FK→tenants, NN",  "Aislamiento multi-tenant"),
            ("observation_id",   "UUID",    "FK→observations, NN, IX","Observacion"),
            ("estado_anterior",  "VARCHAR(30)","NULLABLE",     "Estado previo al cambio"),
            ("estado_nuevo",     "VARCHAR(30)","NN",           "Nuevo estado"),
            ("comentario",       "TEXT",    "NULLABLE",        "Comentario del responsable"),
            ("created_at",       "TIMESTAMPTZ","NN DEFAULT NOW()","Timestamp del cambio"),
            ("created_by",       "UUID",    "NN, FK→users",    "Usuario que hizo el cambio"),
        ]
    },
    {
        "table": "reinspections",
        "group": "Observaciones",
        "color": "#E65100",
        "description": "Reinspeccion asociada a una observacion para verificar la correccion.",
        "columns": [
            ("id",               "UUID",         "PK",              "Identificador unico"),
            ("observation_id",   "UUID",         "FK→observations, NN, IX","Observacion a verificar"),
            ("inspector_id",     "UUID",         "FK→users, NN",    "Inspector asignado"),
            ("fecha_programada", "TIMESTAMPTZ",  "NN",              "Fecha planificada"),
            ("fecha_ejecucion",  "TIMESTAMPTZ",  "NULLABLE",        "Fecha real de ejecucion"),
            ("resultado",        "VARCHAR(30)",  "NULLABLE",        "Aprobada / Rechazada"),
            ("comentario",       "TEXT",         "NULLABLE",        "Observaciones del reinspector"),
        ] + AUDIT_FIELDS
    },

    # ── DOCUMENTOS ──────────────────────────────────────────────
    {
        "table": "project_documents",
        "group": "Documentos",
        "color": "#1B5E20",
        "description": "Documentos adjuntos al proyecto: planos, contratos, especificaciones.",
        "columns": [
            ("id",             "UUID",         "PK",              "Identificador unico"),
            ("project_id",     "UUID",         "FK→projects, NN, IX","Proyecto"),
            ("nombre",         "VARCHAR(200)", "NN",              "Nombre del documento"),
            ("tipo",           "VARCHAR(80)",  "NULLABLE",        "Plano / Contrato / Especificacion / Otro"),
            ("file_path",      "VARCHAR(500)", "NN",              "Ruta en MinIO"),
            ("mime_type",      "VARCHAR(100)", "NN",              "Tipo MIME del archivo"),
            ("tamano_bytes",   "BIGINT",       "NN",              "Tamano del archivo"),
            ("descripcion",    "TEXT",         "NULLABLE",        "Descripcion del documento"),
        ] + AUDIT_FIELDS
    },
]

# ─────────────────────────────────────────────────────────────
# RELACIONES (para el diagrama visual)
# ─────────────────────────────────────────────────────────────

RELATIONS = [
    # (tabla_origen, tabla_destino, cardinalidad, etiqueta)
    ("application_users",    "tenants",              "N:1",  "pertenece a"),
    ("project_user_assignments", "projects",         "N:1",  "asignado a"),
    ("project_user_assignments", "application_users","N:1",  "usuario"),
    ("companies",            "tenants",              "N:1",  "tenant"),
    ("projects",             "companies",            "N:1",  "pertenece a"),
    ("project_stages",       "projects",             "N:1",  "pertenece a"),
    ("project_sectors",      "project_stages",       "N:1",  "pertenece a"),
    ("project_units",        "project_sectors",      "N:1",  "pertenece a"),
    ("project_contractors",  "projects",             "N:1",  "en proyecto"),
    ("project_contractors",  "contractors",          "N:1",  "contratista"),
    ("template_versions",    "inspection_templates", "N:1",  "version de"),
    ("template_sections",    "template_versions",    "N:1",  "seccion de"),
    ("template_questions",   "template_sections",    "N:1",  "pregunta de"),
    ("inspections",          "projects",             "N:1",  "en proyecto"),
    ("inspections",          "template_versions",    "N:1",  "usa plantilla"),
    ("inspections",          "application_users",    "N:1",  "inspector"),
    ("inspection_answers",   "inspections",          "N:1",  "respuesta de"),
    ("inspection_answers",   "template_questions",   "N:1",  "responde a"),
    ("inspection_evidences", "inspections",          "N:1",  "evidencia de"),
    ("inspection_evidences", "inspection_answers",   "N:1",  "adjunto a"),
    ("observations",         "inspections",          "N:1",  "originada en"),
    ("observations",         "contractors",          "N:1",  "responsable"),
    ("observation_histories","observations",         "N:1",  "historial de"),
    ("reinspections",        "observations",         "N:1",  "verifica"),
    ("reinspections",        "application_users",    "N:1",  "inspector"),
    ("project_documents",    "projects",             "N:1",  "adjunto a"),
]

# ─────────────────────────────────────────────────────────────
# POSICIONES para el diagrama visual (x, y) en unidades de plot
# ─────────────────────────────────────────────────────────────

POSITIONS = {
    "tenants":                  (2.0,  17.0),
    "application_users":        (6.0,  17.0),
    "project_user_assignments": (6.0,  14.5),

    "companies":                (0.0,  13.5),
    "projects":                 (3.5,  13.5),
    "project_stages":           (3.5,  11.0),
    "project_sectors":          (3.5,   8.5),
    "project_units":            (3.5,   6.0),
    "contractors":              (7.5,  11.0),
    "project_contractors":      (7.5,  13.5),

    "inspection_templates":     (11.5, 17.0),
    "template_versions":        (11.5, 14.5),
    "template_sections":        (11.5, 12.0),
    "template_questions":       (11.5,  9.5),

    "inspections":              (7.5,   8.5),
    "inspection_answers":       (7.5,   6.0),
    "inspection_evidences":     (7.5,   3.5),

    "observations":             (3.5,   3.5),
    "observation_histories":    (0.5,   1.5),
    "reinspections":            (3.5,   1.0),

    "project_documents":        (0.0,  10.5),
}

GROUP_COLORS = {
    "Identidad":     "#1565C0",
    "Proyectos":     "#00695C",
    "Plantillas":    "#6A1B9A",
    "Inspecciones":  "#B71C1C",
    "Observaciones": "#E65100",
    "Documentos":    "#1B5E20",
}

# ─────────────────────────────────────────────────────────────
# GENERADOR DE IMAGEN ERD
# ─────────────────────────────────────────────────────────────

BOX_W = 2.8   # ancho de la caja
BOX_H = 0.7   # alto de la caja

def draw_erd(output_path):
    fig_w, fig_h = 16, 20
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.set_xlim(-0.5, 15.0)
    ax.set_ylim(-0.5, 19.0)
    ax.axis('off')
    ax.set_facecolor('#F8F9FA')
    fig.patch.set_facecolor('#F8F9FA')

    # Titulo
    ax.text(7.0, 18.6, "ITO Cloud — Diagrama Entidad-Relacion",
            ha='center', va='center', fontsize=14, fontweight='bold',
            color='#1B3A5C')
    ax.text(7.0, 18.25, f"Base de datos PostgreSQL 16  |  {datetime.date.today().strftime('%d/%m/%Y')}",
            ha='center', va='center', fontsize=9, color='#666666', style='italic')

    # Leyenda de grupos
    legend_x = 0.0
    for i, (group, color) in enumerate(GROUP_COLORS.items()):
        rect = mpatches.FancyBboxPatch(
            (legend_x + i * 2.55, 18.0), 2.3, 0.3,
            boxstyle="round,pad=0.05", linewidth=1,
            edgecolor='white', facecolor=color, alpha=0.85
        )
        ax.add_patch(rect)
        ax.text(legend_x + i * 2.55 + 1.15, 18.15, group,
                ha='center', va='center', fontsize=7.5, color='white', fontweight='bold')

    # Dibujar relaciones (flechas)
    for (src, dst, card, label) in RELATIONS:
        if src not in POSITIONS or dst not in POSITIONS:
            continue
        sx, sy = POSITIONS[src]
        dx, dy = POSITIONS[dst]
        cx = sx + BOX_W / 2
        cy = sy + BOX_H / 2
        ex = dx + BOX_W / 2
        ey = dy + BOX_H / 2

        ax.annotate("",
            xy=(ex, ey), xytext=(cx, cy),
            arrowprops=dict(
                arrowstyle="-|>",
                color='#90A4AE',
                lw=0.9,
                connectionstyle="arc3,rad=0.0",
            )
        )

    # Dibujar cajas de entidades
    for entity in ENTITIES:
        table = entity["table"]
        group = entity["group"]
        color = entity["color"]

        if table not in POSITIONS:
            continue

        x, y = POSITIONS[table]

        # Sombra
        shadow = mpatches.FancyBboxPatch(
            (x + 0.05, y - 0.05), BOX_W, BOX_H,
            boxstyle="round,pad=0.07", linewidth=0,
            facecolor='#CCCCCC', alpha=0.5, zorder=1
        )
        ax.add_patch(shadow)

        # Caja principal
        box = mpatches.FancyBboxPatch(
            (x, y), BOX_W, BOX_H,
            boxstyle="round,pad=0.07", linewidth=1.5,
            edgecolor=color, facecolor='white', zorder=2
        )
        ax.add_patch(box)

        # Banda de color superior
        band = mpatches.FancyBboxPatch(
            (x, y + BOX_H * 0.55), BOX_W, BOX_H * 0.45,
            boxstyle="round,pad=0.04", linewidth=0,
            facecolor=color, zorder=3
        )
        ax.add_patch(band)

        # Nombre de la tabla
        display_name = table.replace('_', '_\n') if len(table) > 20 else table
        ax.text(x + BOX_W / 2, y + BOX_H * 0.78, table,
                ha='center', va='center',
                fontsize=7.2, fontweight='bold', color='white',
                zorder=4, fontfamily='monospace')

        # PK principal
        pk_cols = [c[0] for c in entity["columns"] if "PK" in c[2]]
        fk_cols = [c[0] for c in entity["columns"] if "FK" in c[2] and "PK" not in c[2]]
        other_cols = [c[0] for c in entity["columns"] if "PK" not in c[2] and "FK" not in c[2]]

        preview = []
        if pk_cols:
            preview.append(f"PK  {pk_cols[0]}")
        for fk in fk_cols[:3]:
            preview.append(f"FK  {fk}")
        for oc in other_cols[:max(0, 2 - len(fk_cols[:3]))]:
            preview.append(f"      {oc}")

        preview_text = "  |  ".join(preview[:2]) if preview else ""
        ax.text(x + BOX_W / 2, y + BOX_H * 0.25, preview_text,
                ha='center', va='center',
                fontsize=5.8, color='#333333',
                zorder=4, fontfamily='monospace')

    plt.tight_layout(pad=0.5)
    plt.savefig(output_path, dpi=180, bbox_inches='tight',
                facecolor='#F8F9FA', edgecolor='none')
    plt.close()
    print(f"[OK] Diagrama ERD generado: {output_path}")


# ─────────────────────────────────────────────────────────────
# UTILIDADES WORD
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_hr(doc, color='1B5E91'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)
    return p

def heading2(doc, text, color_hex=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    if color_hex:
        r = int(color_hex[1:3], 16)
        g = int(color_hex[3:5], 16)
        b = int(color_hex[5:7], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    else:
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    return p

def body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = italic
    return p

# ─────────────────────────────────────────────────────────────
# GENERADOR DOCUMENTO WORD
# ─────────────────────────────────────────────────────────────

def build_word_document(erd_image_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ── PORTADA ─────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ITO CLOUD")
    run.bold = True
    run.font.size = Pt(34)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Diagrama de Base de Datos")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"PostgreSQL 16  |  Version 1.0  |  {datetime.date.today().strftime('%d/%m/%Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    doc.add_paragraph()
    # Resumen en portada
    stats_tbl = doc.add_table(rows=1, cols=4)
    stats_tbl.style = 'Table Grid'
    stats_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(4):
        stats_tbl.columns[i].width = Cm(3.5)

    stats = [
        (str(len(ENTITIES)),   "Tablas"),
        (str(len(RELATIONS)),  "Relaciones"),
        (str(sum(len(e['columns']) for e in ENTITIES)), "Columnas totales"),
        (str(len(GROUP_COLORS)), "Grupos / Modulos"),
    ]
    row = stats_tbl.rows[0]
    for i, (val, lbl) in enumerate(stats):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_v = p.add_run(val + "\n")
        run_v.bold = True
        run_v.font.size = Pt(22)
        run_v.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)
        run_l = p.add_run(lbl)
        run_l.font.size = Pt(9)
        run_l.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        set_cell_bg(cell, 'EBF3FB')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_page_break()

    # ── SECCION 1: DIAGRAMA ERD ──────────────────
    heading1(doc, "1. Diagrama Entidad-Relacion (ERD)")
    add_hr(doc)
    body(doc,
        "El siguiente diagrama muestra las 20 tablas del sistema agrupadas por modulo funcional. "
        "Las flechas indican relaciones de clave foranea (muchos a uno). "
        "Todas las tablas de negocio incluyen los campos de auditoria estandar y tenant_id."
    )

    if os.path.exists(erd_image_path):
        doc.add_picture(erd_image_path, width=Cm(17.0))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Leyenda de grupos
    doc.add_paragraph()
    leyenda_tbl = doc.add_table(rows=1, cols=6)
    leyenda_tbl.style = 'Table Grid'
    leyenda_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(6):
        leyenda_tbl.columns[i].width = Cm(2.8)

    grp_row = leyenda_tbl.rows[0]
    for i, (group, color) in enumerate(GROUP_COLORS.items()):
        cell = grp_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(group)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hex_c = color.lstrip('#')
        set_cell_bg(cell, hex_c)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_page_break()

    # ── SECCION 2: CAMPOS DE AUDITORIA ESTANDAR ──
    heading1(doc, "2. Campos de Auditoria Estandar")
    add_hr(doc)
    body(doc,
        "Todas las tablas de negocio incluyen los siguientes campos de forma obligatoria. "
        "Son gestionados automaticamente por el AuditInterceptor de EF Core."
    )

    audit_tbl = doc.add_table(rows=1, cols=4)
    audit_tbl.style = 'Table Grid'
    audit_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    audit_tbl.columns[0].width = Cm(3.5)
    audit_tbl.columns[1].width = Cm(3.5)
    audit_tbl.columns[2].width = Cm(4.0)
    audit_tbl.columns[3].width = Cm(6.0)
    _header_row(audit_tbl, ['Campo', 'Tipo', 'Restricciones', 'Descripcion'])

    for i, (col, tipo, constr, desc) in enumerate(AUDIT_FIELDS):
        row = audit_tbl.add_row()
        vals = [col, tipo, constr, desc]
        bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = val
            set_cell_bg(cell, bg)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Consolas' if j < 2 else 'Calibri'
        for r in row.cells[0].paragraphs[0].runs:
            r.bold = True

    doc.add_paragraph()

    # ── SECCION 3: TABLAS POR MODULO ────────────
    heading1(doc, "3. Definicion de Tablas")
    add_hr(doc)

    current_group = None
    for entity in ENTITIES:
        group = entity["group"]
        table = entity["table"]
        color = entity["color"]
        desc  = entity["description"]
        cols  = entity["columns"]

        if group != current_group:
            current_group = group
            doc.add_paragraph()
            # Encabezado de grupo
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(f"  MODULO: {group.upper()}")
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            hex_c = color.lstrip('#')
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex_c)
            pPr.append(shd)

        # Nombre de la tabla
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(10)
        p2.paragraph_format.space_after  = Pt(2)
        run_t = p2.add_run(f"Tabla: {table}")
        run_t.bold = True
        run_t.font.size = Pt(11)
        hex_c = color.lstrip('#')
        r_v = int(hex_c[0:2], 16)
        g_v = int(hex_c[2:4], 16)
        b_v = int(hex_c[4:6], 16)
        run_t.font.color.rgb = RGBColor(r_v, g_v, b_v)
        run_t.font.name = 'Consolas'

        # Descripcion
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after  = Pt(4)
        run_d = p3.add_run(desc)
        run_d.font.size = Pt(9)
        run_d.italic = True
        run_d.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # Tabla de columnas
        col_tbl = doc.add_table(rows=1, cols=4)
        col_tbl.style = 'Table Grid'
        col_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        col_tbl.columns[0].width = Cm(4.0)
        col_tbl.columns[1].width = Cm(3.5)
        col_tbl.columns[2].width = Cm(4.0)
        col_tbl.columns[3].width = Cm(5.5)
        _header_row(col_tbl, ['Columna', 'Tipo', 'Restricciones', 'Descripcion'],
                    bg=hex_c)

        for i, (col_name, col_type, constraints, col_desc) in enumerate(cols):
            row = col_tbl.add_row()
            vals = [col_name, col_type, constraints, col_desc]

            # Color de fila segun tipo de columna
            if "PK" in constraints:
                bg = 'FFF9C4'  # amarillo suave para PK
            elif "FK" in constraints:
                bg = 'E8F5E9'  # verde suave para FK
            else:
                bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'

            for j, val in enumerate(vals):
                cell = row.cells[j]
                cell.text = val
                set_cell_bg(cell, bg)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(8.5)
                        r.font.name = 'Consolas' if j < 2 else 'Calibri'

            # PK en negrita
            if "PK" in constraints:
                for r in row.cells[0].paragraphs[0].runs:
                    r.bold = True

        doc.add_paragraph()

    # ── SECCION 4: RELACIONES ────────────────────
    doc.add_page_break()
    heading1(doc, "4. Tabla de Relaciones (Foreign Keys)")
    add_hr(doc)
    body(doc, f"Total: {len(RELATIONS)} relaciones de clave foranea en el modelo.")

    rel_tbl = doc.add_table(rows=1, cols=4)
    rel_tbl.style = 'Table Grid'
    rel_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rel_tbl.columns[0].width = Cm(5.0)
    rel_tbl.columns[1].width = Cm(5.0)
    rel_tbl.columns[2].width = Cm(2.0)
    rel_tbl.columns[3].width = Cm(5.0)
    _header_row(rel_tbl, ['Tabla Origen', 'Tabla Destino', 'Cardinalidad', 'Descripcion'])

    for i, (src, dst, card, label) in enumerate(RELATIONS):
        row = rel_tbl.add_row()
        vals = [src, dst, card, label]
        bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = val
            set_cell_bg(cell, bg)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Consolas' if j < 2 else 'Calibri'
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── SECCION 5: CONVENCIONES ──────────────────
    heading1(doc, "5. Convenciones y Patrones")
    add_hr(doc)

    convenciones = [
        ("Nombres de tablas",        "snake_case, plural (inspection_templates, project_units)"),
        ("Claves primarias",         "UUID v4 en todas las tablas. Necesario para generacion offline en Android sin colisiones."),
        ("Soft delete",              "Todas las tablas tienen is_deleted BOOLEAN + deleted_at TIMESTAMPTZ + deleted_by UUID."),
        ("Multi-tenancy",            "Todas las tablas de negocio tienen tenant_id UUID NOT NULL con indice. EF Core aplica Global Query Filter automatico."),
        ("Auditoria",                "created_at/by, updated_at/by, deleted_at/by en todas las entidades. Gestionado por AuditInterceptor."),
        ("Indices",                  "IX sobre todas las FK, tenant_id y columnas de busqueda frecuente (email, estado, fecha)."),
        ("JSONB",                    "Usado en template_questions.opciones e inspection_answers.respuesta_opciones para flexibilidad de plantillas dinamicas."),
        ("Timestamps",               "TIMESTAMPTZ (con timezone) en todos los campos de fecha/hora para correcta gestion de zonas horarias."),
        ("Resolucion de conflictos", "updated_at en cada entidad permite implementar estrategia server-wins para sincronizacion offline Android."),
    ]

    conv_tbl = doc.add_table(rows=1, cols=2)
    conv_tbl.style = 'Table Grid'
    conv_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    conv_tbl.columns[0].width = Cm(4.5)
    conv_tbl.columns[1].width = Cm(12.5)
    _header_row(conv_tbl, ['Convencion', 'Descripcion'])

    for i, (conv, desc) in enumerate(convenciones):
        row = conv_tbl.add_row()
        row.cells[0].text = conv
        row.cells[1].text = desc
        bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
        for c in row.cells:
            set_cell_bg(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        for r in row.cells[0].paragraphs[0].runs:
            r.bold = True

    doc.add_paragraph()

    # ── PIE ─────────────────────────────────────
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"ITO Cloud - Diagrama de Base de Datos v1.0  |  "
        f"{datetime.date.today().strftime('%d/%m/%Y')}  |  Confidencial"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    return doc


def _header_row(table, headers, bg='1B5E91'):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ─────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────

if __name__ == '__main__':
    base = r"F:\SISTEMA DE INSPECCION DE OBRA\docs"
    erd_png  = os.path.join(base, "ITO_Cloud_ERD.png")
    word_out = os.path.join(base, "ITO_Cloud_DB_Diagram.docx")

    print("[1/2] Generando diagrama ERD (imagen)...")
    draw_erd(erd_png)

    print("[2/2] Generando documento Word...")
    doc = build_word_document(erd_png)
    doc.save(word_out)

    print(f"[OK] Word generado: {word_out}")
    total_cols = sum(len(e['columns']) for e in ENTITIES)
    print(f"[OK] Tablas: {len(ENTITIES)}  |  Relaciones: {len(RELATIONS)}  |  Columnas totales: {total_cols}")
