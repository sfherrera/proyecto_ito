# -*- coding: utf-8 -*-
"""
Generador del Set de Pruebas QA - ITO Cloud Web
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# UTILIDADES (reutilizadas del generador anterior)
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_hr(doc, color='1B5E91'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)
    return p

def heading2(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    rgb = color or RGBColor(0x0D, 0x47, 0xA1)
    run.font.color.rgb = rgb
    return p

def body(doc, text, italic=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    return p

# ─────────────────────────────────────────────
# DATOS: CASOS DE PRUEBA
# ─────────────────────────────────────────────
# Estructura por caso:
# (id, modulo, nombre, prioridad, tipo, precondicion, pasos, resultado_esperado)
#
# Prioridad:  Alta / Media / Baja
# Tipo:       Funcional / Negativo / Borde / Seguridad / UI

TEST_CASES = [

    # ───────────────────────────────────────────
    # MOD-01  AUTENTICACION E IDENTIDAD
    # ───────────────────────────────────────────
    (
        "TC-AUTH-001", "Autenticacion", "Login exitoso con credenciales validas",
        "Alta", "Funcional",
        "Usuario registrado activo existe en el sistema.",
        "1. Ir a la pantalla de Login.\n2. Ingresar email y contrasena validos.\n3. Hacer clic en 'Iniciar sesion'.",
        "El sistema redirige al Dashboard. La barra lateral muestra el nombre del usuario y su rol."
    ),
    (
        "TC-AUTH-002", "Autenticacion", "Login fallido con contrasena incorrecta",
        "Alta", "Negativo",
        "Usuario registrado activo existe en el sistema.",
        "1. Ir a la pantalla de Login.\n2. Ingresar email valido y contrasena incorrecta.\n3. Hacer clic en 'Iniciar sesion'.",
        "Se muestra el mensaje 'Credenciales invalidas'. No se concede acceso. El sistema no revela si el email existe."
    ),
    (
        "TC-AUTH-003", "Autenticacion", "Login con email no registrado",
        "Alta", "Negativo",
        "Ninguna.",
        "1. Ir a la pantalla de Login.\n2. Ingresar un email inexistente.\n3. Hacer clic en 'Iniciar sesion'.",
        "Se muestra mensaje de error generico. No se indica si el email existe o no (prevencion de enumeracion)."
    ),
    (
        "TC-AUTH-004", "Autenticacion", "Cierre de sesion",
        "Alta", "Funcional",
        "Usuario autenticado en el sistema.",
        "1. Hacer clic en el menu de usuario (avatar/nombre).\n2. Seleccionar 'Cerrar sesion'.",
        "La sesion se cierra. El usuario es redirigido a Login. Intentar acceder a URL protegida redirige a Login."
    ),
    (
        "TC-AUTH-005", "Autenticacion", "Acceso directo a URL protegida sin sesion",
        "Alta", "Seguridad",
        "Usuario no autenticado (sesion cerrada o expirada).",
        "1. Ingresar directamente una URL de pagina interna (ej: /proyectos) en el navegador.",
        "El sistema redirige automaticamente a la pantalla de Login. No se muestra contenido protegido."
    ),
    (
        "TC-AUTH-006", "Autenticacion", "Bloqueo por intentos fallidos (brute force)",
        "Media", "Seguridad",
        "Usuario registrado existe.",
        "1. Intentar Login con contrasena incorrecta 5 veces consecutivas.",
        "El sistema bloquea temporalmente la cuenta o muestra CAPTCHA. Se muestra mensaje de bloqueo con tiempo de espera."
    ),
    (
        "TC-AUTH-007", "Autenticacion", "Aislamiento de datos entre tenants",
        "Alta", "Seguridad",
        "Dos tenants distintos (empresa A y empresa B) con usuarios separados.",
        "1. Autenticarse como usuario de Empresa A.\n2. Intentar acceder por URL a recursos de Empresa B (ej: /api/projects?tenantId=B).",
        "El sistema deniega el acceso. No se muestra ningun dato de la Empresa B. Se retorna HTTP 403 o datos vacios."
    ),

    # ───────────────────────────────────────────
    # MOD-02  EMPRESAS
    # ───────────────────────────────────────────
    (
        "TC-EMP-001", "Empresas", "Crear empresa nueva con datos completos",
        "Alta", "Funcional",
        "Usuario autenticado con rol Administrador.",
        "1. Ir a menu Empresas > Nueva Empresa.\n2. Completar todos los campos: nombre, RUT, direccion, telefono, contacto.\n3. Hacer clic en 'Guardar'.",
        "La empresa se crea exitosamente. Aparece en el listado de empresas. Se muestra mensaje de confirmacion."
    ),
    (
        "TC-EMP-002", "Empresas", "Crear empresa sin campo requerido (nombre)",
        "Alta", "Negativo",
        "Usuario autenticado con rol Administrador.",
        "1. Ir a menu Empresas > Nueva Empresa.\n2. Dejar el campo Nombre vacio.\n3. Completar resto de campos.\n4. Hacer clic en 'Guardar'.",
        "El sistema muestra error de validacion indicando que el nombre es requerido. La empresa NO se crea."
    ),
    (
        "TC-EMP-003", "Empresas", "Editar empresa existente",
        "Alta", "Funcional",
        "Empresa creada previamente en el sistema.",
        "1. Ir a Empresas.\n2. Seleccionar una empresa.\n3. Hacer clic en 'Editar'.\n4. Modificar el nombre.\n5. Guardar cambios.",
        "Los cambios se persisten. El listado y el detalle de empresa muestran el nuevo nombre."
    ),
    (
        "TC-EMP-004", "Empresas", "Buscar empresa por nombre",
        "Media", "Funcional",
        "Al menos 5 empresas creadas en el sistema.",
        "1. Ir a menu Empresas.\n2. En el buscador, escribir parte del nombre de una empresa existente.",
        "El listado filtra y muestra unicamente las empresas cuyo nombre coincide con el criterio ingresado."
    ),
    (
        "TC-EMP-005", "Empresas", "Eliminar empresa (soft delete)",
        "Alta", "Funcional",
        "Empresa sin proyectos activos.",
        "1. Ir a Empresas.\n2. Seleccionar empresa.\n3. Hacer clic en 'Eliminar' y confirmar.",
        "La empresa desaparece del listado activo. Al consultar la API, is_deleted=true. Los datos no se borran fisicamente."
    ),
    (
        "TC-EMP-006", "Empresas", "Paginacion del listado de empresas",
        "Media", "Funcional",
        "Mas de 20 empresas creadas.",
        "1. Ir a menu Empresas.\n2. Desplazarse al final de la primera pagina.\n3. Cambiar de pagina.",
        "El listado pagina correctamente. Cada pagina muestra el numero correcto de registros. Los controles de paginacion funcionan."
    ),

    # ───────────────────────────────────────────
    # MOD-03  PROYECTOS
    # ───────────────────────────────────────────
    (
        "TC-PROY-001", "Proyectos", "Crear proyecto nuevo con todos los campos",
        "Alta", "Funcional",
        "Empresa creada. Usuario con rol Administrador.",
        "1. Ir a Proyectos > Nuevo Proyecto.\n2. Completar: nombre, empresa, ubicacion, fechas inicio/fin, descripcion.\n3. Guardar.",
        "El proyecto se crea y aparece en el listado. Muestra los datos ingresados correctamente."
    ),
    (
        "TC-PROY-002", "Proyectos", "Agregar etapas a un proyecto",
        "Alta", "Funcional",
        "Proyecto creado.",
        "1. Abrir detalle del proyecto.\n2. Ir a pestana Etapas.\n3. Crear tres etapas: Fundacion, Estructura, Terminaciones.",
        "Las tres etapas se crean y se listan dentro del proyecto. El orden se respeta."
    ),
    (
        "TC-PROY-003", "Proyectos", "Agregar sectores a una etapa",
        "Alta", "Funcional",
        "Proyecto con al menos una etapa creada.",
        "1. Dentro de una etapa, ir a Sectores.\n2. Crear dos sectores: Torre A, Torre B.",
        "Los sectores se crean correctamente y se asocian a la etapa."
    ),
    (
        "TC-PROY-004", "Proyectos", "Agregar unidades (viviendas) a un sector",
        "Alta", "Funcional",
        "Sector creado dentro de una etapa.",
        "1. Dentro del sector, ir a Unidades.\n2. Crear unidades: Depto 101, 102, 103.",
        "Las unidades se crean y se listan dentro del sector con su nombre e identificador."
    ),
    (
        "TC-PROY-005", "Proyectos", "Crear proyecto con fecha fin anterior a fecha inicio",
        "Alta", "Negativo",
        "Usuario con rol Administrador.",
        "1. Ir a Proyectos > Nuevo Proyecto.\n2. Ingresar fecha fin anterior a la fecha de inicio.\n3. Guardar.",
        "El sistema muestra error de validacion. El proyecto NO se crea."
    ),
    (
        "TC-PROY-006", "Proyectos", "Asignar contratista a un proyecto",
        "Media", "Funcional",
        "Proyecto creado. Contratista registrado en el sistema.",
        "1. Ir al detalle del proyecto.\n2. En pestana Contratistas, buscar y agregar un contratista.",
        "El contratista queda asociado al proyecto y aparece en la lista de contratistas del mismo."
    ),

    # ───────────────────────────────────────────
    # MOD-04  USUARIOS Y PERMISOS
    # ───────────────────────────────────────────
    (
        "TC-USR-001", "Usuarios y Permisos", "Crear usuario nuevo con rol Inspector",
        "Alta", "Funcional",
        "Usuario con rol Administrador autenticado.",
        "1. Ir a Configuracion > Usuarios > Nuevo Usuario.\n2. Ingresar nombre, email, contrasena temporal.\n3. Asignar rol Inspector.\n4. Guardar.",
        "El usuario se crea. Recibe email de bienvenida con instrucciones de primer acceso."
    ),
    (
        "TC-USR-002", "Usuarios y Permisos", "Crear usuario con email ya registrado",
        "Alta", "Negativo",
        "Usuario con email 'test@empresa.cl' ya registrado.",
        "1. Ir a Nuevo Usuario.\n2. Ingresar el mismo email ya registrado.\n3. Guardar.",
        "El sistema muestra error indicando que el email ya esta en uso. No se crea duplicado."
    ),
    (
        "TC-USR-003", "Usuarios y Permisos", "Cambiar rol de usuario",
        "Alta", "Funcional",
        "Usuario existente con rol Inspector.",
        "1. Ir a Usuarios.\n2. Seleccionar el usuario.\n3. Cambiar rol de Inspector a Supervisor.\n4. Guardar.",
        "El rol queda actualizado. Al hacer login el usuario solo ve funciones de Supervisor."
    ),
    (
        "TC-USR-004", "Usuarios y Permisos", "Usuario con rol Inspector no puede acceder a configuracion",
        "Alta", "Seguridad",
        "Usuario autenticado con rol Inspector.",
        "1. Autenticarse con usuario Inspector.\n2. Intentar navegar a Configuracion > Usuarios por URL directa.",
        "El sistema muestra pagina de acceso denegado (403) o redirige al Dashboard. El menu de Configuracion no es visible."
    ),
    (
        "TC-USR-005", "Usuarios y Permisos", "Desactivar usuario",
        "Alta", "Funcional",
        "Usuario activo en el sistema.",
        "1. Ir a Usuarios.\n2. Seleccionar un usuario activo.\n3. Hacer clic en 'Desactivar' y confirmar.",
        "El usuario queda inactivo. Al intentar login con ese usuario, el sistema rechaza el acceso con mensaje 'Cuenta inactiva'."
    ),
    (
        "TC-USR-006", "Usuarios y Permisos", "Asignar usuario a proyecto especifico",
        "Alta", "Funcional",
        "Usuario y proyecto existentes.",
        "1. Ir al detalle del Proyecto.\n2. En pestana Equipo, buscar y agregar un usuario.\n3. Asignar rol en el proyecto.",
        "El usuario queda asignado al proyecto con el rol especificado. Solo puede ver este proyecto en su vista."
    ),

    # ───────────────────────────────────────────
    # MOD-05  PLANTILLAS DE INSPECCION
    # ───────────────────────────────────────────
    (
        "TC-TPL-001", "Plantillas", "Crear plantilla con secciones y preguntas",
        "Alta", "Funcional",
        "Usuario con rol Administrador.",
        "1. Ir a Plantillas > Nueva Plantilla.\n2. Ingresar nombre y descripcion.\n3. Agregar 2 secciones.\n4. En cada seccion agregar 3 preguntas de distintos tipos (Si/No, Texto, Escala).\n5. Guardar.",
        "La plantilla se crea con todas las secciones y preguntas. Se muestra correctamente en el editor."
    ),
    (
        "TC-TPL-002", "Plantillas", "Crear plantilla sin secciones",
        "Alta", "Negativo",
        "Usuario con rol Administrador.",
        "1. Ir a Plantillas > Nueva Plantilla.\n2. Ingresar solo nombre.\n3. No agregar secciones.\n4. Intentar guardar.",
        "El sistema muestra error indicando que la plantilla debe tener al menos una seccion. No se crea."
    ),
    (
        "TC-TPL-003", "Plantillas", "Editar plantilla: agregar nueva seccion",
        "Alta", "Funcional",
        "Plantilla existente con 2 secciones.",
        "1. Abrir plantilla existente en modo edicion.\n2. Agregar una tercera seccion con 2 preguntas.\n3. Guardar.",
        "La plantilla se actualiza con la nueva seccion. Las secciones previas no se alteran."
    ),
    (
        "TC-TPL-004", "Plantillas", "Versionar una plantilla",
        "Alta", "Funcional",
        "Plantilla usada en inspecciones previas.",
        "1. Abrir plantilla activa.\n2. Hacer clic en 'Nueva Version'.\n3. Modificar una pregunta.\n4. Publicar nueva version.",
        "Se crea una nueva version de la plantilla. La version anterior queda como historico. Inspecciones previas conservan la version original."
    ),
    (
        "TC-TPL-005", "Plantillas", "Tipos de pregunta disponibles",
        "Media", "Funcional",
        "Editor de plantilla abierto.",
        "1. Agregar una pregunta de cada tipo: Si/No, Texto libre, Escala numerica, Seleccion multiple, Lista desplegable.\n2. Guardar plantilla.",
        "Todos los tipos de pregunta se guardan correctamente y se renderizan en la ejecucion de inspeccion."
    ),
    (
        "TC-TPL-006", "Plantillas", "Eliminar plantilla en uso",
        "Alta", "Negativo",
        "Plantilla asociada a al menos una inspeccion existente.",
        "1. Ir a Plantillas.\n2. Intentar eliminar plantilla que ya tiene inspecciones.",
        "El sistema impide la eliminacion y muestra mensaje informando que la plantilla esta en uso."
    ),
    (
        "TC-TPL-007", "Plantillas", "Reordenar secciones en el editor",
        "Media", "Funcional",
        "Plantilla con 3 secciones.",
        "1. Abrir plantilla en editor.\n2. Arrastrar la seccion 3 al primer lugar.\n3. Guardar.",
        "El nuevo orden se persiste. Al visualizar la plantilla en inspeccion, las secciones aparecen en el orden correcto."
    ),

    # ───────────────────────────────────────────
    # MOD-06  PROGRAMACION DE INSPECCIONES
    # ───────────────────────────────────────────
    (
        "TC-PROG-001", "Programacion", "Crear inspeccion programada",
        "Alta", "Funcional",
        "Proyecto, plantilla e inspector existentes.",
        "1. Ir a Inspecciones > Nueva Inspeccion.\n2. Seleccionar proyecto, etapa/sector/unidad.\n3. Seleccionar plantilla.\n4. Asignar inspector.\n5. Definir fecha y hora programada.\n6. Guardar.",
        "La inspeccion queda en estado 'Programada'. El inspector asignado puede verla en su lista de tareas."
    ),
    (
        "TC-PROG-002", "Programacion", "Crear inspeccion sin asignar inspector",
        "Alta", "Negativo",
        "Proyecto y plantilla existentes.",
        "1. Ir a Nueva Inspeccion.\n2. Completar todos los campos excepto inspector.\n3. Intentar guardar.",
        "El sistema muestra error de validacion. La inspeccion no se crea hasta asignar un inspector."
    ),
    (
        "TC-PROG-003", "Programacion", "Crear inspeccion con fecha pasada",
        "Media", "Borde",
        "Proyecto y plantilla existentes.",
        "1. Ir a Nueva Inspeccion.\n2. Ingresar fecha programada de ayer.\n3. Guardar.",
        "El sistema muestra advertencia de fecha pasada. Debe confirmarse explicitamente o bloquearse segun regla de negocio definida."
    ),
    (
        "TC-PROG-004", "Programacion", "Ver calendario de inspecciones",
        "Media", "Funcional",
        "Al menos 5 inspecciones programadas en distintas fechas.",
        "1. Ir a Inspecciones > Vista Calendario.\n2. Navegar entre semanas/meses.",
        "El calendario muestra las inspecciones en las fechas correctas. Hacer clic en una abre su detalle."
    ),
    (
        "TC-PROG-005", "Programacion", "Reasignar inspector de inspeccion programada",
        "Alta", "Funcional",
        "Inspeccion en estado 'Programada' con inspector asignado.",
        "1. Abrir detalle de inspeccion.\n2. Cambiar inspector asignado por otro disponible.\n3. Guardar.",
        "El nuevo inspector queda asignado. El inspector anterior ya no ve la inspeccion en su lista."
    ),
    (
        "TC-PROG-006", "Programacion", "Cancelar inspeccion programada",
        "Media", "Funcional",
        "Inspeccion en estado 'Programada'.",
        "1. Abrir detalle de inspeccion.\n2. Hacer clic en 'Cancelar Inspeccion'.\n3. Ingresar motivo.\n4. Confirmar.",
        "La inspeccion queda en estado 'Cancelada'. Se registra el motivo y quien la cancelo con timestamp."
    ),

    # ───────────────────────────────────────────
    # MOD-07  OBSERVACIONES / NO CONFORMIDADES
    # ───────────────────────────────────────────
    (
        "TC-OBS-001", "Observaciones", "Registrar observacion manual desde web",
        "Alta", "Funcional",
        "Inspeccion en estado 'Ejecutada' o 'En Revision'.",
        "1. Ir al detalle de una inspeccion.\n2. Hacer clic en 'Agregar Observacion'.\n3. Completar: descripcion, severidad (Leve/Grave/Critico), contratista responsable.\n4. Guardar.",
        "La observacion se registra con estado 'Abierta'. Aparece en el listado de NC del proyecto."
    ),
    (
        "TC-OBS-002", "Observaciones", "Asignar observacion a contratista",
        "Alta", "Funcional",
        "Observacion en estado 'Abierta'. Contratista asignado al proyecto.",
        "1. Ir al detalle de la observacion.\n2. Cambiar estado a 'En Correccion'.\n3. Asignar contratista responsable.\n4. Definir fecha maxima de correccion.\n5. Guardar.",
        "La observacion queda asignada con estado 'En Correccion'. El contratista recibe notificacion por email."
    ),
    (
        "TC-OBS-003", "Observaciones", "Cerrar observacion con evidencia",
        "Alta", "Funcional",
        "Observacion en estado 'En Correccion'.",
        "1. Abrir la observacion.\n2. Cambiar estado a 'Resuelta'.\n3. Agregar descripcion del cierre.\n4. Adjuntar imagen de evidencia.\n5. Guardar.",
        "La observacion queda en estado 'Resuelta'. Se registra quien la cerro, fecha y evidencia adjunta."
    ),
    (
        "TC-OBS-004", "Observaciones", "Filtrar observaciones por severidad",
        "Media", "Funcional",
        "Al menos 10 observaciones de distintas severidades en el proyecto.",
        "1. Ir a Observaciones del proyecto.\n2. Aplicar filtro por severidad 'Critico'.",
        "El listado muestra unicamente observaciones criticas. El contador refleja el numero filtrado."
    ),
    (
        "TC-OBS-005", "Observaciones", "Filtrar observaciones por estado",
        "Media", "Funcional",
        "Observaciones en multiples estados.",
        "1. Ir a Observaciones.\n2. Filtrar por estado 'Abierta'.",
        "Se muestran solo las observaciones abiertas."
    ),
    (
        "TC-OBS-006", "Observaciones", "Ver historial de cambios de una observacion",
        "Alta", "Funcional",
        "Observacion con varios cambios de estado registrados.",
        "1. Abrir detalle de observacion.\n2. Ir a pestana 'Historial'.",
        "Se muestra la linea de tiempo de cambios: quien cambio el estado, a que estado, con que comentario y en que fecha."
    ),
    (
        "TC-OBS-007", "Observaciones", "Observacion vencida (fecha limite superada)",
        "Media", "Borde",
        "Observacion en estado 'En Correccion' con fecha limite de ayer.",
        "1. Ir al listado de observaciones del proyecto.",
        "La observacion aparece destacada visualmente (ej: color rojo o icono de alerta) indicando que esta vencida."
    ),

    # ───────────────────────────────────────────
    # MOD-08  REINSPECCIONES
    # ───────────────────────────────────────────
    (
        "TC-REINSP-001", "Reinspecciones", "Crear reinspeccion desde observacion resuelta",
        "Alta", "Funcional",
        "Observacion en estado 'Resuelta' pendiente de verificacion.",
        "1. Abrir la observacion resuelta.\n2. Hacer clic en 'Solicitar Reinspeccion'.\n3. Asignar inspector y fecha.\n4. Confirmar.",
        "Se crea una reinspeccion asociada a la observacion original. El inspector asignado la ve en su lista."
    ),
    (
        "TC-REINSP-002", "Reinspecciones", "Aprobar correccion en reinspeccion",
        "Alta", "Funcional",
        "Reinspeccion ejecutada.",
        "1. Abrir la reinspeccion.\n2. Cambiar resultado a 'Correccion Aprobada'.\n3. Agregar comentario.\n4. Guardar.",
        "La observacion original cambia a estado 'Cerrada'. La reinspeccion queda como 'Aprobada'. Se registra trazabilidad completa."
    ),
    (
        "TC-REINSP-003", "Reinspecciones", "Rechazar correccion en reinspeccion",
        "Alta", "Funcional",
        "Reinspeccion ejecutada.",
        "1. Abrir la reinspeccion.\n2. Cambiar resultado a 'Correccion Rechazada'.\n3. Ingresar descripcion del rechazo.\n4. Guardar.",
        "La observacion original vuelve al estado 'En Correccion'. Se registra el rechazo con motivo y fecha."
    ),
    (
        "TC-REINSP-004", "Reinspecciones", "Trazabilidad completa: inspeccion > observacion > reinspeccion",
        "Alta", "Funcional",
        "Flujo completo ejecutado: inspeccion con observacion cerrada por reinspeccion.",
        "1. Desde el detalle de la inspeccion original, navegar a sus observaciones.\n2. Abrir una observacion cerrada.\n3. Ir a su historial y reinspecciones asociadas.",
        "Se puede navegar el flujo completo: inspeccion > respuesta > observacion > reinspeccion > cierre. Todo con timestamps y responsables."
    ),

    # ───────────────────────────────────────────
    # MOD-09  DOCUMENTOS DE PROYECTO
    # ───────────────────────────────────────────
    (
        "TC-DOC-001", "Documentos", "Subir documento PDF a proyecto",
        "Alta", "Funcional",
        "Proyecto creado. Archivo PDF disponible (< 20 MB).",
        "1. Ir al detalle del proyecto.\n2. Ir a pestana 'Documentos'.\n3. Hacer clic en 'Subir Documento'.\n4. Seleccionar archivo PDF.\n5. Ingresar nombre y tipo (ej: Plano Arquitectonico).\n6. Confirmar.",
        "El documento se sube y aparece en la lista. El nombre, tipo y fecha de carga se muestran correctamente."
    ),
    (
        "TC-DOC-002", "Documentos", "Descargar documento subido",
        "Alta", "Funcional",
        "Documento subido al proyecto.",
        "1. Ir a Documentos del proyecto.\n2. Hacer clic en el documento.\n3. Hacer clic en 'Descargar'.",
        "El archivo se descarga correctamente y puede abrirse. El archivo descargado es identico al subido."
    ),
    (
        "TC-DOC-003", "Documentos", "Subir archivo con formato no permitido",
        "Alta", "Negativo",
        "Proyecto existente.",
        "1. Ir a Documentos del proyecto.\n2. Intentar subir un archivo .exe o .bat.",
        "El sistema rechaza el archivo y muestra mensaje indicando los formatos permitidos."
    ),
    (
        "TC-DOC-004", "Documentos", "Subir archivo que supera el limite de tamano",
        "Alta", "Negativo",
        "Proyecto existente. Archivo > 50 MB disponible.",
        "1. Ir a Documentos del proyecto.\n2. Intentar subir un archivo que supera el limite definido.",
        "El sistema muestra error indicando el limite de tamano antes o durante la subida. El archivo no se almacena."
    ),
    (
        "TC-DOC-005", "Documentos", "Eliminar documento de proyecto",
        "Media", "Funcional",
        "Documento subido al proyecto.",
        "1. Ir a Documentos del proyecto.\n2. Seleccionar un documento.\n3. Hacer clic en 'Eliminar' y confirmar.",
        "El documento desaparece del listado. El archivo se elimina del almacenamiento MinIO (soft delete o hard delete segun configuracion)."
    ),

    # ───────────────────────────────────────────
    # MOD-10  REPORTES
    # ───────────────────────────────────────────
    (
        "TC-RPT-001", "Reportes", "Generar reporte PDF de inspeccion",
        "Alta", "Funcional",
        "Inspeccion completada con respuestas, evidencias y observaciones.",
        "1. Ir al detalle de la inspeccion.\n2. Hacer clic en 'Generar Reporte PDF'.\n3. Esperar generacion.",
        "Se descarga un PDF con: encabezado de empresa/proyecto, datos del inspector, fecha, secciones y respuestas, evidencias fotograficas, listado de observaciones y firma."
    ),
    (
        "TC-RPT-002", "Reportes", "Exportar listado de observaciones a Excel",
        "Alta", "Funcional",
        "Proyecto con al menos 10 observaciones.",
        "1. Ir a Observaciones del proyecto.\n2. Aplicar filtros deseados.\n3. Hacer clic en 'Exportar Excel'.",
        "Se descarga un archivo .xlsx con todas las columnas relevantes: ID, descripcion, severidad, estado, responsable, fechas. Los filtros aplicados se reflejan en la exportacion."
    ),
    (
        "TC-RPT-003", "Reportes", "Reporte PDF muestra evidencias fotograficas",
        "Alta", "Funcional",
        "Inspeccion con al menos 3 fotos de evidencia adjuntas.",
        "1. Generar reporte PDF de la inspeccion.",
        "El PDF incluye las imagenes de evidencia con su descripcion y la pregunta/seccion a la que corresponden."
    ),
    (
        "TC-RPT-004", "Reportes", "Reporte PDF de inspeccion sin observaciones",
        "Media", "Borde",
        "Inspeccion completada sin observaciones registradas.",
        "1. Generar reporte PDF de la inspeccion.",
        "El PDF se genera correctamente. La seccion de observaciones indica 'Sin observaciones registradas' o similar."
    ),
    (
        "TC-RPT-005", "Reportes", "Exportar Excel con cero registros (filtro sin resultados)",
        "Media", "Borde",
        "Proyecto con observaciones.",
        "1. Ir a Observaciones.\n2. Aplicar filtro que no retorne resultados.\n3. Hacer clic en 'Exportar Excel'.",
        "El sistema genera un Excel con los encabezados de columna pero sin filas de datos, o muestra mensaje indicando que no hay datos para exportar."
    ),

    # ───────────────────────────────────────────
    # MOD-11  DASHBOARD
    # ───────────────────────────────────────────
    (
        "TC-DASH-001", "Dashboard", "KPIs se muestran correctamente",
        "Alta", "Funcional",
        "Al menos 1 proyecto activo con inspecciones y observaciones.",
        "1. Autenticarse y navegar al Dashboard.\n2. Revisar los indicadores: inspecciones programadas, ejecutadas, observaciones abiertas, % cumplimiento.",
        "Los KPIs muestran valores consistentes con los datos del sistema. Los numeros coinciden con los listados de inspecciones y observaciones."
    ),
    (
        "TC-DASH-002", "Dashboard", "Graficos se renderizan correctamente",
        "Alta", "Funcional",
        "Datos historicos de inspecciones y observaciones cargados.",
        "1. Ir al Dashboard.\n2. Revisar graficos de barras, torta y linea de tendencia.",
        "Todos los graficos se renderizan sin errores. Los datos son consistentes con los KPIs numericos."
    ),
    (
        "TC-DASH-003", "Dashboard", "Filtrar dashboard por proyecto",
        "Alta", "Funcional",
        "Multiple proyectos con datos.",
        "1. Ir al Dashboard.\n2. Cambiar el selector de proyecto.",
        "Los KPIs y graficos se actualizan para mostrar solo datos del proyecto seleccionado."
    ),
    (
        "TC-DASH-004", "Dashboard", "Dashboard con cero datos (proyecto nuevo)",
        "Media", "Borde",
        "Proyecto recien creado sin inspecciones.",
        "1. Ir al Dashboard.\n2. Seleccionar el proyecto nuevo.",
        "El Dashboard se muestra con KPIs en cero y graficos vacios. No hay errores ni pantallas en blanco."
    ),
    (
        "TC-DASH-005", "Dashboard", "Dashboard es visible segun rol",
        "Media", "Seguridad",
        "Usuarios con roles distintos: Admin, Supervisor, Inspector.",
        "1. Autenticarse como Inspector.\n2. Revisar que el Dashboard solo muestre datos de inspecciones propias.",
        "Cada rol ve solo la informacion pertinente a sus permisos. Un Inspector no ve KPIs globales de la empresa."
    ),

    # ───────────────────────────────────────────
    # MOD-12  NOTIFICACIONES
    # ───────────────────────────────────────────
    (
        "TC-NOTIF-001", "Notificaciones", "Email al asignar inspeccion a inspector",
        "Alta", "Funcional",
        "Inspector con email valido configurado.",
        "1. Crear y asignar una nueva inspeccion a un inspector.\n2. Guardar.",
        "El inspector recibe email con asunto, detalle de la inspeccion (proyecto, fecha, plantilla) y enlace directo."
    ),
    (
        "TC-NOTIF-002", "Notificaciones", "Email al registrar observacion con responsable",
        "Alta", "Funcional",
        "Contratista con email valido asignado al proyecto.",
        "1. Registrar observacion y asignar a un contratista responsable.\n2. Cambiar estado a 'En Correccion'.",
        "El contratista responsable recibe email con detalle de la observacion, severidad y fecha limite."
    ),
    (
        "TC-NOTIF-003", "Notificaciones", "No enviar email si el usuario no tiene email configurado",
        "Media", "Borde",
        "Usuario sin email configurado.",
        "1. Asignar una inspeccion a un usuario sin email.",
        "El sistema no genera error. La inspeccion se crea correctamente. El sistema registra en log que no se pudo enviar la notificacion."
    ),

    # ───────────────────────────────────────────
    # MOD-13  AUDITORIA
    # ───────────────────────────────────────────
    (
        "TC-AUD-001", "Auditoria", "Registro de creacion de entidad",
        "Alta", "Funcional",
        "Usuario autenticado.",
        "1. Crear un nuevo proyecto.\n2. Ir al log de auditoria.",
        "Se registra la accion: entidad creada, usuario que la creo, timestamp y tenant. El campo created_by coincide con el usuario autenticado."
    ),
    (
        "TC-AUD-002", "Auditoria", "Registro de modificacion de entidad",
        "Alta", "Funcional",
        "Proyecto existente.",
        "1. Editar el nombre de un proyecto.\n2. Ir al log de auditoria del proyecto.",
        "Se registra la modificacion: campo modificado, valor anterior, valor nuevo, usuario que modifico, timestamp."
    ),
    (
        "TC-AUD-003", "Auditoria", "Registro de eliminacion (soft delete)",
        "Alta", "Funcional",
        "Empresa sin proyectos activos.",
        "1. Eliminar una empresa.\n2. Consultar el log de auditoria.",
        "Se registra la eliminacion con: usuario que la ejecuto, timestamp, is_deleted=true. La empresa no aparece en listados activos pero existe en BD."
    ),
    (
        "TC-AUD-004", "Auditoria", "Usuario no admin no puede acceder al log completo de auditoria",
        "Alta", "Seguridad",
        "Usuario con rol Inspector.",
        "1. Autenticarse como Inspector.\n2. Intentar acceder al modulo de Auditoria por URL directa.",
        "El sistema deniega el acceso. El menu de Auditoria no es visible para roles sin permiso."
    ),

    # ───────────────────────────────────────────
    # MOD-14  RENDIMIENTO Y UI GENERAL
    # ───────────────────────────────────────────
    (
        "TC-UI-001", "UI General", "Tiempo de carga inicial del Dashboard",
        "Media", "Funcional",
        "Sistema con datos de al menos 1 proyecto activo.",
        "1. Autenticarse.\n2. Medir el tiempo hasta que el Dashboard muestra datos completos.",
        "El Dashboard carga en menos de 3 segundos en condiciones normales de red."
    ),
    (
        "TC-UI-002", "UI General", "Responsive en tablet (1024px)",
        "Media", "UI",
        "Acceso desde navegador en resolucion 1024x768 o tablet.",
        "1. Acceder al sistema desde una tablet o emular 1024px.\n2. Navegar por Dashboard, Proyectos y Observaciones.",
        "El layout se adapta correctamente. Los elementos no se superponen ni desbordan el viewport."
    ),
    (
        "TC-UI-003", "UI General", "Mensajes de validacion son claros y en espanol",
        "Media", "UI",
        "Formulario con campos requeridos.",
        "1. Intentar guardar cualquier formulario con campos obligatorios vacios.",
        "Los mensajes de error son en espanol, descriptivos e indican exactamente que campo necesita correccion."
    ),
    (
        "TC-UI-004", "UI General", "Confirmacion antes de eliminar registros",
        "Alta", "UI",
        "Cualquier entidad eliminable en el sistema.",
        "1. Hacer clic en el boton 'Eliminar' de cualquier entidad (empresa, proyecto, usuario, etc.).",
        "Aparece un dialogo de confirmacion que describe la accion y sus consecuencias. Solo al confirmar se ejecuta la eliminacion."
    ),
    (
        "TC-UI-005", "UI General", "Indicador de carga durante operaciones largas",
        "Media", "UI",
        "Operacion que tome mas de 1 segundo (generacion de reporte, carga de lista grande).",
        "1. Iniciar la operacion larga (ej: generar PDF, cargar listado de 500 registros).",
        "Se muestra un indicador de progreso o spinner mientras la operacion se ejecuta. No queda la interfaz congelada sin feedback."
    ),
]

# ─────────────────────────────────────────────
# GENERADOR DEL DOCUMENTO
# ─────────────────────────────────────────────

PRIORITY_COLORS = {
    "Alta":   ("C62828", "FFFFFF"),
    "Media":  ("F57C00", "FFFFFF"),
    "Baja":   ("388E3C", "FFFFFF"),
}

TYPE_COLORS = {
    "Funcional":  "E3F2FD",
    "Negativo":   "FCE4EC",
    "Borde":      "FFF8E1",
    "Seguridad":  "F3E5F5",
    "UI":         "E8F5E9",
}

MODULE_COLORS = {
    "Autenticacion":     "1B5E91",
    "Empresas":          "1565C0",
    "Proyectos":         "0D47A1",
    "Usuarios y Permisos": "283593",
    "Plantillas":        "1A237E",
    "Programacion":      "006064",
    "Observaciones":     "BF360C",
    "Reinspecciones":    "4E342E",
    "Documentos":        "33691E",
    "Reportes":          "1B5E20",
    "Dashboard":         "827717",
    "Notificaciones":    "E65100",
    "Auditoria":         "37474F",
    "UI General":        "4A148C",
}


def build_qa_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ─── PORTADA ───────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ITO CLOUD")
    run.bold = True
    run.font.size = Pt(34)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Set de Pruebas QA - Modulos Web")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version 1.0  |  Fecha: {datetime.date.today().strftime('%d/%m/%Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento confidencial de uso interno QA")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    doc.add_page_break()

    # ─── INDICE RESUMEN ────────────────────────
    heading1(doc, "1. Resumen del Set de Pruebas")
    add_hr(doc)
    body(doc,
        "Este documento contiene los casos de prueba para la validacion funcional, de seguridad "
        "y de interfaz de usuario de la plataforma web ITO Cloud. Cada caso de prueba esta "
        "estructurado con precondiciones, pasos detallados y resultado esperado para facilitar "
        "la ejecucion por el equipo QA."
    )

    # Tabla resumen por modulo
    modules_summary = {}
    for tc in TEST_CASES:
        mod = tc[1]
        modules_summary[mod] = modules_summary.get(mod, 0) + 1

    total = len(TEST_CASES)
    alta  = sum(1 for tc in TEST_CASES if tc[3] == "Alta")
    media = sum(1 for tc in TEST_CASES if tc[3] == "Media")
    baja  = sum(1 for tc in TEST_CASES if tc[3] == "Baja")

    doc.add_paragraph()
    tbl_resumen = doc.add_table(rows=1, cols=3)
    tbl_resumen.style = 'Table Grid'
    tbl_resumen.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate([4, 4, 4]):
        tbl_resumen.columns[i].width = Cm(w)
    table_header_row_qa(tbl_resumen, ['Total Casos', 'Prioridad Alta', 'Prioridad Media'])
    row = tbl_resumen.add_row()
    for i, val in enumerate([str(total), str(alta), str(media)]):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)
        set_cell_bg(cell, 'EBF3FB')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()

    # Tabla por modulo
    heading2(doc, "Casos por Modulo")
    tbl_mod = doc.add_table(rows=1, cols=3)
    tbl_mod.style = 'Table Grid'
    tbl_mod.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_mod.columns[0].width = Cm(6.5)
    tbl_mod.columns[1].width = Cm(2.5)
    tbl_mod.columns[2].width = Cm(7.0)
    table_header_row_qa(tbl_mod, ['Modulo', 'N. Casos', 'Prefijo ID'])

    prefijos = {
        "Autenticacion":       "TC-AUTH",
        "Empresas":            "TC-EMP",
        "Proyectos":           "TC-PROY",
        "Usuarios y Permisos": "TC-USR",
        "Plantillas":          "TC-TPL",
        "Programacion":        "TC-PROG",
        "Observaciones":       "TC-OBS",
        "Reinspecciones":      "TC-REINSP",
        "Documentos":          "TC-DOC",
        "Reportes":            "TC-RPT",
        "Dashboard":           "TC-DASH",
        "Notificaciones":      "TC-NOTIF",
        "Auditoria":           "TC-AUD",
        "UI General":          "TC-UI",
    }

    for i, (mod, count) in enumerate(modules_summary.items()):
        row = tbl_mod.add_row()
        row.cells[0].text = mod
        row.cells[1].text = str(count)
        row.cells[2].text = prefijos.get(mod, "TC-???")
        bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
        for c in row.cells:
            set_cell_bg(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ─── LEYENDA ───────────────────────────────
    heading2(doc, "Leyenda de Estados para QA")
    tbl_leyenda = doc.add_table(rows=1, cols=3)
    tbl_leyenda.style = 'Table Grid'
    tbl_leyenda.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_leyenda.columns[0].width = Cm(3.5)
    tbl_leyenda.columns[1].width = Cm(3.5)
    tbl_leyenda.columns[2].width = Cm(9.0)
    table_header_row_qa(tbl_leyenda, ['Estado', 'Indicar con', 'Descripcion'])
    leyenda = [
        ("Pendiente",  "PEND",   "El caso aun no ha sido ejecutado"),
        ("Aprobado",   "PASS",   "El resultado observado coincide con el resultado esperado"),
        ("Fallido",    "FAIL",   "El resultado observado NO coincide con el resultado esperado"),
        ("Bloqueado",  "BLOCK",  "No se puede ejecutar por un defecto previo o falta de prerequisito"),
        ("Omitido",    "SKIP",   "Se omite intencionalmente con justificacion documentada"),
    ]
    estado_bg = {
        "Pendiente": "F5F5F5",
        "Aprobado":  "E8F5E9",
        "Fallido":   "FFEBEE",
        "Bloqueado": "FFF8E1",
        "Omitido":   "F3E5F5",
    }
    for estado, sigla, desc in leyenda:
        row = tbl_leyenda.add_row()
        row.cells[0].text = estado
        row.cells[1].text = sigla
        row.cells[2].text = desc
        for c in row.cells:
            set_cell_bg(c, estado_bg.get(estado, 'FFFFFF'))
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        for r in row.cells[0].paragraphs[0].runs:
            r.bold = True

    doc.add_paragraph()

    # ─── CASOS DE PRUEBA ───────────────────────
    doc.add_page_break()
    heading1(doc, "2. Casos de Prueba")
    add_hr(doc)

    current_module = None
    for tc in TEST_CASES:
        (tc_id, modulo, nombre, prioridad, tipo,
         precondicion, pasos, resultado_esperado) = tc

        # Cabecera de modulo
        if modulo != current_module:
            current_module = modulo
            doc.add_paragraph()
            hex_c = MODULE_COLORS.get(modulo, '1B5E91')
            r_val = int(hex_c[0:2], 16)
            g_val = int(hex_c[2:4], 16)
            b_val = int(hex_c[4:6], 16)

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(f"  MODULO: {modulo.upper()}")
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # fondo del parrafo
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex_c)
            pPr.append(shd)

        # Tabla de caso de prueba
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.columns[0].width = Cm(3.5)
        tbl.columns[1].width = Cm(12.5)

        # Fila: ID + Nombre (cabecera del caso)
        row = tbl.add_row()
        row.cells[0].merge(row.cells[1])
        cell = row.cells[0]
        cell.text = ''
        p = cell.paragraphs[0]
        run_id = p.add_run(f"{tc_id}  ")
        run_id.bold = True
        run_id.font.size = Pt(10)
        run_id.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run_name = p.add_run(nombre)
        run_name.bold = True
        run_name.font.size = Pt(10)
        run_name.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, '1B5E91')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)

        # Fila: Prioridad + Tipo
        row2 = tbl.add_row()
        c_lbl = row2.cells[0]
        c_val = row2.cells[1]
        c_lbl.text = ''
        p_lbl = c_lbl.paragraphs[0]
        run_l = p_lbl.add_run("Prioridad / Tipo")
        run_l.bold = True
        run_l.font.size = Pt(9)
        run_l.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        set_cell_bg(c_lbl, 'ECEFF1')

        c_val.text = ''
        p_val = c_val.paragraphs[0]
        # Prioridad con color
        prio_bg, prio_fg = PRIORITY_COLORS.get(prioridad, ("888888", "FFFFFF"))
        run_prio = p_val.add_run(f" {prioridad} ")
        run_prio.bold = True
        run_prio.font.size = Pt(9)
        run_prio.font.color.rgb = RGBColor(
            int(prio_fg[0:2], 16), int(prio_fg[2:4], 16), int(prio_fg[4:6], 16)
        )
        # Tipo
        run_sep = p_val.add_run("   |   Tipo: ")
        run_sep.font.size = Pt(9)
        run_tipo = p_val.add_run(tipo)
        run_tipo.bold = True
        run_tipo.font.size = Pt(9)
        run_tipo.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
        set_cell_bg(c_val, TYPE_COLORS.get(tipo, 'FFFFFF'))

        def add_row_tc(label, content):
            r = tbl.add_row()
            r.cells[0].text = ''
            p0 = r.cells[0].paragraphs[0]
            run0 = p0.add_run(label)
            run0.bold = True
            run0.font.size = Pt(9)
            run0.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            set_cell_bg(r.cells[0], 'ECEFF1')
            r.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

            r.cells[1].text = ''
            p1 = r.cells[1].paragraphs[0]
            run1 = p1.add_run(content)
            run1.font.size = Pt(9)
            set_cell_bg(r.cells[1], 'FFFFFF')
            r.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p1.paragraph_format.space_before = Pt(2)
            p1.paragraph_format.space_after  = Pt(2)

        add_row_tc("Precondicion", precondicion)
        add_row_tc("Pasos", pasos)
        add_row_tc("Resultado Esperado", resultado_esperado)

        # Fila de resultado QA (para rellenar)
        row_res = tbl.add_row()
        row_res.cells[0].text = ''
        p_res_l = row_res.cells[0].paragraphs[0]
        run_res_l = p_res_l.add_run("Resultado QA")
        run_res_l.bold = True
        run_res_l.font.size = Pt(9)
        run_res_l.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        set_cell_bg(row_res.cells[0], 'ECEFF1')
        row_res.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        row_res.cells[1].text = ''
        p_res_v = row_res.cells[1].paragraphs[0]
        run_status = p_res_v.add_run("Estado: _______________   ")
        run_status.bold = True
        run_status.font.size = Pt(9)
        run_status.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)
        run_ejecutado = p_res_v.add_run("Ejecutado por: ___________________________   ")
        run_ejecutado.font.size = Pt(9)
        run_fecha = p_res_v.add_run("Fecha: ____________")
        run_fecha.font.size = Pt(9)

        p_obs_v = row_res.cells[1].add_paragraph()
        run_obs_l = p_obs_v.add_run("Observaciones / Defectos encontrados: ")
        run_obs_l.bold = True
        run_obs_l.font.size = Pt(9)
        run_obs_v = p_obs_v.add_run("_" * 60)
        run_obs_v.font.size = Pt(9)
        run_obs_v.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        set_cell_bg(row_res.cells[1], 'FFFDE7')
        row_res.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ─── PIE DE PAGINA ─────────────────────────
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"ITO Cloud - Set de Pruebas QA Modulos Web v1.0  |  "
        f"{datetime.date.today().strftime('%d/%m/%Y')}  |  Confidencial"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    return doc


def table_header_row_qa(table, headers, bg='1B5E91'):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == '__main__':
    output_path = r"F:\SISTEMA DE INSPECCION DE OBRA\docs\ITO_Cloud_QA_Tests.docx"
    doc = build_qa_document()
    doc.save(output_path)
    print(f"[OK] Documento generado: {output_path}")
    print(f"[OK] Total casos de prueba: {len(TEST_CASES)}")
