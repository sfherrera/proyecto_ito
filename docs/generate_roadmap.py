# -*- coding: utf-8 -*-
"""
Generador del Roadmap ITO Cloud en formato Word (.docx)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# UTILIDADES
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Establece color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom', 'start', 'end', 'left', 'right'):
        tag = 'w:{}'.format(edge)
        element = OxmlElement(tag)
        for key, value in kwargs.items():
            element.set(qn('w:{}'.format(key)), value)
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_hr(doc):
    """Línea horizontal separadora."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1B5E91')
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    return p

def body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if italic:
        run.italic = True
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def table_header_row(table, headers, bg='1B5E91'):
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def table_row(table, values, alt=False):
    row = table.add_row()
    bg = 'EBF3FB' if alt else 'FFFFFF'
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        set_cell_bg(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return row

# ─────────────────────────────────────────────
# DOCUMENTO
# ─────────────────────────────────────────────

def build_document():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Fuente por defecto
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ────────────────────────────
    # PORTADA
    # ────────────────────────────
    # Espacio superior
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ITO CLOUD")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Plataforma de Inspección Técnica de Obras")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ROADMAP DE SOLUCIÓN")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Versión 1.0  |  Fecha: {datetime.date.today().strftime('%d/%m/%Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SaaS B2B Multiempresa — Construcción e Inmobiliaria")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Página nueva
    doc.add_page_break()

    # ────────────────────────────
    # 1. RESUMEN EJECUTIVO
    # ────────────────────────────
    heading1(doc, "1. Resumen Ejecutivo")
    add_hr(doc)
    body(doc,
        "ITO Cloud es una plataforma SaaS B2B multiempresa diseñada para digitalizar y gestionar "
        "el proceso completo de Inspección Técnica de Obras (ITO) en proyectos de construcción "
        "e inmobiliarios. La solución reemplaza los procesos manuales basados en papel, Excel "
        "y WhatsApp por un sistema integrado con trazabilidad completa, dashboards ejecutivos "
        "en tiempo real y una aplicación móvil con capacidad offline."
    )

    heading2(doc, "Problema que resuelve")
    bullets_problema = [
        "Formularios en papel o Excel sin estructura ni trazabilidad",
        "Evidencia fotográfica gestionada por WhatsApp sin vínculo al registro",
        "Observaciones y no conformidades enviadas por correo sin seguimiento formal",
        "Ausencia de KPIs de calidad y cumplimiento para la gerencia",
        "Imposibilidad de reportes automáticos de inspección",
    ]
    for b in bullets_problema:
        bullet(doc, b)

    heading2(doc, "Propuesta de valor por actor")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(11)
    table_header_row(tbl, ['Actor', 'Beneficio Principal'])
    actores = [
        ("ITO / Inspector",       "App móvil offline, registro ágil en terreno, captura de evidencias"),
        ("Supervisor de Obra",    "Trazabilidad de hallazgos y no conformidades en tiempo real"),
        ("Constructora",          "Dashboard ejecutivo, reportes automáticos PDF/Excel, KPIs"),
        ("Mandante / Cliente",    "Transparencia total y evidencia digital del proceso ITO"),
        ("Contratista",           "Seguimiento formal de las observaciones asignadas a su cuadrilla"),
    ]
    for i, (actor, beneficio) in enumerate(actores):
        row = tbl.add_row()
        row.cells[0].text = actor
        row.cells[1].text = beneficio
        set_cell_bg(row.cells[0], 'EBF3FB' if i % 2 == 0 else 'FFFFFF')
        set_cell_bg(row.cells[1], 'EBF3FB' if i % 2 == 0 else 'FFFFFF')
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    # ────────────────────────────
    # 2. STACK TECNOLÓGICO
    # ────────────────────────────
    heading1(doc, "2. Stack Tecnológico")
    add_hr(doc)

    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.columns[0].width = Cm(4.5)
    tbl2.columns[1].width = Cm(7.0)
    tbl2.columns[2].width = Cm(4.5)
    table_header_row(tbl2, ['Capa', 'Tecnología', 'Versión / Notas'])
    stack_rows = [
        ("Backend API",         "ASP.NET Core + Clean Architecture + CQRS (MediatR)",     ".NET 8"),
        ("ORM",                 "Entity Framework Core",                                   "EF Core 8"),
        ("Base de Datos",       "PostgreSQL",                                              "v16"),
        ("Frontend Web",        "Blazor Server + MudBlazor",                              ".NET 8"),
        ("App Móvil",           "Android nativo — Kotlin, Jetpack Compose, Room DB",      "SDK 34+"),
        ("Autenticación",       "ASP.NET Identity + JWT",                                 "Bearer Token"),
        ("Almacenamiento",      "MinIO (S3 compatible)",                                   "Self-hosted"),
        ("Reportes PDF",        "QuestPDF",                                                "MIT License"),
        ("Reportes Excel",      "ClosedXML",                                               "MIT License"),
        ("Logging",             "Serilog",                                                 "Sinks: consola/archivo"),
        ("Infraestructura",     "Docker + docker-compose",                                 "Local + producción"),
        ("Sesiones distribuidas","Redis",                                                  "Blazor sticky sessions"),
        ("Sync Offline",        "WorkManager (Android)",                                   "Offline-first"),
    ]
    for i, row_data in enumerate(stack_rows):
        row = tbl2.add_row()
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            set_cell_bg(cell, 'EBF3FB' if i % 2 == 0 else 'FFFFFF')
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    # ────────────────────────────
    # 3. ARQUITECTURA
    # ────────────────────────────
    heading1(doc, "3. Arquitectura de la Solución")
    add_hr(doc)

    heading2(doc, "3.1 Clean Architecture en capas")
    capas = [
        ("ITO.Cloud.Domain",          "Entidades, interfaces, enums, value objects, reglas de negocio. SIN dependencias externas."),
        ("ITO.Cloud.Application",     "Casos de uso CQRS (Commands/Queries), DTOs, validadores FluentValidation, mapeos AutoMapper."),
        ("ITO.Cloud.Infrastructure",  "EF Core + PostgreSQL, MinIO, SMTP, JWT, auditoría. Implementa las interfaces del dominio."),
        ("ITO.Cloud.API",             "ASP.NET Core Web API REST, controllers, middleware de tenant/errores, Swagger, JWT."),
        ("ITO.Cloud.Web",             "Blazor Server + MudBlazor. Páginas, layouts, servicios de consumo de la API."),
        ("ITO.Cloud.Android",         "App Android Kotlin. UI con Jetpack Compose, Room DB offline, SyncWorker WorkManager."),
    ]
    tbl3 = doc.add_table(rows=1, cols=2)
    tbl3.style = 'Table Grid'
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.columns[0].width = Cm(5.5)
    tbl3.columns[1].width = Cm(10.5)
    table_header_row(tbl3, ['Proyecto / Capa', 'Responsabilidad'])
    for i, (proj, resp) in enumerate(capas):
        row = tbl3.add_row()
        row.cells[0].text = proj
        row.cells[1].text = resp
        for c in row.cells:
            set_cell_bg(c, 'EBF3FB' if i % 2 == 0 else 'FFFFFF')
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    heading2(doc, "3.2 Multi-Tenant Lógico")
    body(doc,
        "Se implementa multi-tenancy lógico: base de datos compartida con discriminador TenantId "
        "en todas las tablas de negocio. EF Core Global Query Filters aseguran aislamiento automático "
        "de datos entre empresas (tenants)."
    )
    bullet(doc, "TenantId incluido en el JWT como claim firmado — no manipulable por el cliente")
    bullet(doc, "TenantMiddleware resuelve el tenant en cada request y lo inyecta en el contexto")
    bullet(doc, "ApplicationDbContext aplica HasQueryFilter global a todas las entidades ITenantEntity")
    bullet(doc, "Validación explícita adicional en comandos críticos (doble verificación)")

    doc.add_paragraph()

    heading2(doc, "3.3 Offline-First en App Android")
    body(doc,
        "El inspector puede trabajar sin conectividad. La arquitectura garantiza que ningún dato "
        "se pierda y que la sincronización sea automática al recuperar señal."
    )
    bullet(doc, "Room DB (SQLite local) como fuente de verdad — siempre se lee/escribe local primero")
    bullet(doc, "SyncQueue: tabla local de operaciones pendientes de sincronización")
    bullet(doc, "SyncWorker (WorkManager) envía la cola al servidor cuando hay red, con retry automático")
    bullet(doc, "PKs UUID en todas las entidades → creación offline sin colisión de IDs")
    bullet(doc, "Resolución de conflictos: estrategia server-wins con last_modified_at")
    bullet(doc, "Fotos: comprimidas en la app antes del upload, referencia temporal reemplazada al sincronizar")

    doc.add_paragraph()

    heading2(doc, "3.4 Despliegue con Docker")
    svcs = [
        ("ito-api",   "ASP.NET Core Web API",  "5000"),
        ("ito-web",   "Blazor Server",          "5001"),
        ("postgres",  "PostgreSQL 16",          "5432"),
        ("minio",     "MinIO S3",               "9000 / 9001"),
        ("redis",     "Redis (sesiones)",        "6379"),
    ]
    tbl4 = doc.add_table(rows=1, cols=3)
    tbl4.style = 'Table Grid'
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl4.columns[0].width = Cm(4)
    tbl4.columns[1].width = Cm(8)
    tbl4.columns[2].width = Cm(4)
    table_header_row(tbl4, ['Servicio', 'Descripción', 'Puerto'])
    for i, (svc, desc, port) in enumerate(svcs):
        row = tbl4.add_row()
        row.cells[0].text = svc
        row.cells[1].text = desc
        row.cells[2].text = port
        for c in row.cells:
            set_cell_bg(c, 'EBF3FB' if i % 2 == 0 else 'FFFFFF')
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    # ────────────────────────────
    # 4. MÓDULOS DEL SISTEMA
    # ────────────────────────────
    heading1(doc, "4. Módulos del Sistema")
    add_hr(doc)

    modulos = [
        ("1",  "Identity & Tenants",          "Login, sesiones, multiempresa, planes SaaS",                                    "MVP",     "✅"),
        ("2",  "Empresas & Proyectos",         "CRUD empresa, obra, etapas, sectores, unidades",                               "MVP",     "✅"),
        ("3",  "Usuarios & Permisos",          "Gestión de usuarios por tenant, roles por proyecto",                           "MVP",     "✅"),
        ("4",  "Plantillas de Inspección",     "Builder dinámico de formularios, secciones, preguntas, versiones",             "MVP",     "✅"),
        ("5",  "Programación",                 "Agenda de inspecciones, asignación a ITO, calendario",                         "MVP",     "✅"),
        ("6",  "Ejecución (Android)",          "App offline, checklist, evidencia fotográfica, sincronización",                "MVP",     "✅"),
        ("7",  "Observaciones / NC",           "Registro de no conformidades, seguimiento, estados, responsables",             "MVP",     "✅"),
        ("8",  "Reinspecciones",               "Validación de correcciones, trazabilidad completa",                            "MVP",     "✅"),
        ("9",  "Documentos",                   "Adjuntos a proyectos: planos, contratos, especificaciones",                    "MVP",     "✅"),
        ("10", "Reportes",                     "Informe PDF de inspección (QuestPDF) + Excel (ClosedXML)",                     "MVP",     "✅"),
        ("11", "Dashboard",                    "KPIs ejecutivos, gráficos de avance, estado del proyecto",                     "MVP",     "✅"),
        ("12", "Notificaciones",               "Email automático por observación/asignación, alertas internas",                "MVP",     "✅"),
        ("13", "Auditoría",                    "Log de acciones, historial de cambios, quién hizo qué y cuándo",               "MVP",     "✅"),
        ("14", "Módulo IA",                    "Interfaces preparadas: clasificación severidad, descripción automática NC",    "Futuro",  "⏳"),
    ]

    tbl5 = doc.add_table(rows=1, cols=5)
    tbl5.style = 'Table Grid'
    tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl5.columns[0].width = Cm(0.8)
    tbl5.columns[1].width = Cm(4.5)
    tbl5.columns[2].width = Cm(7.5)
    tbl5.columns[3].width = Cm(2.0)
    tbl5.columns[4].width = Cm(1.2)
    table_header_row(tbl5, ['#', 'Módulo', 'Descripción', 'Alcance', 'MVP'])

    for i, (num, mod, desc, alcance, estado) in enumerate(modulos):
        row = tbl5.add_row()
        vals = [num, mod, desc, alcance, estado]
        bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
        if alcance == 'Futuro':
            bg = 'FFF8E1'
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = val
            set_cell_bg(cell, bg)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if j == 1:
                        r.bold = True

    doc.add_paragraph()

    # ────────────────────────────
    # 5. ETAPAS DE DESARROLLO
    # ────────────────────────────
    heading1(doc, "5. Etapas de Desarrollo")
    add_hr(doc)

    etapas = [
        {
            "num": "Etapa 1",
            "titulo": "Descubrimiento y Arquitectura",
            "estado": "✅ COMPLETA",
            "estado_bg": "E8F5E9",
            "entregables": [
                "Arquitectura Clean Architecture documentada",
                "ADR-001: Decisión Blazor Server vs WASM",
                "ADR-002: Multi-tenant lógico con TenantId",
                "Estructura completa de la solución .NET",
                "Roadmap y módulos definidos",
            ]
        },
        {
            "num": "Etapa 2",
            "titulo": "Modelo de Dominio y Base de Datos",
            "estado": "🔄 SIGUIENTE",
            "estado_bg": "FFF9C4",
            "entregables": [
                "Solución .NET con todos los proyectos de capas creados",
                "Entidades de dominio completas con relaciones",
                "ApplicationDbContext con configuraciones EF Core",
                "Script SQL de creación de tablas PostgreSQL",
                "docker-compose.yml funcional con todos los servicios",
                "Primera migración EF Core aplicada",
            ]
        },
        {
            "num": "Etapa 3",
            "titulo": "Backend API — Núcleo de Negocio",
            "estado": "⏳ PENDIENTE",
            "estado_bg": "FFFFFF",
            "entregables": [
                "ASP.NET Identity + JWT configurado",
                "CRUD: Tenants, Empresas, Proyectos, Etapas, Sectores, Unidades",
                "CRUD: Usuarios, Roles, Permisos por proyecto",
                "TenantMiddleware + ExceptionMiddleware",
                "Swagger / OpenAPI documentado",
                "Tests unitarios de Application layer",
            ]
        },
        {
            "num": "Etapa 4",
            "titulo": "Frontend Web (Blazor Server)",
            "estado": "⏳ PENDIENTE",
            "estado_bg": "FFFFFF",
            "entregables": [
                "Layout base, login y gestión de sesiones",
                "Dashboard con KPIs (MudBlazor Charts)",
                "CRUD Empresas y Proyectos con UI completa",
                "Builder dinámico de Plantillas de Inspección",
                "Módulo de Programación y asignación de inspecciones",
                "Módulo de Observaciones / No Conformidades",
                "Módulo de Reportes PDF y Excel",
            ]
        },
        {
            "num": "Etapa 5",
            "titulo": "App Móvil Android",
            "estado": "⏳ PENDIENTE",
            "estado_bg": "FFFFFF",
            "entregables": [
                "Login + sincronización inicial de datos",
                "Lista de inspecciones asignadas al inspector",
                "Ejecución offline de checklist dinámico",
                "Captura y compresión de fotografías de evidencia",
                "SyncWorker: sincronización automática con WorkManager",
                "Manejo de conflictos y cola de reintentos",
            ]
        },
        {
            "num": "Etapa 6",
            "titulo": "Reportes y Dashboard Avanzado",
            "estado": "⏳ PENDIENTE",
            "estado_bg": "FFFFFF",
            "entregables": [
                "Informe de inspección en PDF con QuestPDF (logo, firmas, evidencias)",
                "Exportación Excel con ClosedXML",
                "Dashboard ejecutivo con gráficos de KPIs",
                "Módulo de auditoría y log de acciones",
                "Notificaciones por email automáticas",
            ]
        },
        {
            "num": "Etapa 7",
            "titulo": "Infraestructura y Deploy",
            "estado": "⏳ PENDIENTE",
            "estado_bg": "FFFFFF",
            "entregables": [
                "Docker-compose de producción optimizado",
                "Configuración Nginx reverse proxy",
                "CI/CD pipeline básico",
                "Estrategia de backup de PostgreSQL",
                "Monitoreo con Serilog + Seq",
                "Documentación de operaciones",
            ]
        },
        {
            "num": "Etapa 8",
            "titulo": "Módulo IA (Fase futura)",
            "estado": "⏳ FUTURO",
            "estado_bg": "FFF8E1",
            "entregables": [
                "Clasificación automática de severidad de observaciones",
                "Generación automática de descripciones de NC",
                "Resumen ejecutivo automático de inspecciones",
                "Búsqueda semántica sobre observaciones históricas",
                "Integración con Claude API (Anthropic)",
            ]
        },
    ]

    for etapa in etapas:
        # Encabezado de etapa
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run_num = p.add_run(f"{etapa['num']}: ")
        run_num.bold = True
        run_num.font.size = Pt(12)
        run_num.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)
        run_tit = p.add_run(etapa['titulo'])
        run_tit.bold = True
        run_tit.font.size = Pt(12)
        run_tit.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        run_est = p.add_run(f"   {etapa['estado']}")
        run_est.font.size = Pt(10)
        run_est.bold = False
        run_est.italic = True
        run_est.font.color.rgb = RGBColor(0x33, 0x88, 0x33) if "COMPLETA" in etapa['estado'] \
            else RGBColor(0xCC, 0x88, 0x00) if "SIGUIENTE" in etapa['estado'] \
            else RGBColor(0x66, 0x66, 0x66)

        for entregable in etapa['entregables']:
            bullet(doc, entregable)

    doc.add_paragraph()

    # ────────────────────────────
    # 6. ROADMAP SPRINTS
    # ────────────────────────────
    heading1(doc, "6. Roadmap — Sprints del MVP")
    add_hr(doc)
    body(doc,
        "Estimación para 1 equipo full-stack. Total MVP completo: 15–18 semanas.",
        italic=True
    )

    sprints = [
        {
            "sprint": "Sprint 0",
            "titulo": "Fundación",
            "duracion": "1–2 semanas",
            "tareas": [
                "Crear solución .NET con Clean Architecture",
                "Configurar PostgreSQL + EF Core + Migrations",
                "Configurar Identity + JWT",
                "Configurar MinIO",
                "docker-compose local funcional",
                "CI básico",
            ]
        },
        {
            "sprint": "Sprint 1",
            "titulo": "Núcleo de Negocio",
            "duracion": "2–3 semanas",
            "tareas": [
                "CRUD: Tenants, Empresas, Proyectos, Etapas, Sectores, Unidades",
                "CRUD: Usuarios, Roles, Permisos",
                "Blazor Web: Layout, Login, Dashboard vacío, CRUD Empresas y Proyectos",
            ]
        },
        {
            "sprint": "Sprint 2",
            "titulo": "Plantillas de Inspección",
            "duracion": "2 semanas",
            "tareas": [
                "Motor de plantillas dinámicas (secciones, preguntas, tipos)",
                "Builder de secciones y preguntas en web",
                "Versionado de plantillas",
                "Frontend: editor visual de plantillas",
            ]
        },
        {
            "sprint": "Sprint 3",
            "titulo": "Inspecciones Web",
            "duracion": "2 semanas",
            "tareas": [
                "Programación de inspecciones",
                "Asignación a ITO / inspector",
                "Registro básico de inspecciones desde web",
                "Estados y flujo de inspección",
            ]
        },
        {
            "sprint": "Sprint 4",
            "titulo": "App Android",
            "duracion": "3 semanas",
            "tareas": [
                "Login + sincronización inicial de catálogos",
                "Lista de inspecciones asignadas al inspector",
                "Ejecución offline del checklist dinámico",
                "Captura y gestión de fotografías",
                "SyncWorker: sincronización con WorkManager",
            ]
        },
        {
            "sprint": "Sprint 5",
            "titulo": "Observaciones y Reinspecciones",
            "duracion": "2 semanas",
            "tareas": [
                "Módulo completo de No Conformidades",
                "Flujo de asignación, corrección y cierre",
                "Reinspección con trazabilidad completa",
                "Historial de cambios de estado",
            ]
        },
        {
            "sprint": "Sprint 6",
            "titulo": "Reportes y Dashboard",
            "duracion": "2 semanas",
            "tareas": [
                "PDF de inspección con QuestPDF (logo, evidencias, firmas)",
                "Export Excel con ClosedXML",
                "Dashboard KPIs con gráficos MudBlazor Charts",
                "Módulo de auditoría visible en web",
            ]
        },
        {
            "sprint": "Sprint 7",
            "titulo": "Pulido MVP y Lanzamiento",
            "duracion": "1–2 semanas",
            "tareas": [
                "Notificaciones email automáticas",
                "Validaciones finales y edge cases",
                "Documentación API Swagger completa",
                "Testing crítico de seguridad multi-tenant",
                "Deploy en servidor de producción",
            ]
        },
    ]

    tbl6 = doc.add_table(rows=1, cols=4)
    tbl6.style = 'Table Grid'
    tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl6.columns[0].width = Cm(2.5)
    tbl6.columns[1].width = Cm(4.0)
    tbl6.columns[2].width = Cm(2.5)
    tbl6.columns[3].width = Cm(7.0)
    table_header_row(tbl6, ['Sprint', 'Título', 'Duración', 'Tareas Clave'])

    for i, sp in enumerate(sprints):
        row = tbl6.add_row()
        bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
        row.cells[0].text = sp['sprint']
        row.cells[1].text = sp['titulo']
        row.cells[2].text = sp['duracion']
        tareas_text = "\n".join(f"• {t}" for t in sp['tareas'])
        row.cells[3].text = tareas_text
        for c in row.cells:
            set_cell_bg(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        # Bold sprint name
        for r in row.cells[0].paragraphs[0].runs:
            r.bold = True

    doc.add_paragraph()

    # ────────────────────────────
    # 7. DECISIONES TÉCNICAS CLAVE
    # ────────────────────────────
    heading1(doc, "7. Decisiones Técnicas Clave")
    add_hr(doc)

    decisiones = [
        ("ORM",           "EF Core 8",             "Dapper",         "Productividad + migraciones + Global Query Filters para multi-tenant"),
        ("CQRS",          "MediatR",                "Servicio directo","Separación de responsabilidades, mejor testabilidad"),
        ("Validación",    "FluentValidation",       "DataAnnotations","Más expresivo, lógica separada del modelo"),
        ("Mapeo",         "AutoMapper",             "Manual",         "Reduce boilerplate en DTOs"),
        ("PDF",           "QuestPDF",               "iTextSharp",     "Licencia MIT, API moderna, mejor DX"),
        ("Excel",         "ClosedXML",              "EPPlus",         "Licencia MIT, sin restricción comercial"),
        ("Storage",       "MinIO",                  "Azure Blob",     "Auto-hosteable, S3 compatible, sin vendor lock-in"),
        ("Auth",          "ASP.NET Identity + JWT", "Keycloak",       "Simplicidad, control total, sin dependencia externa"),
        ("Logging",       "Serilog",                "ILogger nativo", "Sinks configurables: consola, archivo, Seq"),
        ("Frontend web",  "Blazor Server",          "Blazor WASM",    "SignalR nativo para tiempo real, integración directa con Identity"),
        ("Multi-tenant",  "Lógico (TenantId)",      "Schema por tenant","Simplicidad operacional, menor costo infra inicial"),
        ("PKs",           "UUID / GUID",            "int autoincrement","Necesario para generación offline en Android sin colisiones"),
    ]

    tbl7 = doc.add_table(rows=1, cols=4)
    tbl7.style = 'Table Grid'
    tbl7.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl7.columns[0].width = Cm(3.0)
    tbl7.columns[1].width = Cm(4.0)
    tbl7.columns[2].width = Cm(3.5)
    tbl7.columns[3].width = Cm(5.5)
    table_header_row(tbl7, ['Decisión', 'Opción Elegida', 'Alternativa Descartada', 'Razón'])

    for i, (dec, elegida, descartada, razon) in enumerate(decisiones):
        row = tbl7.add_row()
        row.cells[0].text = dec
        row.cells[1].text = elegida
        row.cells[2].text = descartada
        row.cells[3].text = razon
        bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
        for c in row.cells:
            set_cell_bg(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        for r in row.cells[1].paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)

    doc.add_paragraph()

    # ────────────────────────────
    # 8. RIESGOS TÉCNICOS
    # ────────────────────────────
    heading1(doc, "8. Riesgos Técnicos y Mitigación")
    add_hr(doc)

    riesgos = [
        ("Conflictos de sync offline Android",            "Alta",   "Alto",    "Server-wins + UUID + timestamps last_modified_at"),
        ("Rendimiento Blazor Server bajo alta concurrencia","Media", "Alto",    "Redis para state distribuido + sticky sessions"),
        ("Archivos multimedia grandes (fotos HD en obra)","Alta",   "Medio",   "Compresión en app antes del upload + límites en API"),
        ("Complejidad del motor de plantillas dinámicas", "Alta",   "Alto",    "Diseñar modelo de datos flexible desde Sprint 0"),
        ("Multi-tenant: query sin filtro de tenant",      "Media",  "Crítico", "EF Global Query Filters + tests de seguridad de tenant"),
        ("Migración de esquema con clientes en producción","Media", "Alto",    "Migraciones no destructivas + estrategia blue-green"),
        ("Latencia en generación de PDFs grandes",        "Baja",   "Medio",   "Generación asíncrona con cola de trabajos en background"),
    ]

    tbl8 = doc.add_table(rows=1, cols=4)
    tbl8.style = 'Table Grid'
    tbl8.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl8.columns[0].width = Cm(5.5)
    tbl8.columns[1].width = Cm(2.0)
    tbl8.columns[2].width = Cm(2.0)
    tbl8.columns[3].width = Cm(6.5)
    table_header_row(tbl8, ['Riesgo', 'Probabilidad', 'Impacto', 'Mitigación'])

    color_map = {
        "Alta": "FF5252", "Media": "FFB300", "Baja": "66BB6A",
        "Crítico": "B71C1C", "Alto": "E64A19", "Medio": "F57C00", "Bajo": "388E3C"
    }

    for i, (riesgo, prob, impacto, mitigacion) in enumerate(riesgos):
        row = tbl8.add_row()
        row.cells[0].text = riesgo
        row.cells[3].text = mitigacion
        bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[3], bg)
        for c in [row.cells[0], row.cells[3]]:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

        # Probabilidad con color
        cell_prob = row.cells[1]
        cell_prob.text = ''
        p = cell_prob.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(prob)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_bg(cell_prob, color_map.get(prob, 'FFFFFF'))
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Impacto con color
        cell_imp = row.cells[2]
        cell_imp.text = ''
        p2 = cell_imp.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(impacto)
        run2.bold = True
        run2.font.size = Pt(9)
        set_cell_bg(cell_imp, color_map.get(impacto, 'FFFFFF'))
        run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # ────────────────────────────
    # 9. PREPARACIÓN PARA IA
    # ────────────────────────────
    heading1(doc, "9. Preparación para Módulo IA (Futuro)")
    add_hr(doc)
    body(doc,
        "Se definen interfaces vacías desde el inicio para no acoplar a ningún proveedor de IA. "
        "El sistema queda listo para integrar Claude API (Anthropic) u otro LLM sin modificar "
        "la capa de dominio ni aplicación."
    )
    funciones_ia = [
        ("GenerateObservationDescriptionAsync", "Genera descripción de NC a partir del contexto visual y datos del checklist"),
        ("SummarizeInspectionAsync",            "Resume automáticamente una inspección completa para el informe ejecutivo"),
        ("ClassifySeverityAsync",               "Clasifica la severidad de una observación basado en su descripción"),
        ("SemanticSearchAsync",                 "Búsqueda semántica sobre el historial de observaciones de un proyecto"),
    ]
    tbl9 = doc.add_table(rows=1, cols=2)
    tbl9.style = 'Table Grid'
    tbl9.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl9.columns[0].width = Cm(6.5)
    tbl9.columns[1].width = Cm(9.5)
    table_header_row(tbl9, ['Función IAIService', 'Descripción'])
    for i, (fn, desc) in enumerate(funciones_ia):
        row = tbl9.add_row()
        row.cells[0].text = fn
        row.cells[1].text = desc
        bg = 'FFF8E1' if i % 2 == 0 else 'FFFDE7'
        for c in row.cells:
            set_cell_bg(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        for r in row.cells[0].paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)

    doc.add_paragraph()
    body(doc,
        "Actualmente IAIService es implementado por AIServiceStub (Infrastructure) que retorna "
        "NotImplementedException. El contenedor DI inyecta el stub hasta que se implemente la integración real.",
        italic=True
    )

    # ────────────────────────────
    # PIE DE PÁGINA
    # ────────────────────────────
    doc.add_paragraph()
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"ITO Cloud — Roadmap de Solución v1.0  |  {datetime.date.today().strftime('%d/%m/%Y')}  |  Confidencial")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    return doc


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
if __name__ == '__main__':
    output_path = r"F:\SISTEMA DE INSPECCION DE OBRA\docs\ITO_Cloud_Roadmap.docx"
    doc = build_document()
    doc.save(output_path)
    print(f"[OK] Documento generado: {output_path}")
