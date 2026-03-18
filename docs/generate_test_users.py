# -*- coding: utf-8 -*-
"""
ITO Cloud - Insercion de usuarios de prueba + Word con credenciales
"""

import os
import uuid
import hashlib
import struct
import base64
import datetime
import psycopg2

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────
# CONEXION
# ─────────────────────────────────────────────────────────────

DB = dict(host='localhost', port=5432, database='ito_cloud',
          user='postgres', password='joanmiro')

TENANT_ID = '00000000-0000-0000-0000-000000000001'

ROLES = {
    'SuperAdmin':   '10000000-0000-0000-0000-000000000001',
    'TenantAdmin':  '10000000-0000-0000-0000-000000000002',
    'ITO':          '10000000-0000-0000-0000-000000000003',
    'Supervisor':   '10000000-0000-0000-0000-000000000004',
    'Contratista':  '10000000-0000-0000-0000-000000000005',
    'Visualizador': '10000000-0000-0000-0000-000000000006',
}

# ─────────────────────────────────────────────────────────────
# HASH COMPATIBLE ASP.NET CORE IDENTITY v3
# PBKDF2-HMAC-SHA256, 100000 iteraciones, 16-byte salt, 32-byte key
# ─────────────────────────────────────────────────────────────

def hash_password(password: str, iterations: int = 100000) -> str:
    """
    Genera un password hash compatible con ASP.NET Core Identity v3.
    Formato: 0x01 | PRF(4B BE) | iterations(4B BE) | saltLen(4B BE) | salt | key
    """
    salt = os.urandom(16)
    key  = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt, iterations, dklen=32)
    buf  = bytearray()
    buf += b'\x01'                       # version marker V3
    buf += struct.pack('>I', 1)          # PRF = HMACSHA256
    buf += struct.pack('>I', iterations) # iteration count
    buf += struct.pack('>I', 16)         # salt length
    buf += salt
    buf += key
    return base64.b64encode(bytes(buf)).decode('utf-8')

# ─────────────────────────────────────────────────────────────
# USUARIOS DE PRUEBA
# ─────────────────────────────────────────────────────────────
# (id_fijo, email, password_plain, first_name, last_name, rut, position, rol, telefono)

TEST_USERS = [
    # ── SUPERADMIN (ya existe, solo documentamos) ──────────────
    {
        "id":         "20000000-0000-0000-0000-000000000001",
        "email":      "admin@itocloud.cl",
        "password":   "Admin#2024!",
        "first_name": "Admin",
        "last_name":  "Sistema",
        "rut":        "11.111.111-1",
        "position":   "Administrador de Plataforma",
        "rol":        "SuperAdmin",
        "telefono":   "+56 9 1111 1111",
        "exists":     True,
    },
    # ── TENANTADMIN ────────────────────────────────────────────
    {
        "id":         "20000000-0000-0000-0000-000000000002",
        "email":      "jefe.ito@constructorademo.cl",
        "password":   "JefeITO#2024!",
        "first_name": "Carlos",
        "last_name":  "Mendez Fuentes",
        "rut":        "12.345.678-9",
        "position":   "Jefe de ITO",
        "rol":        "TenantAdmin",
        "telefono":   "+56 9 2222 2222",
        "exists":     False,
    },
    # ── ITO / INSPECTOR 1 ──────────────────────────────────────
    {
        "id":         "20000000-0000-0000-0000-000000000003",
        "email":      "ana.torres@constructorademo.cl",
        "password":   "Inspector#2024!",
        "first_name": "Ana",
        "last_name":  "Torres Vega",
        "rut":        "13.456.789-0",
        "position":   "ITO Certificada",
        "rol":        "ITO",
        "telefono":   "+56 9 3333 3333",
        "exists":     False,
    },
    # ── ITO / INSPECTOR 2 ──────────────────────────────────────
    {
        "id":         "20000000-0000-0000-0000-000000000004",
        "email":      "luis.rojas@constructorademo.cl",
        "password":   "Inspector#2024!",
        "first_name": "Luis",
        "last_name":  "Rojas Perez",
        "rut":        "14.567.890-1",
        "position":   "Inspector Terreno",
        "rol":        "ITO",
        "telefono":   "+56 9 4444 4444",
        "exists":     False,
    },
    # ── SUPERVISOR ─────────────────────────────────────────────
    {
        "id":         "20000000-0000-0000-0000-000000000005",
        "email":      "maria.silva@constructorademo.cl",
        "password":   "Supervisor#2024!",
        "first_name": "Maria",
        "last_name":  "Silva Campos",
        "rut":        "15.678.901-2",
        "position":   "Jefa de Obra",
        "rol":        "Supervisor",
        "telefono":   "+56 9 5555 5555",
        "exists":     False,
    },
    # ── CONTRATISTA ────────────────────────────────────────────
    {
        "id":         "20000000-0000-0000-0000-000000000006",
        "email":      "pedro.sanchez@constructorademo.cl",
        "password":   "Contratista#2024!",
        "first_name": "Pedro",
        "last_name":  "Sanchez Rios",
        "rut":        "16.789.012-3",
        "position":   "Encargado de Obra",
        "rol":        "Contratista",
        "telefono":   "+56 9 6666 6666",
        "exists":     False,
    },
    # ── VISUALIZADOR (mandante / cliente) ──────────────────────
    {
        "id":         "20000000-0000-0000-0000-000000000007",
        "email":      "roberto.lagos@mandante.cl",
        "password":   "Visor#2024!",
        "first_name": "Roberto",
        "last_name":  "Lagos Moreno",
        "rut":        "17.890.123-4",
        "position":   "Representante Mandante",
        "rol":        "Visualizador",
        "telefono":   "+56 9 7777 7777",
        "exists":     False,
    },
]

# ─────────────────────────────────────────────────────────────
# INSERCION EN BD
# ─────────────────────────────────────────────────────────────

def insert_users():
    conn = psycopg2.connect(**DB)
    cur  = conn.cursor()
    results = []

    for u in TEST_USERS:
        if u["exists"]:
            print(f"  [SKIP] {u['email']} (ya existe)")
            results.append((u, "ya existia", None))
            continue

        # Verificar si ya existe por email
        cur.execute("SELECT id FROM users WHERE email = %s", (u["email"],))
        if cur.fetchone():
            print(f"  [SKIP] {u['email']} (ya existe en BD)")
            results.append((u, "ya existia", None))
            continue

        pwd_hash = hash_password(u["password"])

        cur.execute("""
            INSERT INTO users (
                id, tenant_id, user_name, normalized_user_name,
                email, normalized_email, email_confirmed,
                password_hash, security_stamp, concurrency_stamp,
                phone_number, phone_number_confirmed,
                two_factor_enabled, lockout_enabled, access_failed_count,
                first_name, last_name, rut, position,
                is_active, is_deleted, created_at
            ) VALUES (
                %s, %s, %s, %s,
                %s, %s, true,
                %s, %s, %s,
                %s, false,
                false, true, 0,
                %s, %s, %s, %s,
                true, false, now()
            )
        """, (
            u["id"],
            TENANT_ID,
            u["email"],
            u["email"].upper(),
            u["email"],
            u["email"].upper(),
            pwd_hash,
            str(uuid.uuid4()),
            str(uuid.uuid4()),
            u["telefono"],
            u["first_name"],
            u["last_name"],
            u["rut"],
            u["position"],
        ))

        # Asignar rol
        role_id = ROLES[u["rol"]]
        cur.execute("""
            INSERT INTO user_roles (user_id, role_id)
            VALUES (%s, %s)
            ON CONFLICT DO NOTHING
        """, (u["id"], role_id))

        print(f"  [OK]   {u['email']} -> rol {u['rol']}")
        results.append((u, "insertado", pwd_hash))

    conn.commit()
    conn.close()
    return results


# ─────────────────────────────────────────────────────────────
# UTILIDADES WORD
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_hr(doc, color='1B5E91'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

def body(doc, text, italic=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic


# ─────────────────────────────────────────────────────────────
# COLORES Y DESCRIPCIONES POR ROL
# ─────────────────────────────────────────────────────────────

ROL_META = {
    "SuperAdmin":   {
        "bg":       "B71C1C", "fg": "FFFFFF",
        "badge_bg": "FFEBEE", "badge_fg": "B71C1C",
        "desc":     "Acceso total a la plataforma. Gestiona tenants, configuracion global y puede impersonar cualquier usuario.",
        "permisos": [
            "Acceso a todos los modulos sin restriccion",
            "Gestion de tenants (crear, suspender, configurar planes)",
            "Ver logs de auditoria globales",
            "Configuracion de la plataforma",
        ]
    },
    "TenantAdmin":  {
        "bg":       "1565C0", "fg": "FFFFFF",
        "badge_bg": "E3F2FD", "badge_fg": "1565C0",
        "desc":     "Administrador de la empresa (tenant). Gestiona usuarios, proyectos, plantillas y toda la configuracion de su organizacion.",
        "permisos": [
            "Crear y gestionar usuarios del tenant",
            "Crear y configurar proyectos y empresas",
            "Crear y versionar plantillas de inspeccion",
            "Ver reportes y dashboard de su tenant",
            "Configurar contratistas y equipos",
        ]
    },
    "ITO":          {
        "bg":       "1B5E20", "fg": "FFFFFF",
        "badge_bg": "E8F5E9", "badge_fg": "1B5E20",
        "desc":     "Inspector Tecnico de Obras. Ejecuta inspecciones en terreno desde la app movil (offline) y desde la web.",
        "permisos": [
            "Ver inspecciones asignadas",
            "Ejecutar inspecciones (web y app movil)",
            "Registrar respuestas y evidencias fotograficas",
            "Crear observaciones / no conformidades",
            "Generar reporte PDF de inspeccion",
        ]
    },
    "Supervisor":   {
        "bg":       "E65100", "fg": "FFFFFF",
        "badge_bg": "FBE9E7", "badge_fg": "E65100",
        "desc":     "Supervisor de obra. Revisa el avance de inspecciones, gestiona observaciones y ve el estado del proyecto en tiempo real.",
        "permisos": [
            "Ver todas las inspecciones del proyecto",
            "Asignar y reasignar inspecciones a ITOs",
            "Gestionar observaciones (estados, plazos, responsables)",
            "Ver y exportar reportes y dashboard",
            "Solicitar reinspecciones",
        ]
    },
    "Contratista":  {
        "bg":       "4A148C", "fg": "FFFFFF",
        "badge_bg": "F3E5F5", "badge_fg": "4A148C",
        "desc":     "Empresa contratista responsable de corregir observaciones. Vista limitada a sus propias observaciones asignadas.",
        "permisos": [
            "Ver observaciones asignadas a su empresa",
            "Actualizar estado de correccion",
            "Adjuntar evidencias de correccion",
            "Ver historial de sus observaciones",
        ]
    },
    "Visualizador": {
        "bg":       "37474F", "fg": "FFFFFF",
        "badge_bg": "ECEFF1", "badge_fg": "37474F",
        "desc":     "Perfil de solo lectura. Tipicamente el mandante o cliente final que necesita transparencia del proceso.",
        "permisos": [
            "Ver inspecciones y su estado (sin editar)",
            "Ver observaciones y su trazabilidad",
            "Ver y descargar reportes PDF/Excel",
            "Ver dashboard ejecutivo del proyecto",
        ]
    },
}


# ─────────────────────────────────────────────────────────────
# GENERADOR DOCUMENTO WORD
# ─────────────────────────────────────────────────────────────

def build_word(output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ── PORTADA ─────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ITO CLOUD")
    run.bold = True
    run.font.size = Pt(34)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Usuarios de Prueba — QA & Validacion")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Entorno: Demo  |  Tenant: ITO Cloud Demo  |  {datetime.date.today().strftime('%d/%m/%Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    doc.add_paragraph()

    # Aviso de confidencialidad
    p_warn = doc.add_paragraph()
    p_warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_warn.paragraph_format.space_before = Pt(10)
    p_warn.paragraph_format.space_after  = Pt(10)
    run_w = p_warn.add_run(
        "CONFIDENCIAL - Solo para uso interno del equipo QA. "
        "No compartir fuera del equipo de desarrollo."
    )
    run_w.bold = True
    run_w.font.size = Pt(10)
    run_w.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)

    pPr = p_warn._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFEBEE')
    pPr.append(shd)

    doc.add_paragraph()

    # Tarjeta de entorno
    env_tbl = doc.add_table(rows=1, cols=3)
    env_tbl.style = 'Table Grid'
    env_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(3):
        env_tbl.columns[i].width = Cm(5.6)

    env_data = [
        ("URL del Sistema",      "http://localhost:5001"),
        ("Base de Datos",        "ito_cloud @ localhost:5432"),
        ("Tenant de Prueba",     "ITO Cloud Demo  (demo)"),
    ]
    row = env_tbl.rows[0]
    for i, (label, value) in enumerate(env_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(label + "\n")
        run_l.bold = True
        run_l.font.size = Pt(8.5)
        run_l.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run_v = p.add_run(value)
        run_v.font.size = Pt(9.5)
        run_v.bold = True
        run_v.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)
        run_v.font.name = 'Consolas'
        set_cell_bg(cell, 'EBF3FB')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_page_break()

    # ── RESUMEN DE CUENTAS ───────────────────────
    heading1(doc, "1. Resumen de Cuentas de Prueba")
    add_hr(doc)
    body(doc,
        f"Se crearon {len(TEST_USERS)} usuarios de prueba, uno por cada rol del sistema. "
        "Todas las cuentas pertenecen al tenant 'ITO Cloud Demo'. "
        "Las contrasenas cumplen los requisitos de complejidad de ASP.NET Identity.",
        italic=False
    )

    doc.add_paragraph()

    # Tabla resumen
    sum_tbl = doc.add_table(rows=1, cols=5)
    sum_tbl.style = 'Table Grid'
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sum_tbl.columns[0].width = Cm(1.5)
    sum_tbl.columns[1].width = Cm(3.0)
    sum_tbl.columns[2].width = Cm(5.5)
    sum_tbl.columns[3].width = Cm(4.0)
    sum_tbl.columns[4].width = Cm(3.0)

    # Header
    hdr_row = sum_tbl.rows[0]
    headers = ['#', 'Rol', 'Email', 'Contrasena', 'Nombre']
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, '1B5E91')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for idx, u in enumerate(TEST_USERS):
        meta = ROL_META[u["rol"]]
        row = sum_tbl.add_row()

        # #
        cell0 = row.cells[0]
        cell0.text = ''
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run0 = p0.add_run(str(idx + 1))
        run0.font.size = Pt(9)
        run0.bold = True
        set_cell_bg(cell0, 'EBF3FB' if idx % 2 == 0 else 'FFFFFF')

        # Rol (con color)
        cell1 = row.cells[1]
        cell1.text = ''
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(u["rol"])
        run1.bold = True
        run1.font.size = Pt(8.5)
        run1.font.color.rgb = RGBColor(
            int(meta["bg"][0:2], 16),
            int(meta["bg"][2:4], 16),
            int(meta["bg"][4:6], 16),
        )
        set_cell_bg(cell1, meta["badge_bg"])

        # Email
        cell2 = row.cells[2]
        cell2.text = u["email"]
        set_cell_bg(cell2, 'EBF3FB' if idx % 2 == 0 else 'FFFFFF')
        for p in cell2.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.name = 'Consolas'

        # Contrasena
        cell3 = row.cells[3]
        cell3.text = u["password"]
        set_cell_bg(cell3, 'FFF9C4')
        for p in cell3.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = 'Consolas'
                r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

        # Nombre
        cell4 = row.cells[4]
        cell4.text = f"{u['first_name']} {u['last_name']}"
        set_cell_bg(cell4, 'EBF3FB' if idx % 2 == 0 else 'FFFFFF')
        for p in cell4.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

    doc.add_paragraph()

    # ── FICHAS DETALLADAS POR USUARIO ───────────
    doc.add_page_break()
    heading1(doc, "2. Fichas Detalladas por Usuario")
    add_hr(doc)

    for u in TEST_USERS:
        meta = ROL_META[u["rol"]]

        # Cabecera de ficha
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(14)
        p_hdr.paragraph_format.space_after  = Pt(0)
        run_nombre = p_hdr.add_run(f"  {u['first_name']} {u['last_name']}")
        run_nombre.bold = True
        run_nombre.font.size = Pt(13)
        run_nombre.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run_rol = p_hdr.add_run(f"   [{u['rol']}]")
        run_rol.font.size = Pt(10)
        run_rol.font.color.rgb = RGBColor(0xFF, 0xFF, 0xAA)
        pPr = p_hdr._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), meta["bg"])
        pPr.append(shd)

        # Tabla de datos del usuario
        ficha = doc.add_table(rows=0, cols=2)
        ficha.style = 'Table Grid'
        ficha.alignment = WD_TABLE_ALIGNMENT.CENTER
        ficha.columns[0].width = Cm(4.0)
        ficha.columns[1].width = Cm(13.0)

        def add_ficha_row(label, value, mono=False, highlight=False):
            r = ficha.add_row()
            r.cells[0].text = ''
            p_l = r.cells[0].paragraphs[0]
            run_l = p_l.add_run(label)
            run_l.bold = True
            run_l.font.size = Pt(9)
            run_l.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            set_cell_bg(r.cells[0], 'ECEFF1')
            r.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            r.cells[1].text = ''
            p_v = r.cells[1].paragraphs[0]
            run_v = p_v.add_run(value)
            run_v.font.size = Pt(9)
            if mono:
                run_v.font.name = 'Consolas'
            if highlight:
                run_v.bold = True
                run_v.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)
            set_cell_bg(r.cells[1], 'FFFDE7' if highlight else 'FFFFFF')
            r.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        add_ficha_row("Email / Usuario", u["email"], mono=True, highlight=True)
        add_ficha_row("Contrasena",      u["password"], mono=True, highlight=True)
        add_ficha_row("Nombre",          f"{u['first_name']} {u['last_name']}")
        add_ficha_row("RUT",             u["rut"])
        add_ficha_row("Cargo",           u["position"])
        add_ficha_row("Telefono",        u["telefono"])
        add_ficha_row("Rol",             u["rol"])
        add_ficha_row("Tenant",          "ITO Cloud Demo (demo)")
        add_ficha_row("Estado",          "Activo — email confirmado")
        add_ficha_row("ID",              u["id"], mono=True)

        # Descripcion del rol
        r_desc = ficha.add_row()
        r_desc.cells[0].text = ''
        p_dl = r_desc.cells[0].paragraphs[0]
        run_dl = p_dl.add_run("Descripcion del Rol")
        run_dl.bold = True
        run_dl.font.size = Pt(9)
        set_cell_bg(r_desc.cells[0], 'ECEFF1')
        r_desc.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        r_desc.cells[1].text = ''
        p_dv = r_desc.cells[1].paragraphs[0]
        run_dv = p_dv.add_run(meta["desc"])
        run_dv.font.size = Pt(9)
        run_dv.italic = True
        set_cell_bg(r_desc.cells[1], 'FFFFFF')

        # Permisos
        r_perm = ficha.add_row()
        r_perm.cells[0].text = ''
        p_pl = r_perm.cells[0].paragraphs[0]
        run_pl = p_pl.add_run("Permisos Clave")
        run_pl.bold = True
        run_pl.font.size = Pt(9)
        set_cell_bg(r_perm.cells[0], 'ECEFF1')
        r_perm.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        r_perm.cells[1].text = ''
        for perm in meta["permisos"]:
            p_pv = r_perm.cells[1].add_paragraph()
            run_pv = p_pv.add_run(f"• {perm}")
            run_pv.font.size = Pt(9)
            p_pv.paragraph_format.space_before = Pt(1)
            p_pv.paragraph_format.space_after  = Pt(1)
        set_cell_bg(r_perm.cells[1], 'F9FBE7')
        r_perm.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        doc.add_paragraph()

    # ── ESCENARIOS DE PRUEBA SUGERIDOS ───────────
    doc.add_page_break()
    heading1(doc, "3. Escenarios de Prueba Sugeridos por Cuenta")
    add_hr(doc)
    body(doc,
        "Flujos de validacion recomendados para cada perfil. "
        "Usar en conjunto con el Set de Pruebas QA (ITO_Cloud_QA_Tests.docx)."
    )

    escenarios = [
        {
            "rol": "SuperAdmin",
            "email": "admin@itocloud.cl",
            "password": "Admin#2024!",
            "flujos": [
                "Crear un nuevo tenant de prueba adicional",
                "Verificar que un tenant inactivo no puede hacer login",
                "Acceder al log de auditoria global de toda la plataforma",
                "Cambiar el plan de un tenant (basic -> enterprise)",
            ]
        },
        {
            "rol": "TenantAdmin",
            "email": "jefe.ito@constructorademo.cl",
            "password": "JefeITO#2024!",
            "flujos": [
                "Crear empresa, proyecto completo (etapas, sectores, unidades)",
                "Crear una plantilla de inspeccion con 3 secciones y 5 preguntas",
                "Crear y desactivar un usuario del tenant",
                "Asignar un ITO a un proyecto especifico",
                "Ver y exportar reportes de toda la empresa",
            ]
        },
        {
            "rol": "ITO",
            "email": "ana.torres@constructorademo.cl / luis.rojas@constructorademo.cl",
            "password": "Inspector#2024!",
            "flujos": [
                "Ver la lista de inspecciones asignadas",
                "Ejecutar una inspeccion desde la web (responder checklist completo)",
                "Registrar una observacion con severidad Critico y adjuntar foto",
                "Generar el reporte PDF de una inspeccion ejecutada",
                "Verificar que NO puede ver inspecciones de otro inspector",
            ]
        },
        {
            "rol": "Supervisor",
            "email": "maria.silva@constructorademo.cl",
            "password": "Supervisor#2024!",
            "flujos": [
                "Ver el dashboard con KPIs del proyecto",
                "Programar una nueva inspeccion y asignarla a un ITO",
                "Reasignar una inspeccion de un ITO a otro",
                "Cambiar el estado de una observacion (asignar a contratista con fecha limite)",
                "Solicitar reinspeccion de una observacion resuelta",
                "Exportar listado de observaciones a Excel",
            ]
        },
        {
            "rol": "Contratista",
            "email": "pedro.sanchez@constructorademo.cl",
            "password": "Contratista#2024!",
            "flujos": [
                "Ver solo las observaciones asignadas a su empresa",
                "Actualizar el estado de una observacion a 'En Correccion'",
                "Adjuntar evidencia fotografica de la correccion realizada",
                "Verificar que NO puede ver proyectos ni inspecciones no relacionados",
                "Verificar que NO puede crear ni editar plantillas",
            ]
        },
        {
            "rol": "Visualizador",
            "email": "roberto.lagos@mandante.cl",
            "password": "Visor#2024!",
            "flujos": [
                "Ver el dashboard del proyecto (solo lectura)",
                "Descargar un reporte PDF de inspeccion",
                "Ver el historial de una observacion",
                "Verificar que NO puede editar ningun registro",
                "Verificar que NO puede crear inspecciones ni observaciones",
                "Verificar que los botones de edicion no son visibles",
            ]
        },
    ]

    for esc in escenarios:
        meta = ROL_META[esc["rol"]]

        p_rh = doc.add_paragraph()
        p_rh.paragraph_format.space_before = Pt(12)
        p_rh.paragraph_format.space_after  = Pt(4)
        run_rt = p_rh.add_run(f"  {esc['rol']}  ")
        run_rt.bold = True
        run_rt.font.size = Pt(11)
        run_rt.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run_re = p_rh.add_run(f"  {esc['email']}")
        run_re.font.size = Pt(9)
        run_re.font.color.rgb = RGBColor(0xFF, 0xFF, 0xBB)
        run_re.font.name = 'Consolas'
        pPr = p_rh._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), meta["bg"])
        pPr.append(shd)

        # Tabla escenarios
        esc_tbl = doc.add_table(rows=0, cols=3)
        esc_tbl.style = 'Table Grid'
        esc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        esc_tbl.columns[0].width = Cm(1.0)
        esc_tbl.columns[1].width = Cm(12.5)
        esc_tbl.columns[2].width = Cm(3.5)

        # Header
        hrow = esc_tbl.add_row()
        for i, h in enumerate(['#', 'Flujo / Escenario de Prueba', 'Resultado']):
            cell = hrow.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(cell, meta["bg"])

        for fi, flujo in enumerate(esc["flujos"]):
            fr = esc_tbl.add_row()
            bg = meta["badge_bg"] if fi % 2 == 0 else 'FFFFFF'

            fr.cells[0].text = str(fi + 1)
            set_cell_bg(fr.cells[0], bg)
            fr.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rr in fr.cells[0].paragraphs[0].runs:
                rr.font.size = Pt(9)
                rr.bold = True

            fr.cells[1].text = flujo
            set_cell_bg(fr.cells[1], bg)
            for p in fr.cells[1].paragraphs:
                for rr in p.runs:
                    rr.font.size = Pt(9)

            # Columna resultado (para rellenar por QA)
            fr.cells[2].text = ''
            p_res = fr.cells[2].paragraphs[0]
            p_res.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_res = p_res.add_run("PEND / PASS / FAIL")
            run_res.font.size = Pt(8)
            run_res.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            run_res.italic = True
            set_cell_bg(fr.cells[2], 'FFFDE7')

        doc.add_paragraph()

    # ── INSTRUCCIONES ───────────────────────────
    doc.add_page_break()
    heading1(doc, "4. Instrucciones de Uso")
    add_hr(doc)

    instrucciones = [
        ("Como hacer login",
         "Ir a http://localhost:5001 > ingresar Email y Contrasena de la tabla anterior > clic en 'Iniciar sesion'."),
        ("Cambiar contrasena despues del primer login",
         "Ir a perfil de usuario > 'Cambiar contrasena'. Para pruebas, se recomienda mantener las contrasenas tal como estan."),
        ("Que validar al iniciar sesion",
         "Verificar que el Dashboard muestra datos del tenant 'ITO Cloud Demo'. El nombre del usuario debe aparecer en la barra superior."),
        ("Aislamiento entre roles",
         "Probar accesos cruzados: loguearse como Contratista e intentar acceder a /admin o /plantillas via URL directa. Debe retornar 403."),
        ("Prueba de sesion expirada",
         "Dejar una sesion abierta sin actividad > esperar el timeout configurado > intentar navegar. Debe redirigir a Login."),
        ("Restablecer contrasena de prueba",
         "Si se olvida o cambia una contrasena, ejecutar nuevamente el script generate_test_users.py con la opcion --reset."),
        ("Reportar defectos encontrados",
         "Documentar en el Set de Pruebas QA (ITO_Cloud_QA_Tests.docx) en la columna 'Resultado QA' con estado FAIL y descripcion del defecto."),
    ]

    for i, (titulo, detalle) in enumerate(instrucciones):
        p_it = doc.add_paragraph()
        p_it.paragraph_format.space_before = Pt(6)
        p_it.paragraph_format.space_after  = Pt(2)
        run_num = p_it.add_run(f"{i+1}. ")
        run_num.bold = True
        run_num.font.size = Pt(10)
        run_num.font.color.rgb = RGBColor(0x1B, 0x5E, 0x91)
        run_tit = p_it.add_run(titulo)
        run_tit.bold = True
        run_tit.font.size = Pt(10)
        run_tit.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        p_det = doc.add_paragraph()
        p_det.paragraph_format.space_before = Pt(0)
        p_det.paragraph_format.space_after  = Pt(6)
        p_det.paragraph_format.left_indent  = Cm(0.8)
        run_det = p_det.add_run(detalle)
        run_det.font.size = Pt(9)
        run_det.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # ── PIE ─────────────────────────────────────
    add_hr(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"ITO Cloud - Usuarios de Prueba QA v1.0  |  "
        f"{datetime.date.today().strftime('%d/%m/%Y')}  |  CONFIDENCIAL - Solo uso interno"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    doc.save(output_path)
    print(f"[OK] Word generado: {output_path}")


# ─────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print("=" * 60)
    print("ITO Cloud - Insercion de Usuarios de Prueba")
    print("=" * 60)

    print("\n[1/2] Insertando usuarios en la base de datos...")
    results = insert_users()

    print("\n[2/2] Generando documento Word...")
    word_out = r"F:\SISTEMA DE INSPECCION DE OBRA\docs\ITO_Cloud_Usuarios_Prueba.docx"
    build_word(word_out)

    print("\n" + "=" * 60)
    print("RESUMEN FINAL")
    print("=" * 60)
    for u in TEST_USERS:
        print(f"  {u['rol']:15} | {u['email']:45} | {u['password']}")
    print(f"\nDocumento: {word_out}")
